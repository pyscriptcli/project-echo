import sys
import os

# Add root directory to sys.path
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import streamlit as st
import datetime
import json
import re
import subprocess
import tempfile
import time
from io import BytesIO

import docx
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor
import pandas as pd
import PyPDF2
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, ListFlowable, ListItem
import components.doc_fonts as docfonts
docfonts.ensure_pdf_fonts()
import requests
import streamlit.components.v1 as components

# Centralized DB & Components
from utils.db import get_supabase_client, fetch_echo_context
from components.sidebar import setup_page_layout
from utils.auth import require_login, get_current_user
from utils.skills import load_prompt
from utils.minutes_memory import build_style_examples, store_approved_minutes

# 1. Page Configuration (MUST be the first Streamlit command)
st.set_page_config(
    page_title="Project Echo - MoM Generator",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 2. Enforce login before rendering anything
require_login()

# 3. Render Global Navigation
setup_page_layout()

# 3. Custom CSS & Pure CSS SVG Icon Injection (Strictly No Emojis)
CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Cormorant+Garamond:ital,wght@0,500;0,600;1,500;1,600&family=Montserrat:wght@400;500;600&family=Bebas+Neue&display=swap');
html, body, [class*="css"] { font-family: 'Montserrat', sans-serif !important; }
.stApp {
    background-color: #f4f1ec;
    color: #1b1d1e;
}
.stApp > header { display: none !important; }
.block-container { padding-top: 1.5rem !important; padding-right: 2.2rem !important; padding-left: 2.2rem !important; max-width: 100% !important; }
h3 { font-family: 'Cormorant Garamond', serif !important; font-style: italic !important; font-weight: 600 !important; color: #003366 !important; letter-spacing: 0.02em; margin-bottom: 0.25rem; font-size: 1.5rem !important; }
.playfair-label { font-family: 'Cormorant Garamond', serif !important; font-style: italic !important; color: #003366 !important; font-size: 1.05rem !important; margin-bottom: 0.25rem !important; display: block; }

/* Containers & Inputs */
div[data-testid="stVerticalBlockBorderWrapper"] { background-color: #ffffff !important; border-radius: 6px !important; box-shadow: none !important; border: 1px solid rgba(0,51,102,0.12) !important; padding: 1.5rem !important; margin-bottom: 1.25rem !important; }
.stTextArea textarea, .stTextInput input, .stSelectbox select { background-color: #ffffff !important; border: 1px solid rgba(0,51,102,0.2) !important; border-radius: 6px !important; box-shadow: none !important; color: #1b1d1e !important; }
.stTextArea textarea:focus, .stTextInput input:focus, .stSelectbox select:focus { background-color: #ffffff !important; border-color: #003366 !important; }

/* Buttons General */
.stButton > button { background-color: #0c0c0e !important; color: #ffffff !important; border: 1px solid #c9ab4c !important; border-radius: 6px !important; font-family: 'Montserrat', sans-serif !important; font-weight: 600 !important; font-size: 0.68rem !important; height: 28px !important; min-height: 28px !important; padding: 0.1rem 0.5rem !important; width: auto !important; box-shadow: none !important; transition: all 0.2s ease !important; display: flex !important; align-items: center !important; justify-content: center !important; }
.stButton > button:hover { background-color: #003366 !important; border-color: #d9bc5d !important; color: #ffffff !important; box-shadow: none !important; }

/* Settings Button SVG Icon */
.stButton > button[key="card_settings_btn"] {
    width: 36px !important;
    height: 36px !important;
    border-radius: 50% !important;
    padding: 0 !important;
    margin: 0 auto !important;
}
.stButton > button[key="card_settings_btn"]::before {
    content: "";
    display: inline-block;
    width: 18px;
    height: 18px;
    background-color: currentColor;
    -webkit-mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M19.14 12.94c.04-.3.06-.61.06-.94 0-.32-.02-.64-.07-.94l2.03-1.58a.49.49 0 0 0 .12-.61l-1.92-3.32a.488.488 0 0 0-.59-.22l-2.39.96c-.5-.38-1.03-.7-1.62-.94l-.36-2.54a.484.484 0 0 0-.48-.41h-3.84c-.24 0-.43.17-.47.41l-.36 2.54c-.59.24-1.13.57-1.62.94l-2.39-.96c-.22-.08-.47 0-.59.22L2.74 8.87c-.12.21-.08.47.12.61l2.03 1.58c-.05.3-.09.63-.09.94s.02.64.07.94l-2.03 1.58a.49.49 0 0 0-.12.61l1.92 3.32c.12.22.37.29.59.22l2.39-.96c.5.38 1.03.7 1.62.94l.36 2.54c.05.24.24.41.48.41h3.84c.24 0 .44-.17.47-.41l.36-2.54c.59-.24 1.13-.56 1.62-.94l2.39.96c.22.08.47 0 .59-.22l1.92-3.32c.12-.22.07-.47-.12-.61l-2.01-1.58zM12 15.6c-1.98 0-3.6-1.62-3.6-3.6s1.62-3.6 3.6-3.6 3.6 1.62 3.6 3.6-1.62 3.6-3.6 3.6z'/%3E%3C/svg%3E") no-repeat center;
    mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M19.14 12.94c.04-.3.06-.61.06-.94 0-.32-.02-.64-.07-.94l2.03-1.58a.49.49 0 0 0 .12-.61l-1.92-3.32a.488.488 0 0 0-.59-.22l-2.39.96c-.5-.38-1.03-.7-1.62-.94l-.36-2.54a.484.484 0 0 0-.48-.41h-3.84c-.24 0-.43.17-.47.41l-.36 2.54c-.59.24-1.13.57-1.62.94l-2.39-.96c-.22-.08-.47 0-.59.22L2.74 8.87c-.12.21-.08.47.12.61l2.03 1.58c-.05.3-.09.63-.09.94s.02.64.07.94l-2.03 1.58a.49.49 0 0 0-.12.61l1.92 3.32c.12.22.37.29.59.22l2.39-.96c.5.38 1.03.7 1.62.94l.36 2.54c.05.24.24.41.48.41h3.84c.24 0 .44-.17.47-.41l.36-2.54c.59-.24 1.13-.56 1.62-.94l2.39.96c.22.08.47 0 .59-.22l1.92-3.32c.12-.22.07-.47-.12-.61l-2.01-1.58zM12 15.6c-1.98 0-3.6-1.62-3.6-3.6s1.62-3.6 3.6-3.6 3.6 1.62 3.6 3.6-1.62 3.6-3.6 3.6z'/%3E%3C/svg%3E") no-repeat center;
}

/* Save / Webhook Buttons */
.stButton > button[key="btn_save_supabase_bottom"]::before, .stButton > button[key="btn_sync_webhook"]::before {
    content: "";
    display: inline-block;
    width: 16px;
    height: 16px;
    margin-right: 6px;
    background-color: currentColor;
    -webkit-mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M17 3H5c-1.11 0-2 .9-2 2v14c0 1.1.89 2 2 2h14c1.1 0 2-.9 2-2V7l-4-4zm-5 16c-1.66 0-3-1.34-3-3s1.34-3 3-3 3 1.34 3 3-1.34 3-3 3zm3-10H5V5h10v4z'/%3E%3C/svg%3E") no-repeat center;
    mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M17 3H5c-1.11 0-2 .9-2 2v14c0 1.1.89 2 2 2h14c1.1 0 2-.9 2-2V7l-4-4zm-5 16c-1.66 0-3-1.34-3-3s1.34-3 3-3 3 1.34 3 3-1.34 3-3 3zm3-10H5V5h10v4z'/%3E%3C/svg%3E") no-repeat center;
}

/* Delete Row Button */
.stButton > button[key^="del_"] { background-color: #FDF9F9 !important; color: #B23A3A !important; border: 1px solid rgba(178, 58, 58, 0.25) !important; }
.stButton > button[key^="del_"]:hover { background-color: #B23A3A !important; color: #FFFFFF !important; }
.stButton > button[key^="del_"]::before {
    content: "";
    display: inline-block;
    width: 15px;
    height: 15px;
    margin-right: 4px;
    background-color: currentColor;
    -webkit-mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M6 19c0 1.1.9 2 2 2h8c1.1 0 2-.9 2-2V7H6v12zM19 4h-3.5l-1-1h-5l-1 1H5v2h14V4z'/%3E%3C/svg%3E") no-repeat center;
    mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M6 19c0 1.1.9 2 2 2h8c1.1 0 2-.9 2-2V7H6v12zM19 4h-3.5l-1-1h-5l-1 1H5v2h14V4z'/%3E%3C/svg%3E") no-repeat center;
}

/* Chat Styling */
.chat-container { display: flex; flex-direction: column; gap: 0.6rem; margin-top: 0.5rem; padding-bottom: 1rem; max-height: 420px; overflow-y: auto; }
.chat-ai { align-self: flex-start; background-color: #F9FAFB; border: 1px solid rgba(0,51,102,0.12); color: #1b1d1e; padding: 0.6rem 0.85rem; border-radius: 0; max-width: 92%; font-size: 0.88rem; line-height: 1.5; box-shadow: none; }
.chat-user-wrap { display: flex; justify-content: flex-end; width: 100%; margin-bottom: 0.2rem; }
.chat-user { background-color: #222222; color: #FFFFFF; padding: 0.55rem 0.95rem; border-radius: 0; max-width: 82%; font-size: 0.88rem; line-height: 1.45; box-shadow: none; }

/* Evidence & Badges */
.evidence-quote-box {
    background-color: #F9FAFB;
    border-left: 3px solid #003366;
    padding: 0.5rem 0.75rem;
    font-size: 0.82rem;
    color: #1b1d1e;
    margin: 0.4rem 0;
    font-style: italic;
    border-radius: 0;
}
.badge-confidence {
    display: inline-block;
    padding: 2px 8px;
    border-radius: 12px;
    font-size: 0.75rem;
    font-weight: 600;
    margin-left: 6px;
}
.badge-high { background-color: #DEF7EC; color: #03543F; }
.badge-medium { background-color: #FEF08A; color: #713F12; }
.badge-low { background-color: #FDE8E8; color: #9B1C1C; }

.guardrail-alert {
    background-color: #FFFBEB;
    border-left: 3px solid #F59E0B;
    padding: 0.35rem 0.6rem;
    font-size: 0.78rem;
    color: #92400E;
    margin-top: 0.3rem;
    border-radius: 0 4px 4px 0;
}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# 4. SVG Templates
COPY_ICON = '<svg viewBox="0 0 24 24" width="16" height="16" fill="currentColor" style="vertical-align: middle; margin-right: 6px;"><path d="M16 1H4c-1.1 0-2 .9-2 2v14h2V3h12V1zm3 4H8c-1.1 0-2 .9-2 2v14c0 1.1.9 2 2 2h11c1.1 0 2-.9 2-2V7c0-1.1-.9-2-2-2zm0 16H8V7h11v14z"/></svg>'

# 5. Constants & Config
DEEPSEEK_API_KEY = str(st.secrets.get("DEEPSEEK_API_KEY", "")).strip()
DEEPSEEK_CHAT_URL = "https://api.deepseek.com/chat/completions"
GROQ_API_KEY = str(st.secrets.get("GROQ_API_KEY", "")).strip()
GROQ_AUDIO_URL = "https://api.groq.com/openai/v1/audio/transcriptions"
OPENAI_API_KEY = str(st.secrets.get("OPENAI_API_KEY", "")).strip()
OPENAI_AUDIO_URL = "https://api.openai.com/v1/audio/transcriptions"
SLACK_WEBHOOK_URL = str(st.secrets.get("SLACK_WEBHOOK_URL", "")).strip()

CRD_MEMBERS = ["Sondi Tuazon", "Kristina Balajadia", "Meliza Zapata", "Dykstra Pineda", "Cedtrix Rena", "Carlo Medina", "Dave Policarpio", "Irish Rima"]
LOCATION_PRESETS = [
    "GreatWork Mega Tower 32F - Secret Room",
    "GreatWork Mega Tower 32F - Small Meeting Room",
    "GreatWork Mega Tower 24F - Meeting Room",
    "GreatWork Mega Tower 32F - Board Room",
    "GreatWork Mega Tower 32F - Co-working",
    "Online Meeting"
]
MEETING_TYPE_OPTIONS = ["Internal", "External", "Team"]

# 6. Session State Initialization
if "transcript" not in st.session_state: st.session_state["transcript"] = ""
if "df" not in st.session_state: st.session_state["df"] = pd.DataFrame(columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"])
if "other_discussions" not in st.session_state: st.session_state["other_discussions"] = ""
if "show_settings" not in st.session_state: st.session_state["show_settings"] = False
if "tokens_used" not in st.session_state: st.session_state["tokens_used"] = 0
if "last_api_call" not in st.session_state: st.session_state["last_api_call"] = None
if "selected_engine" not in st.session_state: st.session_state["selected_engine"] = "AI - DeepSeek"
if "chat_history" not in st.session_state: st.session_state["chat_history"] = []
if "meeting_date" not in st.session_state: st.session_state["meeting_date"] = datetime.date.today()
if "meeting_location" not in st.session_state: st.session_state["meeting_location"] = ""
if "meeting_type" not in st.session_state: st.session_state["meeting_type"] = "Internal"
if "meeting_client_name" not in st.session_state: st.session_state["meeting_client_name"] = ""
if "meeting_selected_crd" not in st.session_state: st.session_state["meeting_selected_crd"] = []
if "meeting_ext_attendees" not in st.session_state: st.session_state["meeting_ext_attendees"] = ""
if "meeting_prep_name" not in st.session_state: st.session_state["meeting_prep_name"] = ""
if "meeting_prep_desig" not in st.session_state: st.session_state["meeting_prep_desig"] = ""
if "meeting_conf_name" not in st.session_state: st.session_state["meeting_conf_name"] = ""
if "meeting_conf_desig" not in st.session_state: st.session_state["meeting_conf_desig"] = ""

# HITL & Enhancement States
if "user_topics_text" not in st.session_state: st.session_state["user_topics_text"] = ""
if "matched_evidence_items" not in st.session_state: st.session_state["matched_evidence_items"] = []
if "recommended_missed_points" not in st.session_state: st.session_state["recommended_missed_points"] = []
if "entity_corrections_log" not in st.session_state: st.session_state["entity_corrections_log"] = []
if "speaker_mappings" not in st.session_state: st.session_state["speaker_mappings"] = {}
# Auto-flow tracking
if "last_processed_file" not in st.session_state: st.session_state["last_processed_file"] = None
if "_topics_discovered" not in st.session_state: st.session_state["_topics_discovered"] = False
if "_auto_processing" not in st.session_state: st.session_state["_auto_processing"] = False
# Dialog recording
if "_dialog_recorded_bytes" not in st.session_state: st.session_state["_dialog_recorded_bytes"] = None
if "_dialog_record_notes" not in st.session_state: st.session_state["_dialog_record_notes"] = ""

# -------------------------------------------------------------
# Centralized State Reducer
# -------------------------------------------------------------
def set_mom_dataframe(new_df: pd.DataFrame):
    clean_df = new_df.copy()
    for col in ["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]:
        if col not in clean_df.columns: clean_df[col] = ""
    st.session_state["df"] = clean_df.reset_index(drop=True)
    st.session_state["mom_editor_rows"] = st.session_state["df"].to_dict('records')
    st.session_state["_last_df_id"] = id(st.session_state["df"])

def update_mom_field(row_idx: int, field_name: str, new_val: str):
    if 0 <= row_idx < len(st.session_state["df"]):
        st.session_state["df"].at[row_idx, field_name] = new_val
        st.session_state["mom_editor_rows"] = st.session_state["df"].to_dict('records')

def delete_mom_row(row_idx: int):
    if 0 <= row_idx < len(st.session_state["df"]):
        st.session_state["df"] = st.session_state["df"].drop(st.session_state["df"].index[row_idx]).reset_index(drop=True)
        st.session_state["mom_editor_rows"] = st.session_state["df"].to_dict('records')
        st.session_state["_last_df_id"] = id(st.session_state["df"])

def add_mom_row(dp: str = "", ap: str = "", dd: str = "", pic: str = ""):
    new_row = pd.DataFrame([{"Discussion Points": dp, "Action Plan": ap, "Indicative Delivery Date": dd, "Person-in-charge": pic}])
    updated = pd.concat([st.session_state["df"], new_row], ignore_index=True)
    set_mom_dataframe(updated)

# -------------------------------------------------------------
# Deterministic Entity & Jargon Pre-Processor
# -------------------------------------------------------------
def preprocess_transcript_entities(transcript: str):
    if not transcript: return transcript, []
    context_data = fetch_echo_context()
    # Safe default if fetch fails
    if not isinstance(context_data, dict):
        context_data = {}
    replacements = []
    
    phonetic_map = {
        r"\bcool\s*berneties\b": "Kubernetes",
        r"\bcoolbernetes\b": "Kubernetes",
        r"\bmiss\s*meli\b": "Meliza Zapata",
        r"\bsir\s*sondi\b": "Sondi Tuazon",
        r"\bced\b": "Cedtrix Rena",
        r"\bprime\s*ph\b": "PRIME Philippines",
        r"\bgreat\s*work\b": "GreatWork",
        r"\bmom\b": "MOM"
    }
    
    cleaned = transcript
    for pattern, canonical in phonetic_map.items():
        if re.search(pattern, cleaned, flags=re.IGNORECASE):
            cleaned = re.sub(pattern, canonical, cleaned, flags=re.IGNORECASE)
            replacements.append(f"Standardized entity: '{canonical}'")

    jargon = context_data.get('jargon', {})
    if not isinstance(jargon, dict):
        jargon = {}
    for k, v in jargon.items():
        if not isinstance(k, str) or not isinstance(v, str):
            continue
        p = r'\b' + re.escape(k) + r'\b'
        if re.search(p, cleaned, flags=re.IGNORECASE) and k.lower() != v.lower():
            cleaned = re.sub(p, v, cleaned, flags=re.IGNORECASE)
            replacements.append(f"Jargon match: '{k}' -> '{v}'")

    return cleaned, list(set(replacements))

# -------------------------------------------------------------
# Fact-Checking & Guardrail Verification
# -------------------------------------------------------------
def check_row_guardrails(row: dict, valid_attendees: list):
    warnings = []
    dp = str(row.get("Discussion Points", "")).strip()
    ap = str(row.get("Action Plan", "")).strip()
    dd = str(row.get("Indicative Delivery Date", "")).strip()
    pic = str(row.get("Person-in-charge", "")).strip()

    if pic and pic not in ["Unassigned", "None", "TBD", "PRIME Philippines", "Client"]:
        if valid_attendees:
            matched = any(att.lower() in pic.lower() or pic.lower() in att.lower() for att in valid_attendees)
            if not matched:
                warnings.append(f"Assignee '{pic}' not found in confirmed attendee list.")

    action_triggers = ['send', 'prepare', 'submit', 'update', 'review', 'email', 'coordinate', 'finalize', 'present', 'kailangan', 'ipapasa', 'gagawin']
    if (not ap or ap.lower() in ["none", "tbd", "n/a"]) and any(re.search(r'\b' + re.escape(w) + r'\b', dp, re.IGNORECASE) for w in action_triggers):
        warnings.append("Possible commitment detected in discussion point, but Action Plan is empty.")

    if ap and ap.lower() not in ["none", "n/a"] and (not dd or dd.lower() in ["tbd", ""]):
        warnings.append("Missing target delivery date for actionable deliverable.")

    return warnings

# -------------------------------------------------------------
# Speaker Diarization / Identity Quick-Map
# -------------------------------------------------------------
def detect_speaker_tags(transcript: str):
    matches = re.findall(r'\b(Speaker\s*[0-9]+|Speaker\s*[A-Za-z]+)\b', transcript, flags=re.IGNORECASE)
    return sorted(list(set(matches)))

def apply_speaker_remapping(transcript: str, mapping: dict):
    if not mapping or not transcript: return transcript
    updated = transcript
    for spk, name in mapping.items():
        if name and name.strip():
            updated = re.sub(r'\b' + re.escape(spk) + r'\b', name.strip(), updated, flags=re.IGNORECASE)
    return updated

# -------------------------------------------------------------
# Action Item Webhook Sync
# -------------------------------------------------------------
def dispatch_action_items_webhook(webhook_url: str, df: pd.DataFrame, meeting_details: dict):
    if not webhook_url: return False, "No Webhook URL configured."
    try:
        tasks = []
        for idx, row in df.iterrows():
            ap = str(row.get("Action Plan", "")).strip()
            if ap and ap.lower() not in ["none", "n/a"]:
                tasks.append({
                    "task_number": idx + 1,
                    "action_plan": ap,
                    "assignee": str(row.get("Person-in-charge", "Unassigned")),
                    "due_date": str(row.get("Indicative Delivery Date", "TBD")),
                    "context": str(row.get("Discussion Points", ""))
                })
        
        payload = {
            "text": f"*New Action Items Dispatched from Echo MoM*\n*Client/Project:* {meeting_details.get('company_name', 'Internal')}\n*Date:* {meeting_details.get('date', '')}",
            "meeting_id": f"MOM-{datetime.datetime.now().strftime('%Y%m%d-%H%M')}",
            "meeting_details": meeting_details,
            "tasks": tasks
        }
        res = requests.post(webhook_url, json=payload, timeout=15)
        if res.status_code in [200, 201, 204]:
            return True, f"Successfully synced {len(tasks)} action items via Webhook!"
        return False, f"Webhook returned status {res.status_code}: {res.text}"
    except Exception as e:
        return False, f"Webhook dispatch failed: {e}"

# 7. Core Logic Functions
def save_meeting_to_supabase(meeting_details, df, other_discussions, transcript):
    client = get_supabase_client()
    if not client: return False, "Supabase client uninitialized."
    try:
        table_items = [{"Discussion Points": str(row.get("Discussion Points", "")), "Action Plan": str(row.get("Action Plan", "")), "Indicative Delivery Date": str(row.get("Indicative Delivery Date", "")), "Person-in-charge": str(row.get("Person-in-charge", ""))} for _, row in df.iterrows()]
        meeting_id = f"MOM-{datetime.datetime.now().strftime('%Y%m%d-%H%M%S')}"
        client_name = meeting_details.get("company_name", "Unknown Client")
        meeting_date_str = datetime.datetime.now().strftime("%Y-%m-%d")
        if meeting_details.get("date"):
            try: meeting_date_str = datetime.datetime.strptime(meeting_details.get("date"), "%B %d, %Y").strftime("%Y-%m-%d")
            except Exception: pass
        
        payload = {
            "meeting_id": meeting_id, "client_name": client_name, "meeting_date": meeting_date_str,
            "meeting_type": meeting_details.get("meeting_type", "Internal"),
            "location": meeting_details.get("location", ""), "prepared_by": meeting_details.get("prep_name", ""),
            "confirmed_by": meeting_details.get("conf_name", ""), "summary_md": f"### Summary\n{other_discussions}",
            "transcript_md": f"### Transcript\n{transcript[:5000]}", "table_items": table_items,
            "raw_payload": {"meeting_details": meeting_details, "other_discussions": other_discussions}
        }
        client.table("meeting_archives").upsert(payload, on_conflict="meeting_id").execute()

        # Minutes memory: learn this user's preferred style from the approved output
        try:
            _user = get_current_user()
            if _user and _user.get("id"):
                store_approved_minutes(
                    user_id=_user["id"],
                    meeting_id=meeting_id,
                    approved_items=table_items,
                    other_discussions=other_discussions,
                    client_name=client_name,
                )
        except Exception as me:
            pass  # memory loss must never block the save

        return True, "Successfully saved meeting to Supabase!"
    except Exception as e:
        return False, str(e)

def extract_text_from_file(uploaded_file):
    try:
        if uploaded_file.name.endswith('.txt'): return uploaded_file.getvalue().decode("utf-8")
        elif uploaded_file.name.endswith('.pdf'):
            reader = PyPDF2.PdfReader(uploaded_file)
            return "\n".join([page.extract_text() for page in reader.pages])
        elif uploaded_file.name.endswith('.docx'):
            doc = Document(uploaded_file)
            return "\n".join([para.text for para in doc.paragraphs])
        return ""
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return ""

def _call_openai_transcribe(audio_bytes, filename="audio.mp3"):
    """Transcribe with OpenAI returning timestamped [MM:SS] segments."""
    if not OPENAI_API_KEY: return None
    headers = {"Authorization": f"Bearer {OPENAI_API_KEY}"}
    vocab_prompt = f"PRIME Philippines corporate meeting with team: {', '.join(CRD_MEMBERS)}"
    files = {
        "file": (filename, audio_bytes), 
        "model": (None, "gpt-4o-mini-transcribe"), 
        "response_format": (None, "verbose_json"),
        "prompt": (None, vocab_prompt)
    }
    try:
        resp = requests.post(OPENAI_AUDIO_URL, headers=headers, files=files, timeout=180)
        if resp.status_code != 200:
            return None
        data = resp.json()
        segments = data.get("segments")
        if segments and isinstance(segments, list):
            lines = []
            for seg in segments:
                start = seg.get("start", 0)
                text = seg.get("text", "").strip()
                if text:
                    mins, secs = int(start // 60), int(start % 60)
                    lines.append(f"[{mins:02d}:{secs:02d}] {text}")
            if lines:
                return "\n".join(lines)
        return data.get("text", "")
    except Exception: return None

def _call_groq_whisper(audio_bytes, filename="audio.mp3"):
    """Transcribe with Groq Whisper returning timestamped [MM:SS] segments."""
    headers = {"Authorization": f"Bearer {GROQ_API_KEY}"}
    vocab_prompt = f"PRIME Philippines corporate meeting with team: {', '.join(CRD_MEMBERS)}"
    files = {
        "file": (filename, audio_bytes), 
        "model": (None, "whisper-large-v3-turbo"), 
        "response_format": (None, "verbose_json"),
        "prompt": (None, vocab_prompt)
    }
    try:
        resp = requests.post(GROQ_AUDIO_URL, headers=headers, files=files, timeout=60)
        if resp.status_code != 200:
            return None
        data = resp.json()
        # Try verbose_json segments first, fall back to plain text
        segments = data.get("segments") or data.get("chunks")
        if segments and isinstance(segments, list):
            lines = []
            for seg in segments:
                start = seg.get("start", 0)
                text = seg.get("text", "").strip()
                if text:
                    mins, secs = int(start // 60), int(start % 60)
                    lines.append(f"[{mins:02d}:{secs:02d}] {text}")
            if lines:
                return "\n".join(lines)
        return data.get("text", "")
    except Exception: return None

def transcribe_audio_pipeline(audio_bytes, original_filename, progress_bar, status_placeholder):
    progress_bar.progress(10, text="Preprocessing audio container (10%)...")
    ext = os.path.splitext(original_filename)[1] or ".m4a"
    with tempfile.NamedTemporaryFile(suffix=ext, delete=False) as src:
        src.write(audio_bytes)
        src_path = src.name
    compressed_mp3 = src_path + "_compressed.mp3"
    progress_bar.progress(25, text="Compressing audio to 16kHz Mono 24k MP3 (25%)...")
    try:
        res = subprocess.run(["ffmpeg", "-y", "-threads", "1", "-i", src_path, "-vn", "-ac", "1", "-ar", "16000", "-c:a", "libmp3lame", "-b:a", "24k", compressed_mp3], stdout=subprocess.PIPE, stderr=subprocess.PIPE, text=True)
        if res.returncode != 0: return None
        comp_size_mb = os.path.getsize(compressed_mp3) / (1024 * 1024)
        progress_bar.progress(45, text="Evaluating audio duration & routing (45%)...")
        if comp_size_mb <= 10.0 and GROQ_API_KEY:
            status_placeholder.info("Processing via Groq Whisper Primary...")
            progress_bar.progress(70, text="Transcribing via Groq Whisper (70%)...")
            with open(compressed_mp3, "rb") as f: text = _call_groq_whisper(f.read(), "audio.mp3")
            if text:
                progress_bar.progress(100, text="Transcription completed (100%)!")
                status_placeholder.empty()
                return text
        status_placeholder.info("Processing recording via OpenAI...")
        progress_bar.progress(55, text="Preparing audio segments for OpenAI (55%)...")
        segment_pattern = src_path + "_seg_%03d.mp3"
        subprocess.run(["ffmpeg", "-y", "-i", compressed_mp3, "-f", "segment", "-segment_time", "600", "-c", "copy", segment_pattern], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        seg_dir = os.path.dirname(src_path)
        base_name = os.path.basename(src_path) + "_seg_"
        segments = sorted([os.path.join(seg_dir, f) for f in os.listdir(seg_dir) if f.startswith(base_name)])
        full_transcript = []
        for idx, seg in enumerate(segments):
            pct = int(55 + ((idx + 1) / len(segments)) * 40)
            progress_bar.progress(pct, text=f"Transcribing segment {idx + 1} of {len(segments)} ({pct}%)...")
            with open(seg, "rb") as f:
                t = _call_openai_transcribe(f.read(), f"part_{idx}.mp3")
                if t: full_transcript.append(t)
            time.sleep(0.2)
            try: os.remove(seg)
            except Exception: pass
        progress_bar.progress(100, text="Transcription completed successfully (100%)!")
        time.sleep(0.3)
        status_placeholder.empty()
        return " ".join(full_transcript)
    except Exception:
        return None
    finally:
        for path in [src_path, compressed_mp3]:
            if os.path.exists(path):
                try: os.remove(path)
                except Exception: pass

def extract_metadata_with_deepseek(transcript):
    if not DEEPSEEK_API_KEY: return None
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    system_prompt = load_prompt("meeting_metadata")
    user_prompt = f"""Extract metadata from this transcript into valid JSON:
Schema: {{"meeting_type": "Internal, External, or Team", "client_name": "Company/Client name or empty string", "location": "Meeting location preset or custom name or empty string", "crd_attendees": ["Exact matching names from CRD member list"], "external_attendees": "Comma-separated list of external attendee names", "prepared_by": "Name of attendee from PRIME taking notes or empty string", "confirmed_by": "Primary external attendee/client rep or empty string"}}
Transcript: {transcript[:15000]}"""
    payload = {"model": "deepseek-chat", "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], "response_format": {"type": "json_object"}, "temperature": 0.1, "max_tokens": 500}
    try:
        resp = requests.post(DEEPSEEK_CHAT_URL, headers=headers, json=payload, timeout=45)
        if resp.status_code == 200:
            raw_text = resp.json()["choices"][0]["message"]["content"].strip()
            clean_text = re.sub(r"^```(?:json)?\s*", "", raw_text)
            clean_text = re.sub(r"\s*```$", "", clean_text).strip()
            return json.loads(clean_text)
    except Exception: pass
    return None

# -------------------------------------------------------------
# HITL LOGIC: User Topics -> Match Evidence -> Human Approval
# -------------------------------------------------------------
def suggest_discussion_topics_from_transcript(transcript):
    if not DEEPSEEK_API_KEY:
        return "1. Project Status & Progress\n2. Key Deliverables & Timelines\n3. Client Alignment & Action Items"
    
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    system_prompt = load_prompt("topic_extractor")
    user_prompt = f"""Extract 4 to 7 key distinct discussion topics discussed in this transcript as valid JSON.
The transcript has [MM:SS] timestamps — use them to identify meeting flow and group consecutive segments into coherent topics.
Look for natural transitions (e.g. "Let's move to...", "Next item...", "Now for...") as topic boundaries.
Order topics chronologically by when they first appear in the timestamps.

Schema: {{"topics": ["Topic 1 title", "Topic 2 title", "Topic 3 title"]}}
Transcript: {transcript[:20000]}"""
    payload = {"model": "deepseek-chat", "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], "response_format": {"type": "json_object"}, "temperature": 0.1, "max_tokens": 500}
    try:
        resp = requests.post(DEEPSEEK_CHAT_URL, headers=headers, json=payload, timeout=45)
        if resp.status_code == 200:
            raw = resp.json()["choices"][0]["message"]["content"].strip()
            clean = re.sub(r"^```(?:json)?\s*", "", raw)
            clean = re.sub(r"\s*```$", "", clean).strip()
            data = json.loads(clean)
            topics = data.get("topics", [])
            return "\n".join([f"{i+1}. {t}" for i, t in enumerate(topics)])
    except Exception: pass
    return "1. Project Updates\n2. Technical Implementation\n3. Timeline & Target Deadlines\n4. Resource Allocation & Next Steps"

def match_evidence_and_synthesize(transcript, user_topics_str):
    # Safe retrieval of context data
    context_data = fetch_echo_context()
    if not isinstance(context_data, dict):
        context_data = {}
    
    # Extract and sanitize team, jargon, projects
    team = context_data.get('team', [])
    if not isinstance(team, list):
        team = []
    team_list = ", ".join([str(t) for t in team if t])
    
    jargon = context_data.get('jargon', {})
    if not isinstance(jargon, dict):
        jargon = {}
    jargon_lines = []
    for k, v in jargon.items():
        if isinstance(k, str) and isinstance(v, str):
            jargon_lines.append(f"- {k}: {v}")
    jargon_list = "\n".join(jargon_lines)
    
    projects = context_data.get('projects', [])
    if not isinstance(projects, list):
        projects = []
    projects_list = ", ".join([str(p) for p in projects if p])

    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}

    # Style examples from approved past minutes (minutes memory / few-shot learning)
    _uid = get_current_user().get("id") if get_current_user() else None
    style_examples = build_style_examples(_uid, limit=3) if _uid else ""
    st.session_state["mm_style_examples"] = style_examples

    knowledge_section = (
        f"Source Knowledge Base:\nTeam: {team_list}\nProjects: {projects_list}\nJargon:\n{jargon_list}"
        if (team_list or projects_list or jargon_list) else ""
    )
    system_prompt = load_prompt(
        "minutes_generator",
        memory_examples=style_examples,
        knowledge_base=knowledge_section,
    )

    user_prompt = f"""Match and synthesize evidence for each of the following user discussion points using the meeting transcript.
The transcript has [MM:SS] timestamps — use them to:
1. Identify meeting flow and group relevant evidence to the correct discussion topics
2. Note who is speaking when the evidence_quote occurs (look at the speaker tag or surrounding context)
3. Return evidence quotes WITH their [MM:SS] prefix so the user can locate them

In the evidence_quote field, include the timestamp prefix (e.g. "[01:23] The text...") for traceability.

USER DISCUSSION TOPICS:
{user_topics_str}

MEETING TRANSCRIPT:
{transcript[:28000]}

ADDITIONAL CONTEXT — User's Meeting Notes (pre-meeting agenda, notes, observations):
{st.session_state.get("user_notes", "")[:5000]}

Format output strictly as JSON matching this schema:
{{
  "matched_items": [
    {{
      "topic_title": "Original or refined point title",
      "discussion_point": "Polished high-level corporate synthesized summary",
      "evidence_quote": "Exact 1-2 sentence verbatim quote or excerpt from the transcript",
      "action_plan": "Specific executable deliverable, or 'None' if purely informational",
      "indicative_delivery_date": "Target date, time, or 'TBD'",
      "person_in_charge": "Designated individual (from attendees/CRD team) or 'Unassigned'",
      "confidence": "High, Medium, or Low"
    }}
  ],
  "recommended_missed_points": [
    {{
      "topic_title": "Distinct topic found in the transcript but NOT in the user's input",
      "evidence_quote": "Exact 1-2 sentence verbatim quote supporting this topic"
    }}
  ],
  "other_discussions": "Concise summary of peripheral matters, warm-up banter, or general context"
}}"""

    payload = {"model": "deepseek-chat", "messages": [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_prompt}], "response_format": {"type": "json_object"}, "temperature": 0.1, "max_tokens": 2000}
    try:
        resp = requests.post(DEEPSEEK_CHAT_URL, headers=headers, json=payload, timeout=120)
        if resp.status_code == 200:
            res_json = resp.json()
            st.session_state["tokens_used"] += res_json.get("usage", {}).get("total_tokens", len(transcript) // 4)
            st.session_state["last_api_call"] = datetime.datetime.now()
            raw = res_json["choices"][0]["message"]["content"].strip()
            clean = re.sub(r"^```(?:json)?\s*", "", raw)
            clean = re.sub(r"\s*```$", "", clean).strip()
            data = json.loads(clean)
            recs = data.get("recommended_missed_points", []) or []
            if isinstance(recs, list):
                st.session_state["recommended_missed_points"] = recs
            else:
                st.session_state["recommended_missed_points"] = []
            return data.get("matched_items", []), data.get("other_discussions", "")
    except Exception: pass
    st.session_state["recommended_missed_points"] = []
    return [], ""

# -------------------------------------------------------------
# Bidirectional "Ask Echo" Table Mutation
# -------------------------------------------------------------
def ask_deepseek_with_mutation(transcript: str, question: str, chat_history: list, current_df: pd.DataFrame):
    if not DEEPSEEK_API_KEY: return "DeepSeek API key is missing. Please check your configuration.", None
    
    headers = {"Authorization": f"Bearer {DEEPSEEK_API_KEY}", "Content-Type": "application/json"}
    table_json = current_df.to_json(orient="records")

    system_prompt = load_prompt("ask_echo")

    messages = [{"role": "system", "content": system_prompt}]
    for msg in chat_history[-4:]: messages.append({"role": msg["role"], "content": msg["content"]})
    messages.append({
        "role": "user", 
        "content": f"CURRENT MOM TABLE:\n{table_json}\n\nTRANSCRIPT CONTEXT:\n{transcript[:20000]}\n\nUSER REQUEST: {question}"
    })
    
    payload = {"model": "deepseek-chat", "messages": messages, "response_format": {"type": "json_object"}, "temperature": 0.1, "max_tokens": 800}
    try:
        resp = requests.post(DEEPSEEK_CHAT_URL, headers=headers, json=payload, timeout=60)
        if resp.status_code == 200:
            res_json = resp.json()
            st.session_state["tokens_used"] += res_json.get("usage", {}).get("total_tokens", 0)
            st.session_state["last_api_call"] = datetime.datetime.now()
            raw = res_json["choices"][0]["message"]["content"].strip()
            clean = re.sub(r"^```(?:json)?\s*", "", raw)
            clean = re.sub(r"\s*```$", "", clean).strip()
            data = json.loads(clean)
            return data.get("reply", "Understood."), data.get("action", None)
        return f"Service notice ({resp.status_code}): {resp.text}", None
    except Exception as e:
        return f"Connection error: {e}", None

def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

# -------------------------------------------------------------
# TEMPLATE 1: Standard Corporate (Combined Table)
# -------------------------------------------------------------
def export_to_word_template_1(df, meeting_details, other_discussions):
    template_files = ["MOM_Template.docx", "MOM Template.docx"]
    template_path = next((f for f in template_files if os.path.exists(f)), None)
    doc = Document(template_path) if template_path else Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("MINUTES OF THE MEETING")
    r_title.bold = True
    r_title.underline = True
    r_title.font.name = docfonts.WORD_BODY
    r_title.font.size = Pt(11)

    company_target = meeting_details.get("external_attendees", [])
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "CLIENT"
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run(f"PRIME PHILIPPINES & {primary_client_rep.upper()}")
    r_sub.bold = True
    r_sub.font.name = docfonts.WORD_BODY
    r_sub.font.size = Pt(11)

    date_str = meeting_details.get("date", "____________")
    time_str = meeting_details.get("time_range", "")
    full_date = f"Date: {date_str}" + (f", {time_str}" if time_str.strip() else "")
    
    p_date = doc.add_paragraph(full_date)
    p_date.paragraph_format.space_after = Pt(2)
    for r in p_date.runs: r.font.name, r.font.size = docfonts.WORD_BODY, Pt(10)

    p_loc = doc.add_paragraph(f"Location: {meeting_details.get('location', '____________')}")
    p_loc.paragraph_format.space_after = Pt(2)
    for r in p_loc.runs: r.font.name, r.font.size = docfonts.WORD_BODY, Pt(10)

    prime_atts = meeting_details.get("prime_attendees", [])
    ext_atts = meeting_details.get("external_attendees", [])
    
    p_att = doc.add_paragraph()
    p_att.paragraph_format.space_after = Pt(2)
    p_att.paragraph_format.tab_stops.add_tab_stop(Inches(1.35), WD_TAB_ALIGNMENT.LEFT)
    r_att_label = p_att.add_run("Attended by:")
    r_att_label.font.name, r_att_label.font.size = docfonts.WORD_BODY, Pt(10)
    
    first_attendee = True
    for att in ext_atts:
        if not att.strip(): continue
        p = p_att if first_attendee else doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if not first_attendee: p.paragraph_format.left_indent = Inches(1.35)
        else: p.add_run("\t")
        comp_label = f", {meeting_details.get('company_name')}" if meeting_details.get('company_name') else ""
        r = p.add_run(f"{att}{comp_label}")
        r.font.name, r.font.size = docfonts.WORD_BODY, Pt(10)
        first_attendee = False

    for att in prime_atts:
        p = p_att if first_attendee else doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if not first_attendee: p.paragraph_format.left_indent = Inches(1.35)
        else: p.add_run("\t")
        r = p.add_run(f"{att} – PRIME Philippines")
        r.font.name, r.font.size = docfonts.WORD_BODY, Pt(10)
        first_attendee = False

    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(6)
    r_line = p_line.add_run("_________________________________________________________________________________")
    r_line.font.name, r_line.font.color.rgb = docfonts.WORD_BODY, RGBColor(105, 114, 125)

    client_display = meeting_details.get('company_name', '').strip() or "the Client"
    p_intro = doc.add_paragraph(f"During the meeting held last {date_str}, PRIME Philippines, represented by the attendee/s shown above, met with {client_display} to discuss opportunities for collaboration.")
    p_intro.paragraph_format.space_after = Pt(10)
    for r in p_intro.runs: r.font.name, r.font.size = docfonts.WORD_BODY, Pt(9.5)

    table = doc.add_table(rows=len(df)+1, cols=4)
    table.alignment, table.style, table.autofit, table.allow_autofit = WD_TABLE_ALIGNMENT.CENTER, "Table Grid", False, False
    col_widths = [Inches(2.5), Inches(2.2), Inches(1.1), Inches(1.2)]
    headers = ["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width, cell.text = col_widths[i], header
        set_cell_shading(cell, "C9AB4C")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs: p.runs[0].font.bold, p.runs[0].font.size, p.runs[0].font.name = True, Pt(9), docfonts.WORD_BODY

    for i, row in df.iterrows():
        cells = table.rows[i+1].cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = f"{i+1}. {str(row.get('Discussion Points', ''))}", str(row.get("Action Plan", "")), str(row.get("Indicative Delivery Date", "")), str(row.get("Person-in-charge", ""))
        for c_idx, cell in enumerate(cells):
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            if c_idx in [2, 3]: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs: p.runs[0].font.size, p.runs[0].font.name = Pt(8.5), docfonts.WORD_BODY

    doc.add_paragraph()
    p_note = doc.add_paragraph("*Note: The indicative delivery date serves as reference point and still subject to changes. Furthermore, it depends on the progress of both parties.")
    p_note.paragraph_format.space_after = Pt(8)
    p_note.runs[0].italic, p_note.runs[0].font.name, p_note.runs[0].font.size = True, docfonts.WORD_BODY, Pt(8)

    if other_discussions.strip():
        p_od_head = doc.add_paragraph()
        p_od_head.paragraph_format.space_before, p_od_head.paragraph_format.space_after = Pt(6), Pt(4)
        r_od_head = p_od_head.add_run("Other Discussions:")
        r_od_head.bold, r_od_head.font.size, r_od_head.font.name = True, Pt(10), docfonts.WORD_BODY
        p_od = doc.add_paragraph(other_discussions)
        p_od.paragraph_format.space_after = Pt(12)
        for r in p_od.runs: r.font.name, r.font.size = docfonts.WORD_BODY, Pt(9.5)

    p_prep_label = doc.add_paragraph("Prepared by:")
    p_prep_label.paragraph_format.space_before = Pt(12)
    p_prep_label.paragraph_format.space_after = Pt(2)
    p_prep_label.runs[0].font.name, p_prep_label.runs[0].font.bold, p_prep_label.runs[0].font.size = docfonts.WORD_BODY, True, Pt(9.5)
    p_prep_line = doc.add_paragraph("_______________________________")
    p_prep_line.paragraph_format.space_after = Pt(2)
    p_prep_line.runs[0].font.name = docfonts.WORD_BODY
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"
    prep_desig = meeting_details.get("prep_desig", "").strip() or "PRIME Philippines"
    p_prep_info = doc.add_paragraph(f"{prep_name}\n{prep_desig}")
    p_prep_info.paragraph_format.space_after = Pt(12)
    for r in p_prep_info.runs: r.font.name, r.font.size = docfonts.WORD_BODY, Pt(9.5)

    p_conf_label = doc.add_paragraph("Confirmed by:")
    p_conf_label.paragraph_format.space_after = Pt(2)
    p_conf_label.runs[0].font.name, p_conf_label.runs[0].font.bold, p_conf_label.runs[0].font.size = docfonts.WORD_BODY, True, Pt(9.5)
    p_conf_line = doc.add_paragraph("_______________________________")
    p_conf_line.paragraph_format.space_after = Pt(2)
    p_conf_line.runs[0].font.name = docfonts.WORD_BODY
    conf_name = meeting_details.get("conf_name", "").strip() or (ext_atts[0] if ext_atts else "____________________")
    conf_desig = meeting_details.get("conf_desig", "").strip() or (meeting_details.get("company_name", "").strip() or "Client")
    p_conf_info = doc.add_paragraph(f"{conf_name}\n{conf_desig}")
    p_conf_info.paragraph_format.space_after = Pt(6)
    for r in p_conf_info.runs: r.font.name, r.font.size = docfonts.WORD_BODY, Pt(9.5)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_to_pdf_template_1(df, meeting_details, other_discussions):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.6 * inch, rightMargin=0.6 * inch, topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    story, styles = [], getSampleStyleSheet()
    
    style_title = ParagraphStyle('DocTitle', parent=styles['Normal'], fontName=docfonts.FONT_HEAD_BOLD, fontSize=11.5, alignment=1, spaceAfter=2)
    company_target = meeting_details.get("external_attendees", [])
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "CLIENT"
    style_subtitle = ParagraphStyle('DocSubTitle', parent=styles['Normal'], fontName=docfonts.FONT_HEAD_BOLD, fontSize=10.5, alignment=1, spaceAfter=10)
    style_body = ParagraphStyle('DocBody', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=9, leading=12, spaceAfter=3)
    style_th = ParagraphStyle('TableHead', parent=styles['Normal'], fontName=docfonts.FONT_HEAD_BOLD, fontSize=8.5, leading=10, alignment=1)
    style_td = ParagraphStyle('TableData', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=8, leading=10)
    style_td_center = ParagraphStyle('TableDataCenter', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=8, leading=10, alignment=1)

    story.append(Paragraph("<u>MINUTES OF THE MEETING</u>", style_title))
    story.append(Paragraph(f"PRIME PHILIPPINES & {primary_client_rep.upper()}", style_subtitle))
    
    date_str, time_str = meeting_details.get("date", "____________"), meeting_details.get("time_range", "")
    full_date = f"<b>Date:</b> {date_str}" + (f", {time_str}" if time_str.strip() else "")
    story.append(Paragraph(full_date, style_body))
    story.append(Paragraph(f"<b>Location:</b> {meeting_details.get('location', '____________')}", style_body))
    
    prime_atts, ext_atts = meeting_details.get("prime_attendees", []), meeting_details.get("external_attendees", [])
    att_list = []
    for att in ext_atts:
        if att.strip(): att_list.append(f"{att}{f', {meeting_details.get('company_name')}' if meeting_details.get('company_name') else ''}")
    for att in prime_atts: att_list.append(f"{att} – PRIME Philippines")
    
    if att_list:
        story.append(Paragraph(f"<b>Attended by:</b>&nbsp;&nbsp;&nbsp;&nbsp;{att_list[0]}", style_body))
        for a in att_list[1:]: story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{a}", style_body))
    
    story.append(Spacer(1, 4))
    client_display = meeting_details.get('company_name', '').strip() or "the Client"
    story.append(Paragraph(f"During the meeting held last {date_str}, PRIME Philippines, represented by the attendee/s shown above, met with {client_display} to discuss opportunities for collaboration.", style_body))
    story.append(Spacer(1, 6))
    
    table_data = [[Paragraph("<b>Discussion Points</b>", style_th), Paragraph("<b>Action Plan</b>", style_th), Paragraph("<b>Indicative Delivery Date</b>", style_th), Paragraph("<b>Person-in-charge</b>", style_th)]]
    for i, row in df.iterrows():
        table_data.append([
            Paragraph(f"{i+1}. {str(row.get('Discussion Points', ''))}", style_td),
            Paragraph(str(row.get("Action Plan", "")), style_td),
            Paragraph(str(row.get("Indicative Delivery Date", "")), style_td_center),
            Paragraph(str(row.get("Person-in-charge", "")), style_td_center)
        ])
    
    t = Table(table_data, colWidths=[2.4 * inch, 2.3 * inch, 1.1 * inch, 1.0 * inch], repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#C9AB4C')),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#69727d')),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4)
    ]))
    story.append(t)
    
    note_style = ParagraphStyle('Note', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=7.5, leading=9, spaceBefore=4)
    story.append(Paragraph("*Note: The indicative delivery date serves as reference point and still subject to changes. Furthermore, it depends on the progress of both parties.", note_style))
    
    if other_discussions.strip():
        story.append(Spacer(1, 6))
        story.append(Paragraph("<b>Other Discussions:</b>", style_body))
        story.append(Paragraph(other_discussions, style_body))
    
    story.append(Spacer(1, 8))
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"
    prep_desig = meeting_details.get("prep_desig", "").strip() or "PRIME Philippines"
    conf_name = meeting_details.get("conf_name", "").strip() or (ext_atts[0] if ext_atts else "____________________")
    conf_desig = meeting_details.get("conf_desig", "").strip() or (meeting_details.get("company_name", "").strip() or "Client")
    
    sign_data = [
        [Paragraph("<b>Prepared by:</b>", style_body), Paragraph("<b>Confirmed by:</b>", style_body)],
        [Paragraph("_______________________________", style_body), Paragraph("_______________________________", style_body)],
        [Paragraph(f"{prep_name}<br/>{prep_desig}", style_body), Paragraph(f"{conf_name}<br/>{conf_desig}", style_body)]
    ]
    sign_table = Table(sign_data, colWidths=[3.4 * inch, 3.4 * inch])
    sign_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2)
    ]))
    story.append(sign_table)
    
    doc.build(story)
    buffer.seek(0)
    return buffer

# -------------------------------------------------------------
# TEMPLATE 2: Detailed General Meeting (Vertical Layout)
# -------------------------------------------------------------
def export_to_word_template_2(df, meeting_details, other_discussions):
    doc = Document()
    
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("MINUTES OF THE MEETING")
    r_title.bold = True
    r_title.font.name = docfonts.WORD_BODY
    r_title.font.size = Pt(16)
    
    company_target = meeting_details.get("external_attendees", [])
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "General Meeting"
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(16)
    r_sub = p_sub.add_run(f"Project / Client: {primary_client_rep}")
    r_sub.font.name = docfonts.WORD_BODY
    r_sub.font.size = Pt(11)
    r_sub.font.color.rgb = RGBColor(105, 114, 125)

    doc.add_heading('Meeting Details', level=2)
    details_table = doc.add_table(rows=5, cols=2)
    details_table.style = 'Table Grid'
    
    date_str = meeting_details.get("date", "____________")
    time_str = meeting_details.get("time_range", "____________")
    location_str = meeting_details.get('location', '____________')
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"
    
    details_map = [
        ("Date", date_str),
        ("Time", time_str),
        ("Venue", location_str),
        ("Prepared by", prep_name),
        ("Date prepared", datetime.datetime.now().strftime("%B %d, %Y"))
    ]
    
    for i, (key, val) in enumerate(details_map):
        cells = details_table.rows[i].cells
        cells[0].text = key
        cells[1].text = val
        cells[0].paragraphs[0].runs[0].font.bold = True
        for cell in cells:
            cell.paragraphs[0].runs[0].font.name = docfonts.WORD_BODY
            cell.paragraphs[0].runs[0].font.size = Pt(9.5)

    doc.add_paragraph()

    doc.add_heading('Attendees', level=2)
    all_atts = meeting_details.get("prime_attendees", []) + meeting_details.get("external_attendees", [])
    for att in all_atts:
        if att.strip():
            p_att = doc.add_paragraph(style='List Bullet')
            r_att = p_att.add_run(att.strip())
            r_att.font.name = docfonts.WORD_BODY
            r_att.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading('Purpose & Summary', level=2)
    p_purp = doc.add_paragraph(other_discussions if other_discussions.strip() else "To discuss project updates, ongoing deliverables, and establish clear action plans.")
    p_purp.paragraph_format.space_after = Pt(12)
    for r in p_purp.runs: r.font.name, r.font.size = docfonts.WORD_BODY, Pt(10)

    doc.add_heading('Discussion Points', level=2)
    for i, row in df.iterrows():
        p_dp = doc.add_paragraph(style='List Number')
        r_dp = p_dp.add_run(str(row.get('Discussion Points', '')))
        r_dp.font.name = docfonts.WORD_BODY
        r_dp.font.size = Pt(10)

    doc.add_paragraph()

    doc.add_heading('Action Plan', level=2)
    act_table = doc.add_table(rows=len(df)+1, cols=4)
    act_table.style = 'Table Grid'
    act_headers = ["#", "Action Plan", "Owner", "Deadline"]
    col_widths = [Inches(0.5), Inches(3.5), Inches(1.5), Inches(1.5)]
    
    for i, header in enumerate(act_headers):
        cell = act_table.rows[0].cells[i]
        cell.width, cell.text = col_widths[i], header
        set_cell_shading(cell, "003366")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs: 
            p.runs[0].font.bold = True
            p.runs[0].font.color.rgb = RGBColor(255, 255, 255)
            p.runs[0].font.size, p.runs[0].font.name = Pt(9), docfonts.WORD_BODY

    for i, row in df.iterrows():
        cells = act_table.rows[i+1].cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = str(i+1), str(row.get("Action Plan", "")), str(row.get("Person-in-charge", "")), str(row.get("Indicative Delivery Date", ""))
        for c_idx, cell in enumerate(cells):
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            if c_idx in [0, 2, 3]: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs: p.runs[0].font.size, p.runs[0].font.name = Pt(9), docfonts.WORD_BODY

    doc.add_paragraph()

    p_footer = doc.add_paragraph()
    p_footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_footer.paragraph_format.space_before = Pt(24)
    r_footer = p_footer.add_run(f"Prepared for circulation to {primary_client_rep}. Please return corrections before this is treated as the agreed record.")
    r_footer.italic = True
    r_footer.font.color.rgb = RGBColor(105, 114, 125)
    r_footer.font.name = docfonts.WORD_BODY
    r_footer.font.size = Pt(8)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_to_pdf_template_2(df, meeting_details, other_discussions):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.8 * inch, rightMargin=0.8 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    story, styles = [], getSampleStyleSheet()
    
    style_title = ParagraphStyle('Title2', parent=styles['Normal'], fontName=docfonts.FONT_HEAD_BOLD, fontSize=15, alignment=1, spaceAfter=2)
    style_subtitle = ParagraphStyle('SubTitle2', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=10.5, textColor=colors.HexColor('#69727d'), alignment=1, spaceAfter=20)
    style_h2 = ParagraphStyle('Heading2', parent=styles['Normal'], fontName=docfonts.FONT_HEAD_BOLD, fontSize=12, textColor=colors.HexColor("#003366"), spaceBefore=12, spaceAfter=6)
    style_body = ParagraphStyle('Body2', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=9.5, leading=14, spaceAfter=4)
    style_th = ParagraphStyle('TH2', parent=styles['Normal'], fontName=docfonts.FONT_HEAD_BOLD, fontSize=9, textColor=colors.white, alignment=1)
    style_td = ParagraphStyle('TD2', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=9, leading=12)
    style_td_center = ParagraphStyle('TDC2', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=9, leading=12, alignment=1)
    
    company_target = meeting_details.get("external_attendees", [])
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "General Meeting"

    story.append(Paragraph("MINUTES OF THE MEETING", style_title))
    story.append(Paragraph(f"Project / Client: {primary_client_rep}", style_subtitle))
    
    story.append(Paragraph("Meeting Details", style_h2))
    date_str = meeting_details.get("date", "____________")
    time_str = meeting_details.get("time_range", "____________")
    location_str = meeting_details.get('location', '____________')
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"
    
    details_data = [
        [Paragraph("<b>Date</b>", style_body), Paragraph(date_str, style_body)],
        [Paragraph("<b>Time</b>", style_body), Paragraph(time_str, style_body)],
        [Paragraph("<b>Venue</b>", style_body), Paragraph(location_str, style_body)],
        [Paragraph("<b>Prepared by</b>", style_body), Paragraph(prep_name, style_body)],
        [Paragraph("<b>Date prepared</b>", style_body), Paragraph(datetime.datetime.now().strftime("%B %d, %Y"), style_body)]
    ]
    t_details = Table(details_data, colWidths=[2 * inch, 4.5 * inch])
    t_details.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#69727d')),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6)
    ]))
    story.append(t_details)
    story.append(Spacer(1, 10))

    story.append(Paragraph("Attendees", style_h2))
    all_atts = meeting_details.get("prime_attendees", []) + meeting_details.get("external_attendees", [])
    att_items = [ListItem(Paragraph(a.strip(), style_body)) for a in all_atts if a.strip()]
    if att_items:
        story.append(ListFlowable(att_items, bulletType='bullet'))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Purpose & Summary", style_h2))
    story.append(Paragraph(other_discussions if other_discussions.strip() else "To discuss project updates, ongoing deliverables, and establish clear action plans.", style_body))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Discussion Points", style_h2))
    dp_items = [ListItem(Paragraph(str(row.get('Discussion Points', '')), style_body)) for _, row in df.iterrows()]
    if dp_items:
        story.append(ListFlowable(dp_items, bulletType='1'))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Action Plan", style_h2))
    act_data = [[Paragraph("<b>#</b>", style_th), Paragraph("<b>Action Plan</b>", style_th), Paragraph("<b>Owner</b>", style_th), Paragraph("<b>Deadline</b>", style_th)]]
    for i, row in df.iterrows():
        act_data.append([
            Paragraph(str(i+1), style_td_center),
            Paragraph(str(row.get("Action Plan", "")), style_td),
            Paragraph(str(row.get("Person-in-charge", "")), style_td_center),
            Paragraph(str(row.get("Indicative Delivery Date", "")), style_td_center)
        ])
    
    t_act = Table(act_data, colWidths=[0.4 * inch, 3.5 * inch, 1.3 * inch, 1.3 * inch], repeatRows=1)
    t_act.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#003366')),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor('#69727d')),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5)
    ]))
    story.append(t_act)

    story.append(Spacer(1, 24))
    footer_style = ParagraphStyle('Footer2', parent=styles['Normal'], fontName=docfonts.FONT_BODY, fontSize=8, textColor=colors.HexColor('#69727d'), alignment=1)
    story.append(Paragraph(f"Prepared for circulation to {primary_client_rep}. Please return corrections before this is treated as the agreed record.", footer_style))

    doc.build(story)
    buffer.seek(0)
    return buffer

# -------------------------------------------------------------
# Recording Studio Dialog (frees audio_input from Streamlit reruns)
# -------------------------------------------------------------
@st.dialog("Recording Studio", width="large")
def recording_studio_dialog():
    """Modal dialog for recording audio + taking notes — survives reruns."""
    col_left, col_right = st.columns([0.62, 0.38])
    
    # LEFT: Notepad
    with col_left:
        st.markdown("##### Meeting Notes")
        notes = st.text_area(
            "Jot down key points during the meeting...",
            value=st.session_state.get("_dialog_record_notes", ""),
            height=340,
            placeholder="Key decisions, action items, questions to raise...",
            label_visibility="collapsed"
        )
        st.session_state["_dialog_record_notes"] = notes
    
    # RIGHT: Recorder + Timer
    with col_right:
        st.markdown("##### Recorder")
        recorded = st.audio_input("Record meeting audio", label_visibility="collapsed")
        
        # JS timer placeholder — shows elapsed time via browser
        timer_placeholder = st.empty()
        timer_placeholder.markdown(
            '<p style="font-size:0.85rem; color:#69727d; margin-top:0.5rem;">'
            '<span id="rec-timer">[ Ready ]</span></p>',
            unsafe_allow_html=True
        )
        
        if recorded:
            # Store audio bytes
            st.session_state["_dialog_recorded_bytes"] = recorded.read()
            
            # Show captured status with replay
            st.audio(st.session_state["_dialog_recorded_bytes"], format="audio/wav")
            duration_s = len(st.session_state["_dialog_recorded_bytes"]) / 32000  # rough estimate
            timer_placeholder.markdown(
                f'<p style="font-size:0.9rem; color:#03543F; font-weight:600; background:#DEF7EC; '
                f'padding:0.4rem 0.6rem; border-radius:4px;">'
                f'[ Captured (~{max(1, int(duration_s))} sec) ]</p>',
                unsafe_allow_html=True
            )
        else:
            st.markdown(
                '<p style="font-size:0.82rem; color:#69727d; margin-top:0.3rem;">'
                'Click the microphone above to start. '
                'Your browser will ask for microphone permission.</p>',
                unsafe_allow_html=True
            )
    
    # Bottom action bar
    st.markdown("<hr style='margin:0.5rem 0;'>", unsafe_allow_html=True)
    act_c1, act_c2, act_c3 = st.columns([1, 1, 1])
    with act_c1:
        if st.button("Save & Close", use_container_width=True):
            if st.session_state["_dialog_recorded_bytes"] is None:
                st.warning("No recording captured yet.")
            else:
                st.rerun()
    with act_c2:
        if st.button("Discard & Close", use_container_width=True):
            st.session_state["_dialog_recorded_bytes"] = None
            st.session_state["_dialog_record_notes"] = ""
            st.rerun()
    with act_c3:
        if st.button("Clear & Re-record", use_container_width=True):
            st.session_state["_dialog_recorded_bytes"] = None
            st.rerun()

# 8. UI Layout
col_upload, col_details = st.columns(2)

# LEFT CONTAINER: Audio & Text Upload Section
with col_upload:
    with st.container(height=520, border=True):
        st.markdown('<h3>Input & Transcription</h3>', unsafe_allow_html=True)
        tab_upload, tab_record, tab_text = st.tabs(["Upload Audio", "Record Audio", "Upload Text"])
        with tab_upload:
            uploaded_file = st.file_uploader("Upload audio file (200MB limit supported)", type=["wav", "mp3", "m4a", "ogg", "flac", "mp4", "webm"], help="Audio uploads up to 200MB are supported.")
            if uploaded_file and not st.session_state["_auto_processing"]:
                file_key = f"{uploaded_file.name}_{uploaded_file.size}"
                if file_key != st.session_state.get("last_processed_file"):
                    st.session_state["last_processed_file"] = file_key
                    st.session_state["_auto_processing"] = True
                    p_bar = st.progress(0, text="Initializing audio pipeline (0%)...")
                    p_status = st.empty()
                    raw_transcript = transcribe_audio_pipeline(uploaded_file.read(), uploaded_file.name, p_bar, p_status)
                    p_bar.empty()
                    p_status.empty()
                    if raw_transcript:
                        clean_tx, logs = preprocess_transcript_entities(raw_transcript)
                        st.session_state["transcript"] = clean_tx
                        st.session_state["entity_corrections_log"] = logs
                        set_mom_dataframe(pd.DataFrame(columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]))
                        st.session_state["other_discussions"] = ""
                        st.session_state["chat_history"] = []
                        st.session_state["matched_evidence_items"] = []
                        st.session_state["user_topics_text"] = ""
                        st.session_state["_topics_discovered"] = False
                        meta = extract_metadata_with_deepseek(clean_tx)
                        if meta:
                            if meta.get("meeting_type") and meta["meeting_type"] in MEETING_TYPE_OPTIONS: st.session_state["meeting_type"] = meta["meeting_type"]
                            if meta.get("client_name"): st.session_state["meeting_client_name"] = meta["client_name"]
                            if meta.get("location"): st.session_state["meeting_location"] = meta["location"]
                            if meta.get("crd_attendees"):
                                matched_crd = [c for c in meta["crd_attendees"] if c in CRD_MEMBERS]
                                if matched_crd: st.session_state["meeting_selected_crd"] = matched_crd
                            if meta.get("external_attendees"): st.session_state["meeting_ext_attendees"] = meta["external_attendees"]
                            if meta.get("prepared_by"): st.session_state["meeting_prep_name"] = meta["prepared_by"]
                            if meta.get("confirmed_by"): st.session_state["meeting_conf_name"] = meta["confirmed_by"]
                        st.session_state["_auto_processing"] = False
                        st.rerun()
                    else:
                        st.session_state["last_processed_file"] = None
                        st.session_state["_auto_processing"] = False
                        st.error("Transcription failed. Please try again or upload text directly.")
        with tab_record:
            # Open Recording Studio button
            if st.button("Open Recording Studio", key="btn_open_rec_studio", use_container_width=True):
                recording_studio_dialog()
            
            st.markdown("<div style='height:0.5rem;'></div>", unsafe_allow_html=True)
            
            # Status bar if audio was captured via dialog
            stored_bytes = st.session_state.get("_dialog_recorded_bytes")
            if stored_bytes is not None:
                st.audio(stored_bytes, format="audio/wav")
                # Notes preview
                stored_notes = st.session_state.get("_dialog_record_notes", "")
                if stored_notes:
                    preview = stored_notes[:200] + ("..." if len(stored_notes) > 200 else "")
                    with st.expander(f"Meeting Notes ({len(stored_notes)} chars)", expanded=False):
                        st.text_area("Notes", value=stored_notes, height=120, disabled=True, label_visibility="collapsed")
                
                st.markdown(
                    f'<p style="font-size:0.85rem; color:#03543F; font-weight:600; '
                    f'background:#DEF7EC; padding:0.4rem 0.6rem; border-radius:4px;">'
                    f'[ Recording saved ({len(stored_bytes)//32000}s estimated) ]</p>',
                    unsafe_allow_html=True
                )
                
                r_btn1, r_btn2, r_btn3 = st.columns([1.5, 1, 1])
                with r_btn1:
                    if st.button("Transcribe Recording", key="btn_tx_dialog", use_container_width=True):
                        st.session_state["_auto_processing"] = True
                        p_bar = st.progress(0, text="Initializing audio pipeline (0%)...")
                        p_status = st.empty()
                        raw_transcript = transcribe_audio_pipeline(stored_bytes, "recording.wav", p_bar, p_status)
                        p_bar.empty()
                        p_status.empty()
                        if raw_transcript:
                            clean_tx, logs = preprocess_transcript_entities(raw_transcript)
                            st.session_state["transcript"] = clean_tx
                            st.session_state["entity_corrections_log"] = logs
                            # Carry dialog notes into user_notes for AI context
                            if stored_notes:
                                st.session_state["user_notes"] = stored_notes + "\n" + st.session_state.get("user_notes", "")
                            set_mom_dataframe(pd.DataFrame(columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]))
                            st.session_state["other_discussions"] = ""
                            st.session_state["chat_history"] = []
                            st.session_state["matched_evidence_items"] = []
                            st.session_state["user_topics_text"] = ""
                            st.session_state["_topics_discovered"] = False
                            st.session_state["_auto_processing"] = False
                            st.rerun()
                        else:
                            st.session_state["_auto_processing"] = False
                            st.error("Transcription failed. Please try again or upload text directly.")
                with r_btn2:
                    if st.button("Clear Recording", key="btn_clear_dialog", use_container_width=True):
                        st.session_state["_dialog_recorded_bytes"] = None
                        st.session_state["_dialog_record_notes"] = ""
                        st.rerun()
                with r_btn3:
                    st.download_button(label="Download (.wav)", data=stored_bytes, file_name=f"Recording_{datetime.date.today().strftime('%Y%m%d')}.wav", mime="audio/wav", use_container_width=True)
            else:
                st.markdown(
                    '<p style="font-size:0.85rem; color:#69727d; margin-top:1rem; text-align:center;">'
                    'Click "Open Recording Studio" to start. The modal protects the recorder from page refreshes, '
                    'and includes a notepad for live meeting notes.</p>',
                    unsafe_allow_html=True
                )
        with tab_text:
            uploaded_text_file = st.file_uploader("Upload Document (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"])
            pasted_text = st.text_area("Or Paste Transcript Here", height=95, placeholder="Paste transcript text directly here...")
            if st.button("Process Text", key="btn_tx_text"):
                p_bar = st.progress(0, text="Extracting document text (0%)...")
                time.sleep(0.2)
                p_bar.progress(50, text="Reading document stream (50%)...")
                extracted_str = ""
                if uploaded_text_file: extracted_str = extract_text_from_file(uploaded_text_file)
                if pasted_text and pasted_text.strip(): extracted_str += "\n" + pasted_text.strip()
                p_bar.progress(100, text="Document processed (100%)!")
                time.sleep(0.2)
                p_bar.empty()
                if extracted_str.strip():
                    clean_tx, logs = preprocess_transcript_entities(extracted_str.strip())
                    st.session_state["transcript"] = clean_tx
                    st.session_state["entity_corrections_log"] = logs
                    set_mom_dataframe(pd.DataFrame(columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]))
                    st.session_state["other_discussions"] = ""
                    st.session_state["chat_history"] = []
                    st.session_state["matched_evidence_items"] = []
                    st.session_state["user_topics_text"] = ""
                    st.session_state["_topics_discovered"] = False
                    st.rerun()
                else:
                    st.warning("Please upload a file or paste text to proceed.")

        # Transcript expander (hidden by default, inside upload card)
        if st.session_state["transcript"]:
            st.markdown("---")
            with st.expander("Show Full Transcript", expanded=False):
                st.text_area("Transcript Content", st.session_state["transcript"], height=250, label_visibility="collapsed")
                if st.session_state.get("entity_corrections_log"):
                    with st.expander(f"Entity Standardizations ({len(st.session_state['entity_corrections_log'])})"):
                        for item in st.session_state["entity_corrections_log"]:
                            st.caption(f"• {item}")
                t_col1, t_col2 = st.columns(2)
                with t_col1:
                    escaped_tx = json.dumps(st.session_state["transcript"])
                    copy_html = f"""
                    <!DOCTYPE html><html><head><style>body{{margin:0;padding:0;font-family:'Montserrat',sans-serif;}}button{{width:100%;height:36px;background-color:#0c0c0e;color:#ffffff;border:1px solid #c9ab4c;border-radius:6px;font-size:0.82rem;font-weight:600;cursor:pointer;transition:all 0.2s ease;}}button:hover{{background-color:#003366;border-color:#d9bc5d;color:#ffffff;}}</style></head><body><button id="copy-btn">{COPY_ICON} Copy Text</button><script>document.getElementById("copy-btn").addEventListener("click",function(){{navigator.clipboard.writeText({escaped_tx}).then(function(){{document.getElementById("copy-btn").innerHTML = '{COPY_ICON} Copied';setTimeout(() => document.getElementById("copy-btn").innerHTML = '{COPY_ICON} Copy Text', 2000);}});}});</script></body></html>
                    """
                    components.html(copy_html, height=36)
                with t_col2:
                    st.download_button(label="Download Transcript", data=st.session_state["transcript"], file_name=f"Transcript_{st.session_state['meeting_date'].strftime('%Y%m%d')}.txt", mime="text/plain", use_container_width=True)

# User Notes section (between Upload and Meeting Details)
if "user_notes" not in st.session_state:
    st.session_state["user_notes"] = ""
with st.container(border=True):
    st.markdown('<h3>Meeting Notes (Optional)</h3>', unsafe_allow_html=True)
    st.caption("Upload or paste pre-meeting notes, agenda, or observations to enrich the AI context.")
    notes_file = st.file_uploader("Upload notes (.txt, .docx, .pdf)", type=["txt", "docx", "pdf"], key="notes_file_uploader", label_visibility="collapsed")
    notes_text = st.text_area("Or paste notes here", value=st.session_state["user_notes"], height=100, placeholder="Pre-meeting agenda, key objectives, client background...", label_visibility="collapsed")
    
    note_changed = False
    if notes_file:
        extracted = extract_text_from_file(notes_file)
        if extracted:
            st.session_state["user_notes"] = extracted
            note_changed = True
    if notes_text != st.session_state.get("user_notes", ""):
        st.session_state["user_notes"] = notes_text
        note_changed = True
    
    if st.session_state["user_notes"]:
        preview = st.session_state["user_notes"][:300] + ("..." if len(st.session_state["user_notes"]) > 300 else "")
        st.markdown(f'<div style="font-size:0.82rem; color:#333; background:#F9FAFB; padding:0.5rem; border-left:3px solid #003366; margin-top:0.3rem;"><i>{preview}</i></div>', unsafe_allow_html=True)
        if len(st.session_state["user_notes"]) > 300:
            with st.expander("Show full notes"):
                st.text_area("Notes full", value=st.session_state["user_notes"], height=150, disabled=True, label_visibility="collapsed")
        col_c1, _ = st.columns([1, 9])
        with col_c1:
            if st.button("Clear Notes", key="btn_clear_notes"):
                st.session_state["user_notes"] = ""
                st.rerun()

# RIGHT CONTAINER: Meeting Details Card
with col_details:
    with st.container(height=520, border=True):
        if st.session_state["transcript"]:
            head_col1, head_col_auto, head_col2 = st.columns([5.5, 3.5, 1.0])
            with head_col_auto:
                if st.button("Populate from Transcript", key="btn_auto_populate"):
                    with st.spinner("Extracting metadata..."):
                        meta = extract_metadata_with_deepseek(st.session_state["transcript"])
                        if meta:
                            if meta.get("meeting_type") and meta["meeting_type"] in MEETING_TYPE_OPTIONS:
                                st.session_state["meeting_type"] = meta["meeting_type"]
                            if meta.get("client_name"): st.session_state["meeting_client_name"] = meta["client_name"]
                            if meta.get("location"): st.session_state["meeting_location"] = meta["location"]
                            if meta.get("crd_attendees"):
                                matched_crd = [c for c in meta["crd_attendees"] if c in CRD_MEMBERS]
                                if matched_crd: st.session_state["meeting_selected_crd"] = matched_crd
                            if meta.get("external_attendees"): st.session_state["meeting_ext_attendees"] = meta["external_attendees"]
                            if meta.get("prepared_by"): st.session_state["meeting_prep_name"] = meta["prepared_by"]
                            if meta.get("confirmed_by"): st.session_state["meeting_conf_name"] = meta["confirmed_by"]
                            st.rerun()
        else:
            head_col1, head_col2 = st.columns([9.0, 1.0])
        with head_col1:
            st.markdown('<h3 style="margin-top:0.2rem;">Meeting Details</h3>', unsafe_allow_html=True)
        with head_col2:
            if st.button("", key="card_settings_btn", help="Settings"):
                st.session_state["show_settings"] = not st.session_state["show_settings"]
                st.rerun()
        if st.session_state["show_settings"]:
            with st.expander("Settings & Engine Diagnostics", expanded=True):
                set_col1, set_col2 = st.columns(2)
                with set_col1:
                    engine_options = ["AI - DeepSeek", "Non-AI - Python Heuristic"]
                    selected_eng = st.selectbox("MoM Generation Engine", options=engine_options, index=engine_options.index(st.session_state["selected_engine"]) if st.session_state["selected_engine"] in engine_options else 0)
                    st.session_state["selected_engine"] = selected_eng
                with set_col2:
                    st.markdown("**Diagnostics**")
                    st.write(f"• **Session Tokens:** `{st.session_state['tokens_used']:,}`")
                    if st.session_state["last_api_call"]:
                        last_call = st.session_state["last_api_call"]
                        st.write(f"• **Last Call:** `{last_call.strftime('%I:%M:%S %p')}`")
            st.markdown("---")
        
        # Row 1: Date, Location, Meeting Type
        r1_c1, r1_c2, r1_c3 = st.columns([1.1, 1.4, 0.9])
        with r1_c1:
            meeting_date = st.date_input("Date", value=st.session_state["meeting_date"])
            st.session_state["meeting_date"] = meeting_date
        with r1_c2:
            current_loc = st.session_state.get("meeting_location", "")
            loc_options = list(LOCATION_PRESETS)
            if current_loc and current_loc not in loc_options: loc_options.append(current_loc)
            loc_options.append("Other / Custom...")
            default_idx = loc_options.index(current_loc) if current_loc in loc_options else 0
            selected_loc_choice = st.selectbox("Location", options=loc_options, index=default_idx)
            
            if selected_loc_choice == "Other / Custom...":
                custom_loc = st.text_input("Enter Location", value="" if current_loc in LOCATION_PRESETS else current_loc, placeholder="e.g. Boardroom or Client Office", label_visibility="collapsed")
                meeting_location = custom_loc
            else:
                meeting_location = selected_loc_choice
            st.session_state["meeting_location"] = meeting_location if meeting_location else ""
        with r1_c3:
            curr_type = st.session_state.get("meeting_type", "Internal")
            type_idx = MEETING_TYPE_OPTIONS.index(curr_type) if curr_type in MEETING_TYPE_OPTIONS else 0
            meeting_type = st.selectbox("Meeting Type", options=MEETING_TYPE_OPTIONS, index=type_idx)
            st.session_state["meeting_type"] = meeting_type

        # Row 2: Start and End Times
        r2_c1, r2_c2 = st.columns(2)
        with r2_c1:
            st.markdown("<p style='font-size:0.85rem; margin-bottom:0.2rem; color:#333; font-weight:500;'>Start Time</p>", unsafe_allow_html=True)
            sc1, sc2, sc3 = st.columns([1, 1, 1.2])
            sh = sc1.selectbox("SH", [f"{i:02d}" for i in range(1,13)], key="sh", label_visibility="collapsed")
            sm = sc2.selectbox("SM", [f"{i:02d}" for i in range(0,60,5)], key="sm", label_visibility="collapsed")
            sap = sc3.selectbox("SAP", ["AM", "PM"], key="sap", label_visibility="collapsed")
            start_str = f"{sh}:{sm} {sap}"
        with r2_c2:
            st.markdown("<p style='font-size:0.85rem; margin-bottom:0.2rem; color:#333; font-weight:500;'>End Time</p>", unsafe_allow_html=True)
            ec1, ec2, ec3 = st.columns([1, 1, 1.2])
            eh = ec1.selectbox("EH", [f"{i:02d}" for i in range(1,13)], key="eh", label_visibility="collapsed")
            em = ec2.selectbox("EM", [f"{i:02d}" for i in range(0,60,5)], key="em", label_visibility="collapsed")
            eap = ec3.selectbox("EAP", ["AM", "PM"], key="eap", label_visibility="collapsed")
            end_str = f"{eh}:{em} {eap}"
        
        # Row 3: Client, Attendees, Prepared & Confirmed details
        r3_c1, r3_c2 = st.columns(2)
        with r3_c1:
            client_name = st.text_input("Client / Company / Department", value=st.session_state["meeting_client_name"], placeholder="XYZ Company")
            st.session_state["meeting_client_name"] = client_name
            selected_crd = st.multiselect("CRD Team Attendees", options=CRD_MEMBERS, default=st.session_state["meeting_selected_crd"])
            st.session_state["meeting_selected_crd"] = selected_crd
        with r3_c2:
            ext_attendees_raw = st.text_input("External Attendees", value=st.session_state["meeting_ext_attendees"], placeholder="e.g. Mr. ABCD, Jane Doe")
            st.session_state["meeting_ext_attendees"] = ext_attendees_raw
            prep_col, conf_col = st.columns(2)
            with prep_col:
                prep_name = st.text_input("Prepared By", value=st.session_state["meeting_prep_name"], placeholder="Name")
                st.session_state["meeting_prep_name"] = prep_name
                prep_desig = st.text_input("Prep Designation", value=st.session_state["meeting_prep_desig"], placeholder="Designation")
                st.session_state["meeting_prep_desig"] = prep_desig
            with conf_col:
                conf_name = st.text_input("Confirmed By", value=st.session_state["meeting_conf_name"], placeholder="Name")
                st.session_state["meeting_conf_name"] = conf_name
                conf_desig = st.text_input("Conf Designation", value=st.session_state["meeting_conf_desig"], placeholder="Designation")
                st.session_state["meeting_conf_desig"] = conf_desig

# Speaker Identity Mapping Quick-Bar
if st.session_state["transcript"]:
    speakers_found = detect_speaker_tags(st.session_state["transcript"])
    if speakers_found:
        with st.container(border=True):
            st.markdown('<h3>Speaker Identity Mapping</h3>', unsafe_allow_html=True)
            st.caption("Map detected raw speaker tags directly to confirmed attendee names across the entire transcript.")
            spk_cols = st.columns(min(len(speakers_found), 4))
            attendee_candidates = [""] + selected_crd + [x.strip() for x in ext_attendees_raw.split(",") if x.strip()]
            for idx, spk in enumerate(speakers_found):
                col_target = spk_cols[idx % len(spk_cols)]
                with col_target:
                    st.session_state["speaker_mappings"][spk] = st.selectbox(
                        f"Map '{spk}' to:",
                        options=attendee_candidates,
                        index=0,
                        key=f"spk_map_select_{spk}"
                    )
            if st.button("Apply Speaker Replacements to Transcript", key="btn_apply_spk_remap"):
                remapped = apply_speaker_remapping(st.session_state["transcript"], st.session_state["speaker_mappings"])
                st.session_state["transcript"] = remapped
                st.success("Transcript updated with mapped speaker identities!")
                st.session_state["_topics_discovered"] = False
                st.rerun()

# =====================================================================
# PHASE 2 & 3: Meeting Minutes Review + Export
# =====================================================================
if st.session_state["transcript"]:
    
    # Auto-discover topics + match evidence on first load
    if not st.session_state.get("_topics_discovered") and not st.session_state["matched_evidence_items"]:
        with st.spinner("Discovering discussion topics and matching evidence..."):
            st.session_state["_topics_discovered"] = True
            topics_suggested = suggest_discussion_topics_from_transcript(st.session_state["transcript"])
            st.session_state["user_topics_text"] = topics_suggested
            items, other_disc = match_evidence_and_synthesize(st.session_state["transcript"], topics_suggested)
            if items:
                for itm in items:
                    itm["approved"] = True
                st.session_state["matched_evidence_items"] = items
                st.session_state["other_discussions"] = other_disc
    
    with st.container(border=True):
        st.markdown('<h3>Meeting Minutes Review</h3>', unsafe_allow_html=True)
        st.caption("Review and edit discussion items below. Each card includes source evidence and inline fields. Approve items, then export.")
        
        # Action bar
        action_bar1, action_bar2, action_bar3 = st.columns([3, 3, 4])
        with action_bar1:
            if st.button("Re-discover Topics", key="btn_rediscover_topics"):
                with st.spinner("Re-scanning transcript for topics..."):
                    sugg = suggest_discussion_topics_from_transcript(st.session_state["transcript"])
                    st.session_state["user_topics_text"] = sugg
                    st.session_state["matched_evidence_items"] = []
                    st.session_state["_topics_discovered"] = False
                    st.rerun()
        with action_bar2:
            if st.button("Re-match Evidence", key="btn_rematch_evidence"):
                with st.spinner("Re-matching evidence with current topics..."):
                    items, other_disc = match_evidence_and_synthesize(
                        st.session_state["transcript"],
                        st.session_state["user_topics_text"]
                    )
                    if items:
                        for itm in items:
                            itm["approved"] = True
                        st.session_state["matched_evidence_items"] = items
                        st.session_state["other_discussions"] = other_disc
                        st.success(f"Matched {len(items)} items!")
                        st.rerun()
                    else:
                        st.warning("Could not match evidence. Check transcript contents.")
        with action_bar3:
            if st.button("+ Add Item Manually", key="btn_add_review_item"):
                st.session_state["matched_evidence_items"].append({
                    "topic_title": "New Item",
                    "discussion_point": "",
                    "evidence_quote": "",
                    "action_plan": "",
                    "indicative_delivery_date": "TBD",
                    "person_in_charge": "Unassigned",
                    "confidence": "Medium",
                    "approved": True
                })
                st.rerun()
        
        st.markdown("<hr style='margin:0.5rem 0; border:none; border-top:1px solid rgba(0,0,0,0.07);'>", unsafe_allow_html=True)
        
        # Discussion topics editor
        st.markdown('<span class="playfair-label">Discussion Topics</span>', unsafe_allow_html=True)
        st.caption("These topics guide evidence matching. Edits take effect after re-matching.")
        st.session_state["user_topics_text"] = st.text_area(
            "Topics",
            value=st.session_state["user_topics_text"],
            height=80,
            label_visibility="collapsed",
            placeholder="1. Architecture updates\n2. Q3 Delivery Deadlines\n3. Client Integration Requirements"
        )
        
        st.markdown("<hr style='margin:0.5rem 0; border:none; border-top:1px solid rgba(0,0,0,0.07);'>", unsafe_allow_html=True)
        
        # Recommended missed points
        missed_recs = st.session_state.get("recommended_missed_points", []) or []
        if missed_recs:
            with st.expander(f"Echo found topics you may have missed ({len(missed_recs)})", expanded=False):
                add_these = []
                for ridx, rec in enumerate(missed_recs):
                    rec_title = str(rec.get("topic_title") or "").strip()
                    rec_quote = str(rec.get("evidence_quote") or "").strip()
                    if not rec_title:
                        continue
                    col_c, col_q = st.columns([1.5, 8.5])
                    with col_c:
                        pick = st.checkbox("Add", key=f"chip_rec_{ridx}")
                    with col_q:
                        st.markdown(f"**{rec_title}**")
                        if rec_quote:
                            st.markdown(f'<div class="evidence-quote-box" style="font-size:0.78rem;"><b>Quote:</b> "{rec_quote}"</div>', unsafe_allow_html=True)
                    if pick:
                        add_these.append(rec_title)
                if add_these and st.button("Add selected & re-match", key="btn_add_missed_recs"):
                    current = (st.session_state.get("user_topics_text") or "").strip()
                    new_lines = "\n".join(f"- {t}" for t in add_these)
                    st.session_state["user_topics_text"] = (current + "\n" + new_lines).strip()
                    st.session_state["recommended_missed_points"] = []
                    st.session_state["matched_evidence_items"] = []
                    st.session_state["_topics_discovered"] = False
                    st.success("Added topics. Re-running evidence matching...")
                    st.rerun()
        
        # Matched evidence items as cards
        all_valid_attendees = selected_crd + [x.strip() for x in ext_attendees_raw.split(",") if x.strip()]
        
        if not st.session_state["matched_evidence_items"]:
            st.info("No discussion items yet. Click 'Re-discover Topics' to auto-generate, or add items manually.")
        else:
            for idx, item in enumerate(st.session_state["matched_evidence_items"]):
                with st.container(border=True):
                    # Header row
                    h_col1, h_col2, h_col3 = st.columns([6, 2.5, 1.5])
                    with h_col1:
                        topic_title = st.text_input(
                            "Topic",
                            value=item.get("topic_title", f"Point {idx+1}"),
                            key=f"rev_topic_{idx}",
                            label_visibility="collapsed"
                        )
                        item["topic_title"] = topic_title
                    with h_col2:
                        conf = item.get("confidence", "Medium")
                        conf_class = "badge-high" if conf.lower() == "high" else ("badge-low" if conf.lower() == "low" else "badge-medium")
                        st.markdown(f'<span class="badge-confidence {conf_class}">{conf} Grounding</span>', unsafe_allow_html=True)
                    with h_col3:
                        item["approved"] = st.checkbox("Approve", value=item.get("approved", True), key=f"rev_app_{idx}")
                    
                    # Evidence quote
                    eq = item.get("evidence_quote", "").strip()
                    if eq:
                        st.markdown(f'<div class="evidence-quote-box"><b>Verbatim Source Quote:</b> "{eq}"</div>', unsafe_allow_html=True)
                    
                    # Inline editable fields
                    f_col1, f_col2, f_col3 = st.columns([3.5, 3.5, 3.0])
                    with f_col1:
                        item["discussion_point"] = st.text_area(
                            "Discussion Point",
                            value=item.get("discussion_point", ""),
                            key=f"rev_dp_{idx}",
                            height=60,
                            label_visibility="collapsed"
                        )
                    with f_col2:
                        item["action_plan"] = st.text_area(
                            "Action Plan",
                            value=item.get("action_plan", ""),
                            key=f"rev_ap_{idx}",
                            height=60,
                            label_visibility="collapsed"
                        )
                    with f_col3:
                        sub_c1, sub_c2 = st.columns(2)
                        with sub_c1:
                            item["indicative_delivery_date"] = st.text_area(
                                "Delivery Date",
                                value=item.get("indicative_delivery_date", "TBD"),
                                key=f"rev_dd_{idx}",
                                height=60,
                                label_visibility="collapsed"
                            )
                        with sub_c2:
                            item["person_in_charge"] = st.text_area(
                                "PIC",
                                value=item.get("person_in_charge", "Unassigned"),
                                key=f"rev_pic_{idx}",
                                height=60,
                                label_visibility="collapsed"
                            )
                    
                    # Guardrail warnings
                    row_dict = {
                        "Discussion Points": item.get("discussion_point", ""),
                        "Action Plan": item.get("action_plan", ""),
                        "Indicative Delivery Date": item.get("indicative_delivery_date", ""),
                        "Person-in-charge": item.get("person_in_charge", "")
                    }
                    warnings = check_row_guardrails(row_dict, all_valid_attendees)
                    for w in warnings:
                        st.markdown(f'<div class="guardrail-alert">{w}</div>', unsafe_allow_html=True)
                    
                    # Delete button
                    del_col1, _ = st.columns([1, 9])
                    with del_col1:
                        if st.button("Delete", key=f"del_rev_{idx}"):
                            st.session_state["matched_evidence_items"].pop(idx)
                            st.rerun()
        
        st.markdown("<hr style='margin:0.5rem 0; border:none; border-top:1px solid rgba(0,0,0,0.07);'>", unsafe_allow_html=True)
        
        # Other Discussions
        st.markdown('<span class="playfair-label">Other Discussions / Summary</span>', unsafe_allow_html=True)
        st.session_state["other_discussions"] = st.text_area(
            "Other Discussions Content",
            value=st.session_state["other_discussions"],
            height=80,
            label_visibility="collapsed"
        )
        
        # Build DataFrame from approved items (updates df dynamically)
        approved_rows = []
        for item in st.session_state["matched_evidence_items"]:
            if item.get("approved", True):
                approved_rows.append({
                    "Discussion Points": item.get("discussion_point", ""),
                    "Action Plan": item.get("action_plan", ""),
                    "Indicative Delivery Date": item.get("indicative_delivery_date", "TBD"),
                    "Person-in-charge": item.get("person_in_charge", "Unassigned")
                })
        
        if approved_rows:
            set_mom_dataframe(pd.DataFrame(approved_rows))
        
        # =====================================================================
        # EXPORT SECTION (inside Meeting Minutes Review)
        # =====================================================================
        st.markdown("<hr style='margin:0.8rem 0; border:none; border-top:2px solid rgba(0,51,102,0.15);'>", unsafe_allow_html=True)
        st.markdown('<h3 style="margin-top:0.2rem;">Export</h3>', unsafe_allow_html=True)
        
        if st.session_state["df"].empty:
            st.info("Approve at least one discussion item above to enable export.")
        else:
            time_range_str = f"{start_str} to {end_str}"
            meeting_details = {
                "date": meeting_date.strftime("%B %d, %Y"), "time_range": time_range_str,
                "meeting_type": st.session_state.get("meeting_type", "Internal"),
                "location": meeting_location if meeting_location.strip() else "____________",
                "company_name": client_name.strip() if client_name.strip() else "",
                "prime_attendees": selected_crd,
                "external_attendees": [x.strip() for x in ext_attendees_raw.split(",") if x.strip()],
                "prep_name": prep_name.strip(), "prep_desig": prep_desig.strip(),
                "conf_name": conf_name.strip(), "conf_desig": conf_desig.strip()
            }
            
            # Template picker
            template_selection = st.selectbox(
                "Select MoM Template Format",
                options=["Template 1 - Standard Corporate (Combined Table)", "Template 2 - Detailed General Meeting (Vertical Layout)"],
                label_visibility="collapsed"
            )
            
            # Export buttons
            exp_row1, exp_row2, exp_row3, exp_row4 = st.columns([2.5, 2.5, 1.5, 3.5])
            if "Template 1" in template_selection:
                with exp_row1:
                    doc_bio = export_to_word_template_1(st.session_state["df"], meeting_details, st.session_state["other_discussions"])
                    st.download_button(label="Download DOCX", data=doc_bio, file_name=f"MOM_{client_name.replace(' ', '_') if client_name else 'Report'}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="btn_download_docx_1")
                with exp_row2:
                    pdf_bio = export_to_pdf_template_1(st.session_state["df"], meeting_details, st.session_state["other_discussions"])
                    st.download_button(label="Download PDF", data=pdf_bio, file_name=f"MOM_{client_name.replace(' ', '_') if client_name else 'Report'}.pdf", mime="application/pdf", key="btn_download_pdf_1")
            else:
                with exp_row1:
                    doc_bio = export_to_word_template_2(st.session_state["df"], meeting_details, st.session_state["other_discussions"])
                    st.download_button(label="Download DOCX", data=doc_bio, file_name=f"MOM_Detailed_{client_name.replace(' ', '_') if client_name else 'Report'}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key="btn_download_docx_2")
                with exp_row2:
                    pdf_bio = export_to_pdf_template_2(st.session_state["df"], meeting_details, st.session_state["other_discussions"])
                    st.download_button(label="Download PDF", data=pdf_bio, file_name=f"MOM_Detailed_{client_name.replace(' ', '_') if client_name else 'Report'}.pdf", mime="application/pdf", key="btn_download_pdf_2")
            
            with exp_row3:
                if st.button("Save to Archive", key="btn_save_supabase_bottom"):
                    success, msg = save_meeting_to_supabase(meeting_details, st.session_state["df"], st.session_state["other_discussions"], st.session_state["transcript"])
                    if success: st.success(msg)
                    else: st.error(f"Save failed: {msg}")
            
            with exp_row4:
                target_webhook = st.text_input("Webhook URL", value=SLACK_WEBHOOK_URL, placeholder="https://hooks.slack.com/services/...", label_visibility="collapsed")
                if st.button("Sync to Webhook", key="btn_sync_webhook"):
                    ok, msg = dispatch_action_items_webhook(target_webhook, st.session_state["df"], meeting_details)
                    if ok: st.success(msg)
                    else: st.error(msg)
            
            # Ask Echo inline chat
            st.markdown("<hr style='margin:0.5rem 0; border:none; border-top:1px solid rgba(0,0,0,0.07);'>", unsafe_allow_html=True)
            with st.expander("Ask Echo (AI Assistant)", expanded=False):
                st.caption("Ask questions or issue live commands: 'Change row 2 PIC to Kristina', 'Add task for Sondi', or 'Delete row 3'.")
                st.markdown('<div class="chat-container">', unsafe_allow_html=True)
                if not st.session_state["chat_history"]:
                    st.markdown('<div class="chat-ai">Hello. I am Echo. Ask questions or tell me how to refine your Minutes of Meeting table.</div>', unsafe_allow_html=True)
                else:
                    for msg in st.session_state["chat_history"]:
                        if msg["role"] == "assistant":
                            st.markdown(f'<div class="chat-ai">{msg["content"].replace(chr(10), "<br>")}</div>', unsafe_allow_html=True)
                        else:
                            st.markdown(f'<div class="chat-user-wrap"><div class="chat-user">{msg["content"]}</div></div>', unsafe_allow_html=True)
                st.markdown('</div>', unsafe_allow_html=True)
                if prompt := st.chat_input("Ask Echo or command an edit..."):
                    st.session_state["chat_history"].append({"role": "user", "content": prompt})
                    with st.spinner("Echo is analyzing request and table state..."):
                        answer, action = ask_deepseek_with_mutation(st.session_state["transcript"], prompt, st.session_state["chat_history"], st.session_state["df"])
                        if action and isinstance(action, dict):
                            tool_name = action.get("tool")
                            r_idx = int(action.get("row_index", 0))
                            fields = action.get("fields", {})
                            if tool_name == "update_row" and 0 <= r_idx < len(st.session_state["df"]):
                                for f_key, f_val in fields.items():
                                    update_mom_field(r_idx, f_key, str(f_val))
                            elif tool_name == "delete_row" and 0 <= r_idx < len(st.session_state["df"]):
                                delete_mom_row(r_idx)
                            elif tool_name == "add_row":
                                add_mom_row(
                                    fields.get("Discussion Points", ""),
                                    fields.get("Action Plan", ""),
                                    fields.get("Indicative Delivery Date", "TBD"),
                                    fields.get("Person-in-charge", "Unassigned")
                                )
                    st.session_state["chat_history"].append({"role": "assistant", "content": answer})
                    st.rerun()
