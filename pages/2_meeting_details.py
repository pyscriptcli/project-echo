import sys
import os

# Ensure root directory is resolvable
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), "..")))

import streamlit as st
import pandas as pd
from datetime import datetime, date, timedelta, time
from io import BytesIO

# Document generation imports
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml import parse_xml
from docx.shared import Inches, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle, ListFlowable, ListItem

from utils.db import fetch_meeting_archives, get_supabase_client
from components.sidebar import setup_page_layout
from utils.auth import require_login

# CRD team members for the attendee picker when editing meeting details.
CRD_MEMBERS = [
    "Sondi Tuazon", "Kristina Balajadia", "Meliza Zapata", "Dykstra Pineda",
    "Cedtrix Rena", "Carlo Medina", "Dave Policarpio", "Irish Rima",
]


def _parse_time_range(tr):
    """Parse '1:00 AM to 2:00 PM' into (start_time, end_time) or (None, None)."""
    if not tr or " to " not in tr:
        return None, None
    parts = tr.split(" to ")

    def _p(s):
        s = (s or "").strip()
        for fmt in ("%I:%M %p", "%I %p", "%H:%M"):
            try:
                return datetime.strptime(s, fmt).time()
            except ValueError:
                continue
        return None

    start = _p(parts[0]) if len(parts) > 0 else None
    end = _p(parts[1]) if len(parts) > 1 else None
    return start, end


# 1. Page Config (MUST be first)
st.set_page_config(
    page_title="Project Echo - Meetings Workspace",
    layout="wide",
    initial_sidebar_state="expanded"
)
require_login()
setup_page_layout()

# Default date filter to "This Month"
today = date.today()
first_day_of_month = today.replace(day=1)

# 2. Global State for View Mode & Filters
if "view_mode" not in st.session_state:
    st.session_state["view_mode"] = "gallery"
if "selected_meeting_id" not in st.session_state:
    st.session_state["selected_meeting_id"] = None
if "gal_search_q" not in st.session_state:
    st.session_state["gal_search_q"] = ""
if "gal_type_f" not in st.session_state:
    st.session_state["gal_type_f"] = "All Meetings"
if "gal_date_range" not in st.session_state:
    st.session_state["gal_date_range"] = ()
if "edit_meeting_details" not in st.session_state:
    st.session_state["edit_meeting_details"] = False

# 3. Custom CSS & Pure SVG Icon Button Injection
CUSTOM_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Montserrat:wght@400;500;600&family=Playfair+Display:ital,wght@1,400;1,500;1,600&display=swap');

html, body, [class*="css"] { font-family: 'Montserrat', sans-serif !important; }

.stApp {
    background-color: #F3EFE6; 
    background-image: linear-gradient(rgba(0, 0, 0, 0.04) 1px, transparent 1px), linear-gradient(90deg, rgba(0, 0, 0, 0.04) 1px, transparent 1px);
    background-size: 80px 80px;
    color: #2D2D2D;
}
.stApp > header { display: none !important; }
.block-container { padding-top: 1.5rem !important; padding-right: 2.5rem !important; padding-left: 2.5rem !important; }

h3 {
    font-family: 'Playfair Display', serif !important; 
    font-style: italic !important; 
    font-weight: 400 !important; 
    color: #1A2B4C !important; 
    letter-spacing: 0.02em; 
    margin-bottom: 0.25rem; 
    font-size: 1.35rem !important;
}

.playfair-label {
    font-family: 'Playfair Display', serif !important; 
    font-style: italic !important;
    color: #1A2B4C !important; 
    font-size: 1.05rem !important; 
    margin-bottom: 0.25rem !important; 
    display: block;
}

/* 3D Drop Shadow Containers */
div[data-testid="stVerticalBlockBorderWrapper"] {
    background-color: #FFFFFF !important; 
    border-radius: 12px !important;
    box-shadow: 0 10px 24px rgba(0, 0, 0, 0.08), 0 3px 8px rgba(0, 0, 0, 0.04) !important;
    border: 1px solid rgba(0, 0, 0, 0.06) !important; 
    padding: 1.5rem !important; 
    margin-bottom: 1.25rem !important;
    transition: transform 0.2s ease, box-shadow 0.2s ease !important;
}

/* Form Inputs */
.stTextArea textarea, .stTextInput input, [data-baseweb="input"], [data-baseweb="select"] {
    background-color: #FAFAFA !important; 
    border: 1px solid rgba(0,0,0,0.08) !important;
    border-radius: 8px !important; 
    box-shadow: inset 0 2px 4px rgba(0,0,0,0.02) !important;
}

.stTextArea textarea:focus, .stTextInput input:focus {
    background-color: #FFFFFF !important; 
    border-color: #D4AF37 !important;
}

/* Base Buttons */
.stButton > button {
    background-color: #111A2B !important; 
    color: #FFFFFF !important; 
    border: 1px solid #D4AF37 !important; 
    border-radius: 18px !important; 
    font-family: 'Inter', sans-serif !important; 
    font-weight: 600 !important; 
    font-size: 0.82rem !important; 
    height: 36px !important; 
    padding: 0 0.75rem !important;
    box-shadow: 0 4px 10px rgba(26, 43, 76, 0.18) !important; 
    transition: all 0.2s ease !important; 
    width: 100% !important;
}

.stButton > button:hover { 
    background-color: #1A2B4C !important; 
    color: #FFFFFF !important; 
    border-color: #E6C44D !important;
    box-shadow: 0 6px 14px rgba(212, 175, 55, 0.25) !important;
    transform: none !important;
}

/* Center Vertically & Right-Aligned View Meeting Button */
.view-btn-wrapper {
    display: flex !important;
    align-items: center !important;
    justify-content: flex-end !important;
    height: 100% !important;
    min-height: 80px !important;
}

/* Topbar Date Picker Trigger Styling */
div[data-testid="stPopover"] > button {
    background-color: #111A2B !important;
    color: #FFFFFF !important;
    border: 1px solid #D4AF37 !important;
    border-radius: 18px !important;
    font-size: 0.84rem !important;
    font-weight: 600 !important;
    height: 36px !important;
    box-shadow: 0 4px 10px rgba(26, 43, 76, 0.18) !important;
    width: 100% !important;
    text-align: left !important;
    justify-content: flex-start !important;
    padding: 0 0.85rem !important;
}

div[data-testid="stPopover"] > button:hover {
    border-color: #E6C44D !important;
    background-color: #1A2B4C !important;
    color: #FFFFFF !important;
    transform: none !important;
}

/* Popover Content Width for Split-Pane Date Picker */
div[data-testid="stPopoverBody"] {
    min-width: 560px !important;
    max-width: 600px !important;
    padding: 1.25rem !important;
    background-color: #FFFFFF !important;
    border-radius: 10px !important;
    box-shadow: 0 15px 35px rgba(0,0,0,0.18) !important;
}

/* Preset Buttons Inside Date Popover */
.stButton > button[key^="preset_"] {
    background-color: transparent !important;
    color: #4A5568 !important;
    border: 1px solid transparent !important;
    border-radius: 4px !important;
    font-size: 0.85rem !important;
    font-weight: 500 !important;
    text-align: left !important;
    justify-content: flex-start !important;
    padding: 0.45rem 0.65rem !important;
    height: 34px !important;
    margin-bottom: 0.35rem !important;
    box-shadow: none !important;
}

.stButton > button[key^="preset_"]:hover {
    background-color: #EDF2F7 !important;
    color: #1A202C !important;
    transform: none !important;
}

.stButton > button[key="btn_apply_modal_date"] {
    background-color: #111A2B !important;
    color: #FFFFFF !important;
    border: 1px solid #D4AF37 !important;
    border-radius: 18px !important;
    height: 36px !important;
    font-weight: 600 !important;
}
.stButton > button[key="btn_apply_modal_date"]:hover {
    background-color: #1A2B4C !important;
    border-color: #E6C44D !important;
    color: #FFFFFF !important;
    transform: none !important;
}

/* Back Button Pill */
.stButton > button[key="btn_back_gallery"] {
    background-color: transparent !important;
    color: #1A2B4C !important;
    border: 1px solid rgba(26, 43, 76, 0.3) !important;
    width: auto !important;
    min-width: 170px !important;
}
.stButton > button[key="btn_back_gallery"]:hover {
    background-color: #1A2B4C !important;
    color: #FFFFFF !important;
}

/* Details Action Buttons */
.stButton > button[key="btn_toggle_edit_details"], .stButton > button[key="btn_cancel_edit_details"] {
    background-color: #111A2B !important;
    color: #FFFFFF !important;
    border: 1px solid #D4AF37 !important;
    border-radius: 18px !important;
    height: 34px !important;
    font-weight: 600 !important;
}
.stButton > button[key="btn_toggle_edit_details"]:hover, .stButton > button[key="btn_cancel_edit_details"]:hover {
    background-color: #1A2B4C !important;
    border-color: #E6C44D !important;
    color: #FFFFFF !important;
}

/* Gallery Typography */
.card-title {
    font-family: 'Playfair Display', serif !important;
    font-style: italic !important;
    font-size: 1.25rem !important;
    color: #1A2B4C !important;
    margin: 0 0 0.25rem 0 !important;
    line-height: 1.3 !important;
}

.card-meta {
    font-size: 0.84rem !important;
    color: #666666 !important;
    margin-bottom: 0.55rem !important;
}

.card-desc {
    font-size: 0.88rem !important;
    color: #2D2D2D !important;
    line-height: 1.5 !important;
    margin: 0 !important;
}

/* Tabs Header */
button[data-baseweb="tab"] {
    background: transparent !important;
    border: none !important;
    font-size: 1.05rem !important;
    font-weight: 500 !important;
    color: #1A2B4C !important;
    padding: 0.5rem 1rem !important;
}

button[data-baseweb="tab"][aria-selected="true"] {
    color: #FF4B4B !important;
    border-bottom: 2px solid #FF4B4B !important;
}

div[data-baseweb="tab-highlight"] {
    background-color: #FF4B4B !important;
}

/* Delete Row SVG Button */
.stButton > button[key^="del_"] { 
    background-color: #FDF9F9 !important; 
    color: #B23A3A !important; 
    border: 1px solid rgba(178, 58, 58, 0.25) !important; 
}
.stButton > button[key^="del_"]:hover { 
    background-color: #B23A3A !important; 
    color: #FFFFFF !important; 
}
.stButton > button[key^="del_"]::before {
    content: "";
    display: inline-block;
    width: 14px;
    height: 14px;
    margin-right: 4px;
    background-color: currentColor;
    -webkit-mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M6 19c0 1.1.9 2 2 2h8c1.1 0 2-.9 2-2V7H6v12zM19 4h-3.5l-1-1h-5l-1 1H5v2h14V4z'/%3E%3C/svg%3E") no-repeat center;
    mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M6 19c0 1.1.9 2 2 2h8c1.1 0 2-.9 2-2V7H6v12zM19 4h-3.5l-1-1h-5l-1 1H5v2h14V4z'/%3E%3C/svg%3E") no-repeat center;
}

/* Save Icon */
.stButton > button[key^="btn_save_"]::before,
.stButton > button[key="btn_save_meta"]::before {
    content: "";
    display: inline-block;
    width: 15px;
    height: 15px;
    margin-right: 6px;
    background-color: currentColor;
    -webkit-mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M17 3H5c-1.11 0-2 .9-2 2v14c0 1.1.89 2 2 2h14c1.1 0 2-.9 2-2V7l-4-4zm-5 16c-1.66 0-3-1.34-3-3s1.34-3 3-3 3 1.34 3 3-1.34 3-3 3zm3-10H5V5h10v4z'/%3E%3C/svg%3E") no-repeat center;
    mask: url("data:image/svg+xml,%3Csvg xmlns='http://www.w3.org/2000/svg' viewBox='0 0 24 24'%3E%3Cpath d='M17 3H5c-1.11 0-2 .9-2 2v14c0 1.1.89 2 2 2h14c1.1 0 2-.9 2-2V7l-4-4zm-5 16c-1.66 0-3-1.34-3-3s1.34-3 3-3 3 1.34 3 3-1.34 3-3 3zm3-10H5V5h10v4z'/%3E%3C/svg%3E") no-repeat center;
}
</style>
"""
st.markdown(CUSTOM_CSS, unsafe_allow_html=True)

# -------------------------------------------------------------------------
# Document Generation Functions
# -------------------------------------------------------------------------
def set_cell_shading(cell, color_hex):
    shd = parse_xml(f'<w:shd xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main" w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shd)

def export_to_word_template_1(df, meeting_details, other_discussions):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.4)
        section.bottom_margin = Inches(0.4)
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(2)
    r_title = p_title.add_run("MINUTES OF THE MEETING")
    r_title.bold = True
    r_title.underline = True
    r_title.font.name = "Arial"
    r_title.font.size = Pt(11)

    company_target = meeting_details.get("external_attendees", [])
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "CLIENT"
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(12)
    r_sub = p_sub.add_run(f"PRIME PHILIPPINES & {primary_client_rep.upper()}")
    r_sub.bold = True
    r_sub.font.name = "Arial"
    r_sub.font.size = Pt(11)

    date_str = meeting_details.get("date", "____________")
    time_str = meeting_details.get("time_range", "")
    full_date = f"Date: {date_str}" + (f", {time_str}" if time_str.strip() else "")
    
    p_date = doc.add_paragraph(full_date)
    p_date.paragraph_format.space_after = Pt(2)
    for r in p_date.runs: r.font.name, r.font.size = "Arial", Pt(10)

    p_loc = doc.add_paragraph(f"Location: {meeting_details.get('location', '____________')}")
    p_loc.paragraph_format.space_after = Pt(2)
    for r in p_loc.runs: r.font.name, r.font.size = "Arial", Pt(10)

    prime_atts = meeting_details.get("prime_attendees", [])
    ext_atts = meeting_details.get("external_attendees", [])
    
    p_att = doc.add_paragraph()
    p_att.paragraph_format.space_after = Pt(2)
    p_att.paragraph_format.tab_stops.add_tab_stop(Inches(1.35), WD_TAB_ALIGNMENT.LEFT)
    r_att_label = p_att.add_run("Attended by:")
    r_att_label.font.name, r_att_label.font.size = "Arial", Pt(10)
    
    first_attendee = True
    for att in ext_atts:
        if not att.strip(): continue
        p = p_att if first_attendee else doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if not first_attendee: p.paragraph_format.left_indent = Inches(1.35)
        else: p.add_run("\t")
        comp_label = f", {meeting_details.get('company_name')}" if meeting_details.get('company_name') else ""
        r = p.add_run(f"{att}{comp_label}")
        r.font.name, r.font.size = "Arial", Pt(10)
        first_attendee = False

    for att in prime_atts:
        p = p_att if first_attendee else doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        if not first_attendee: p.paragraph_format.left_indent = Inches(1.35)
        else: p.add_run("\t")
        r = p.add_run(f"{att} – PRIME Philippines")
        r.font.name, r.font.size = "Arial", Pt(10)
        first_attendee = False

    p_line = doc.add_paragraph()
    p_line.paragraph_format.space_before = Pt(4)
    p_line.paragraph_format.space_after = Pt(6)
    r_line = p_line.add_run("_________________________________________________________________________________")
    r_line.font.name, r_line.font.color.rgb = "Arial", RGBColor(160, 160, 160)

    client_display = meeting_details.get('company_name', '').strip() or "the Client"
    p_intro = doc.add_paragraph(f"During the meeting held last {date_str}, PRIME Philippines, represented by the attendee/s shown above, met with {client_display} to discuss opportunities for collaboration.")
    p_intro.paragraph_format.space_after = Pt(10)
    for r in p_intro.runs: r.font.name, r.font.size = "Arial", Pt(9.5)

    table = doc.add_table(rows=len(df)+1, cols=4)
    table.alignment, table.style, table.autofit, table.allow_autofit = WD_TABLE_ALIGNMENT.CENTER, "Table Grid", False, False
    col_widths = [Inches(2.5), Inches(2.2), Inches(1.1), Inches(1.2)]
    headers = ["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"]
    
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.width, cell.text = col_widths[i], header
        set_cell_shading(cell, "FFFF00")
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs: p.runs[0].font.bold, p.runs[0].font.size, p.runs[0].font.name = True, Pt(9), "Arial"

    for i, row in df.iterrows():
        cells = table.rows[i+1].cells
        cells[0].text, cells[1].text, cells[2].text, cells[3].text = f"{i+1}. {str(row.get('Discussion Points', ''))}", str(row.get("Action Plan", "")), str(row.get("Indicative Delivery Date", "")), str(row.get("Person-in-charge", ""))
        for c_idx, cell in enumerate(cells):
            cell.width = col_widths[c_idx]
            p = cell.paragraphs[0]
            if c_idx in [2, 3]: p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs: p.runs[0].font.size, p.runs[0].font.name = Pt(8.5), "Arial"

    doc.add_paragraph()
    p_note = doc.add_paragraph("*Note: The indicative delivery date serves as reference point and still subject to changes. Furthermore, it depends on the progress of both parties.")
    p_note.paragraph_format.space_after = Pt(8)
    p_note.runs[0].italic, p_note.runs[0].font.name, p_note.runs[0].font.size = True, "Arial", Pt(8)

    if other_discussions.strip():
        p_od_head = doc.add_paragraph()
        p_od_head.paragraph_format.space_before, p_od_head.paragraph_format.space_after = Pt(6), Pt(4)
        r_od_head = p_od_head.add_run("Other Discussions:")
        r_od_head.bold, r_od_head.font.size, r_od_head.font.name = True, Pt(10), "Arial"
        p_od = doc.add_paragraph(other_discussions)
        p_od.paragraph_format.space_after = Pt(12)
        for r in p_od.runs: r.font.name, r.font.size = "Arial", Pt(9.5)

    p_prep_label = doc.add_paragraph("Prepared by:")
    p_prep_label.paragraph_format.space_before = Pt(12)
    p_prep_label.paragraph_format.space_after = Pt(2)
    p_prep_label.runs[0].font.name, p_prep_label.runs[0].font.bold, p_prep_label.runs[0].font.size = "Arial", True, Pt(9.5)
    p_prep_line = doc.add_paragraph("_______________________________")
    p_prep_line.paragraph_format.space_after = Pt(2)
    p_prep_line.runs[0].font.name = "Arial"
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"
    prep_desig = meeting_details.get("prep_desig", "").strip() or "PRIME Philippines"
    p_prep_info = doc.add_paragraph(f"{prep_name}\n{prep_desig}")
    p_prep_info.paragraph_format.space_after = Pt(12)
    for r in p_prep_info.runs: r.font.name, r.font.size = "Arial", Pt(9.5)

    p_conf_label = doc.add_paragraph("Confirmed by:")
    p_conf_label.paragraph_format.space_after = Pt(2)
    p_conf_label.runs[0].font.name, p_conf_label.runs[0].font.bold, p_conf_label.runs[0].font.size = "Arial", True, Pt(9.5)
    p_conf_line = doc.add_paragraph("_______________________________")
    p_conf_line.paragraph_format.space_after = Pt(2)
    p_conf_line.runs[0].font.name = "Arial"
    conf_name = meeting_details.get("conf_name", "").strip() or (ext_atts[0] if ext_atts else "____________________")
    conf_desig = meeting_details.get("conf_desig", "").strip() or (meeting_details.get("company_name", "").strip() or "Client")
    p_conf_info = doc.add_paragraph(f"{conf_name}\n{conf_desig}")
    p_conf_info.paragraph_format.space_after = Pt(6)
    for r in p_conf_info.runs: r.font.name, r.font.size = "Arial", Pt(9.5)

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_to_pdf_template_1(df, meeting_details, other_discussions):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.6 * inch, rightMargin=0.6 * inch, topMargin=0.5 * inch, bottomMargin=0.5 * inch)
    story, styles = [], getSampleStyleSheet()
    
    style_title = ParagraphStyle('DocTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=11.5, alignment=1, spaceAfter=2)
    company_target = meeting_details.get("external_attendees", [])
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "CLIENT"
    style_subtitle = ParagraphStyle('DocSubTitle', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=10.5, alignment=1, spaceAfter=10)
    style_body = ParagraphStyle('DocBody', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=12, spaceAfter=3)
    style_th = ParagraphStyle('TableHead', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8.5, leading=10, alignment=1)
    style_td = ParagraphStyle('TableData', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=10)
    style_td_center = ParagraphStyle('TableDataCenter', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=10, alignment=1)

    story.append(Paragraph("<u>MINUTES OF THE MEETING</u>", style_title))
    story.append(Paragraph(f"PRIME PHILIPPINES & {primary_client_rep.upper()}", style_subtitle))
    
    date_str, time_str = meeting_details.get("date", "____________"), meeting_details.get("time_range", "")
    full_date = f"<b>Date:</b> {date_str}" + (f", {time_str}" if time_str.strip() else "")
    story.append(Paragraph(full_date, style_body))
    story.append(Paragraph(f"<b>Location:</b> {meeting_details.get('location', '____________')}", style_body))
    
    prime_atts, ext_atts = meeting_details.get("prime_attendees", []), meeting_details.get("external_attendees", [])
    att_list = []
    for att in ext_atts:
        if att.strip(): att_list.append(f"{att}{f', {meeting_details.get('company_name')}' if meeting_details.get('company_name') else ''}")
    for att in prime_atts: att_list.append(f"{att} – PRIME Philippines")
    
    if att_list:
        story.append(Paragraph(f"<b>Attended by:</b>&nbsp;&nbsp;&nbsp;&nbsp;{att_list[0]}", style_body))
        for a in att_list[1:]: story.append(Paragraph(f"&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;&nbsp;{a}", style_body))
    
    story.append(Spacer(1, 4))
    client_display = meeting_details.get('company_name', '').strip() or "the Client"
    story.append(Paragraph(f"During the meeting held last {date_str}, PRIME Philippines, represented by the attendee/s shown above, met with {client_display} to discuss opportunities for collaboration.", style_body))
    story.append(Spacer(1, 6))
    
    table_data = [[Paragraph("<b>Discussion Points</b>", style_th), Paragraph("<b>Action Plan</b>", style_th), Paragraph("<b>Indicative Delivery Date</b>", style_th), Paragraph("<b>Person-in-charge</b>", style_th)]]
    for i, row in df.iterrows():
        table_data.append([
            Paragraph(f"{i+1}. {str(row.get('Discussion Points', ''))}", style_td),
            Paragraph(str(row.get("Action Plan", "")), style_td),
            Paragraph(str(row.get("Indicative Delivery Date", "")), style_td_center),
            Paragraph(str(row.get("Person-in-charge", "")), style_td_center)
        ])
    
    t = Table(table_data, colWidths=[2.4 * inch, 2.3 * inch, 1.1 * inch, 1.0 * inch], repeatRows=1)
    t.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#FFFF00')),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING', (0, 0), (-1, -1), 4),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 4)
    ]))
    story.append(t)
    
    note_style = ParagraphStyle('Note', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=7.5, leading=9, spaceBefore=4)
    story.append(Paragraph("*Note: The indicative delivery date serves as reference point and still subject to changes. Furthermore, it depends on the progress of both parties.", note_style))
    
    if other_discussions.strip():
        story.append(Spacer(1, 6))
        story.append(Paragraph("<b>Other Discussions:</b>", style_body))
        story.append(Paragraph(other_discussions, style_body))
    
    story.append(Spacer(1, 8))
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"
    prep_desig = meeting_details.get("prep_desig", "").strip() or "PRIME Philippines"
    conf_name = meeting_details.get("conf_name", "").strip() or (ext_atts[0] if ext_atts else "____________________")
    conf_desig = meeting_details.get("conf_desig", "").strip() or (meeting_details.get("company_name", "").strip() or "Client")
    
    sign_data = [
        [Paragraph("<b>Prepared by:</b>", style_body), Paragraph("<b>Confirmed by:</b>", style_body)],
        [Paragraph("_______________________________", style_body), Paragraph("_______________________________", style_body)],
        [Paragraph(f"{prep_name}<br/>{prep_desig}", style_body), Paragraph(f"{conf_name}<br/>{conf_desig}", style_body)]
    ]
    sign_table = Table(sign_data, colWidths=[3.4 * inch, 3.4 * inch])
    sign_table.setStyle(TableStyle([
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('LEFTPADDING', (0, 0), (-1, -1), 0),
        ('RIGHTPADDING', (0, 0), (-1, -1), 0),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 2)
    ]))
    story.append(sign_table)
    
    doc.build(story)
    buffer.seek(0)
    return buffer

def export_to_word_template_2(df, meeting_details, other_discussions):
    doc = Document()
    
    # Page setup - Standard 1 inch margins
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Base helper to create styled text matching the document
    def add_clean_para(text="", bold=False, italic=False, font_size=11, space_before=0, space_after=4, color_rgb=(0, 0, 0)):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(space_after)
        p.paragraph_format.line_spacing = 1.15
        if text:
            run = p.add_run(text)
            run.bold = bold
            run.italic = italic
            run.font.name = "Arial"
            run.font.size = Pt(font_size)
            run.font.color.rgb = RGBColor(*color_rgb)
        return p

    # 1. Title Block
    add_clean_para("Minutes of the Meeting", bold=True, font_size=16, space_after=0)
    
    company_name = meeting_details.get("company_name", "").strip() or "General Meeting"
    add_clean_para(company_name, bold=False, font_size=13, space_after=0)
    
    # Meeting topic / sub-headline
    topic = "Planning and Format Confirmation" if "sportsfest" in company_name.lower() else "General Alignment and Status Review"
    add_clean_para(topic, font_size=11, space_after=14)

    # 2. Meeting Details
    add_clean_para("Meeting Details", bold=True, font_size=12, space_before=10, space_after=4)
    
    date_str = meeting_details.get("date", "____________")
    time_str = meeting_details.get("time_range", "____________")
    location_str = meeting_details.get('location', '____________')
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"
    prep_desig = meeting_details.get("prep_desig", "PRIME Philippines").strip()
    
    details_map = [
        ("Date", date_str),
        ("Time", time_str),
        ("Venue", location_str),
        ("Prepared by", f"{prep_name}, {prep_desig}"),
        ("Date prepared", datetime.now().strftime("%B %d, %Y")),
        ("Source", "Meeting notes and transcript")
    ]
    
    table_details = doc.add_table(rows=len(details_map), cols=2)
    table_details.style = 'Table Grid'
    table_details.alignment = WD_TABLE_ALIGNMENT.LEFT
    table_details.autofit = False
    
    for i, (k, v) in enumerate(details_map):
        row = table_details.rows[i]
        c0, c1 = row.cells[0], row.cells[1]
        c0.width = Inches(2.0)
        c1.width = Inches(4.5)
        
        c0.text = k
        c1.text = v
        
        for c in [c0, c1]:
            p = c.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(10)

    add_clean_para("", space_after=8)

    # 3. Attendees
    add_clean_para("Attendees", bold=True, font_size=12, space_before=8, space_after=2)
    add_clean_para(f"PRIME Philippines — {company_name}", font_size=10, space_after=4)
    
    all_atts = meeting_details.get("prime_attendees", []) + meeting_details.get("external_attendees", [])
    if prep_name and prep_name not in all_atts:
        all_atts.insert(0, f"{prep_name} (minutes prepared by)")
    elif all_atts:
        all_atts[0] = f"{all_atts[0]} (minutes prepared by)"
        
    for att in all_atts:
        if att.strip():
            p_att = doc.add_paragraph()
            p_att.paragraph_format.space_before = Pt(0)
            p_att.paragraph_format.space_after = Pt(2)
            p_att.paragraph_format.left_indent = Inches(0.2)
            r = p_att.add_run(att.strip())
            r.font.name = "Arial"
            r.font.size = Pt(10)

    add_clean_para("", space_after=8)

    # 4. Purpose
    add_clean_para("Purpose", bold=True, font_size=12, space_before=8, space_after=2)
    purpose_text = other_discussions if other_discussions.strip() else "To review key discussion items, assign actionable deliverables, and confirm project milestones."
    add_clean_para(purpose_text, font_size=10, space_after=12)

    # 5. Decisions Agreed / Discussion Points
    add_clean_para("Decisions Agreed", bold=True, font_size=12, space_before=8, space_after=4)
    for i, row in df.iterrows():
        dp = str(row.get('Discussion Points', '')).strip()
        if dp:
            p_dp = doc.add_paragraph()
            p_dp.paragraph_format.space_before = Pt(2)
            p_dp.paragraph_format.space_after = Pt(2)
            r_num = p_dp.add_run(f"• ")
            r_num.font.name = "Arial"
            r_num.font.size = Pt(10)
            r_txt = p_dp.add_run(dp)
            r_txt.font.name = "Arial"
            r_txt.font.size = Pt(10)

    add_clean_para("", space_after=8)

    # 6. Action Plan Table
    add_clean_para("Action Plan", bold=True, font_size=12, space_before=8, space_after=4)
    
    act_table = doc.add_table(rows=len(df)+1, cols=4)
    act_table.style = 'Table Grid'
    act_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    act_table.autofit = False
    
    act_widths = [Inches(0.5), Inches(3.6), Inches(1.2), Inches(1.2)]
    act_headers = ["#", "Action", "Owner", "Deadline"]
    
    for idx, head_text in enumerate(act_headers):
        cell = act_table.rows[0].cells[idx]
        cell.width = act_widths[idx]
        cell.text = head_text
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        if idx in [0, 2, 3]:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if p.runs:
            p.runs[0].font.bold = True
            p.runs[0].font.name = "Arial"
            p.runs[0].font.size = Pt(9.5)

    for i, row in df.iterrows():
        row_cells = act_table.rows[i+1].cells
        row_cells[0].text = str(i+1)
        row_cells[1].text = str(row.get("Action Plan", ""))
        row_cells[2].text = str(row.get("Person-in-charge", ""))
        row_cells[3].text = str(row.get("Indicative Delivery Date", ""))
        
        for c_idx, cell in enumerate(row_cells):
            cell.width = act_widths[c_idx]
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(3)
            p.paragraph_format.space_after = Pt(3)
            if c_idx in [0, 2, 3]:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            if p.runs:
                p.runs[0].font.name = "Arial"
                p.runs[0].font.size = Pt(9)

    add_clean_para("", space_after=12)

    # 7. Document Footer Note
    p_foot = add_clean_para(
        f"Owners and deadlines are marked as to be confirmed where the meeting notes did not record them.\nPrepared for circulation to {company_name}. Please return corrections before this is treated as the agreed record.",
        font_size=8.5,
        italic=True,
        color_rgb=(100, 100, 100),
        space_before=14
    )

    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def export_to_pdf_template_2(df, meeting_details, other_discussions):
    buffer = BytesIO()
    doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.8 * inch, rightMargin=0.8 * inch, topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    story, styles = [], getSampleStyleSheet()
    
    style_title = ParagraphStyle('Title2', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=15, alignment=1, spaceAfter=2)
    style_subtitle = ParagraphStyle('SubTitle2', parent=styles['Normal'], fontName='Helvetica', fontSize=10.5, textColor=colors.HexColor("#64748B"), alignment=1, spaceAfter=20)
    style_h2 = ParagraphStyle('Heading2', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=12, textColor=colors.HexColor("#1A2B4C"), spaceBefore=12, spaceAfter=6)
    style_body = ParagraphStyle('Body2', parent=styles['Normal'], fontName='Helvetica', fontSize=9.5, leading=14, spaceAfter=4)
    style_th = ParagraphStyle('TH2', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=9, textColor=colors.white, alignment=1)
    style_td = ParagraphStyle('TD2', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=12)
    style_td_center = ParagraphStyle('TDC2', parent=styles['Normal'], fontName='Helvetica', fontSize=9, leading=12, alignment=1)
    
    company_target = meeting_details.get("external_attendees", [])
    primary_client_rep = company_target[0] if company_target else meeting_details.get("company_name", "").strip() or "General Meeting"

    story.append(Paragraph("MINUTES OF THE MEETING", style_title))
    story.append(Paragraph(f"Project / Client: {primary_client_rep}", style_subtitle))
    
    story.append(Paragraph("Meeting Details", style_h2))
    date_str = meeting_details.get("date", "____________")
    time_str = meeting_details.get("time_range", "____________")
    location_str = meeting_details.get('location', '____________')
    prep_name = meeting_details.get("prep_name", "").strip() or "____________________"
    
    details_data = [
        [Paragraph("<b>Date</b>", style_body), Paragraph(date_str, style_body)],
        [Paragraph("<b>Time</b>", style_body), Paragraph(time_str, style_body)],
        [Paragraph("<b>Venue</b>", style_body), Paragraph(location_str, style_body)],
        [Paragraph("<b>Prepared by</b>", style_body), Paragraph(prep_name, style_body)],
        [Paragraph("<b>Date prepared</b>", style_body), Paragraph(datetime.now().strftime("%B %d, %Y"), style_body)]
    ]
    t_details = Table(details_data, colWidths=[2 * inch, 4.5 * inch])
    t_details.setStyle(TableStyle([
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6)
    ]))
    story.append(t_details)
    story.append(Spacer(1, 10))

    story.append(Paragraph("Attendees", style_h2))
    all_atts = meeting_details.get("prime_attendees", []) + meeting_details.get("external_attendees", [])
    att_items = [ListItem(Paragraph(a.strip(), style_body)) for a in all_atts if a.strip()]
    if att_items:
        story.append(ListFlowable(att_items, bulletType='bullet'))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Purpose & Summary", style_h2))
    story.append(Paragraph(other_discussions if other_discussions.strip() else "To discuss project updates, ongoing deliverables, and establish clear action plans.", style_body))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Discussion Points", style_h2))
    dp_items = [ListItem(Paragraph(str(row.get('Discussion Points', '')), style_body)) for _, row in df.iterrows()]
    if dp_items:
        story.append(ListFlowable(dp_items, bulletType='1'))
    story.append(Spacer(1, 10))

    story.append(Paragraph("Action Plan", style_h2))
    act_data = [[Paragraph("<b>#</b>", style_th), Paragraph("<b>Action Plan</b>", style_th), Paragraph("<b>Owner</b>", style_th), Paragraph("<b>Deadline</b>", style_th)]]
    for i, row in df.iterrows():
        act_data.append([
            Paragraph(str(i+1), style_td_center),
            Paragraph(str(row.get("Action Plan", "")), style_td),
            Paragraph(str(row.get("Person-in-charge", "")), style_td_center),
            Paragraph(str(row.get("Indicative Delivery Date", "")), style_td_center)
        ])
    
    t_act = Table(act_data, colWidths=[0.4 * inch, 3.5 * inch, 1.3 * inch, 1.3 * inch], repeatRows=1)
    t_act.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor('#1A2B4C')),
        ('ALIGN', (0, 0), (-1, 0), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'TOP'),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5)
    ]))
    story.append(t_act)

    story.append(Spacer(1, 24))
    footer_style = ParagraphStyle('Footer2', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=8, textColor=colors.grey, alignment=1)
    story.append(Paragraph(f"Prepared for circulation to {primary_client_rep}. Please return corrections before this is treated as the agreed record.", footer_style))

    doc.build(story)
    buffer.seek(0)
    return buffer

# 4. Data Ingestion & Date Normalization
meetings = fetch_meeting_archives(limit=500)

if not meetings:
    st.info("No meeting records found in Supabase.")
    st.stop()

def parse_meeting_date(raw_date_str):
    if not raw_date_str:
        return None
    raw_s = str(raw_date_str).strip()[:10]
    for fmt in ("%Y-%m-%d", "%B %d, %Y", "%m/%d/%Y"):
        try:
            return datetime.strptime(raw_s, fmt).date()
        except ValueError:
            pass
    return None

def get_iso_date_str(meeting_item):
    parsed = parse_meeting_date(meeting_item.get("meeting_date", ""))
    return parsed.strftime("%Y-%m-%d") if parsed else ""

def categorize_meeting(meeting_item):
    client_name = str(meeting_item.get("client_name", "")).strip().lower()
    raw_payload = meeting_item.get("raw_payload", {}) or {}
    meeting_details = raw_payload.get("meeting_details", {}) if isinstance(raw_payload, dict) else {}
    external_atts = meeting_details.get("external_attendees", [])
    
    if "crd" in client_name:
        return "CRD Team Meetings"
    elif "internal" in client_name or "prime" in client_name or (not external_atts and not client_name):
        return "Internal Meetings"
    return "External Meetings"

# ==============================================================================
# MODE 1: FULL-SCREEN MEETING GALLERY
# ==============================================================================
if st.session_state["view_mode"] == "gallery":
    with st.container(border=True):
        st.markdown("<h3>Meeting Gallery</h3>", unsafe_allow_html=True)
        st.caption("Search across meeting topics, filter by category or date range, and review meetings.")
        
        # Filter Bar Layout
        f_c1, f_c2, f_c3, f_c4 = st.columns([4.2, 2.3, 2.5, 1.0])
        
        with f_c1:
            search_input = st.text_input(
                "Search",
                value=st.session_state["gal_search_q"],
                placeholder="Search by client, ID, topic, transcript, PIC...",
                label_visibility="collapsed",
                key="gal_search_input"
            )
            st.session_state["gal_search_q"] = search_input
            
        with f_c2:
            type_options = ["All Meetings", "Internal Meetings", "External Meetings", "CRD Team Meetings"]
            selected_type = st.selectbox(
                "Meeting Type",
                options=type_options,
                index=type_options.index(st.session_state["gal_type_f"]) if st.session_state["gal_type_f"] in type_options else 0,
                label_visibility="collapsed",
                key="gal_type_select"
            )
            st.session_state["gal_type_f"] = selected_type
            
        with f_c3:
            dr = st.session_state["gal_date_range"]
            if dr and len(dr) == 2:
                btn_label = f"{dr[0].strftime('%b %d, %Y')} — {dr[1].strftime('%b %d, %Y')} •"
            elif dr and len(dr) == 1:
                btn_label = f"{dr[0].strftime('%b %d, %Y')} •"
            else:
                btn_label = "All Dates •"

            with st.popover(btn_label, use_container_width=True):
                pop_left, pop_right = st.columns([1.1, 2.3], gap="medium")
                
                with pop_left:
                    st.markdown("<p style='font-size:0.75rem; color:#888; margin-bottom:0.4rem; text-transform:uppercase;'>Presets</p>", unsafe_allow_html=True)
                    
                    if st.button("This Week", key="preset_this_week"):
                        start_w = today - timedelta(days=today.weekday())
                        st.session_state["gal_date_range"] = (start_w, start_w + timedelta(days=6))
                        st.rerun()
                    if st.button("Last Week", key="preset_last_week"):
                        start_lw = today - timedelta(days=today.weekday() + 7)
                        st.session_state["gal_date_range"] = (start_lw, start_lw + timedelta(days=6))
                        st.rerun()
                    if st.button("This Month", key="preset_this_month"):
                        st.session_state["gal_date_range"] = (today.replace(day=1), today)
                        st.rerun()
                    if st.button("Last Month", key="preset_last_month"):
                        first_this = today.replace(day=1)
                        last_m_end = first_this - timedelta(days=1)
                        st.session_state["gal_date_range"] = (last_m_end.replace(day=1), last_m_end)
                        st.rerun()
                    if st.button("Clear", key="preset_clear"):
                        st.session_state["gal_date_range"] = ()
                        st.rerun()
                
                with pop_right:
                    picked_range = st.date_input(
                        "Custom Range",
                        value=st.session_state["gal_date_range"] if st.session_state["gal_date_range"] else None,
                        label_visibility="collapsed",
                        key="modal_date_picker"
                    )
                    
                    st.write("<div style='height: 10px;'></div>", unsafe_allow_html=True)
                    app_c1, app_c2 = st.columns([5, 5])
                    with app_c2:
                        if st.button("Apply", key="btn_apply_modal_date"):
                            if isinstance(picked_range, tuple):
                                st.session_state["gal_date_range"] = picked_range
                            elif isinstance(picked_range, date):
                                st.session_state["gal_date_range"] = (picked_range, picked_range)
                            st.rerun()

        with f_c4:
            if st.button("Reset", key="btn_reset_all_filters"):
                st.session_state["gal_search_q"] = ""
                st.session_state["gal_type_f"] = "All Meetings"
                st.session_state["gal_date_range"] = ()
                st.rerun()

        # Filtering Logic Execution
        filtered_meetings = []
        q_clean = st.session_state["gal_search_q"].strip().lower()
        active_type = st.session_state["gal_type_f"]
        active_dr = st.session_state["gal_date_range"]

        for m in meetings:
            if active_type != "All Meetings":
                if categorize_meeting(m) != active_type:
                    continue
            
            if active_dr:
                m_date_obj = parse_meeting_date(m.get("meeting_date", ""))
                if not m_date_obj:
                    continue
                if len(active_dr) == 1 and m_date_obj != active_dr[0]:
                    continue
                elif len(active_dr) == 2 and not (active_dr[0] <= m_date_obj <= active_dr[1]):
                    continue
            
            if q_clean:
                searchable_corpus = " ".join([
                    str(m.get("client_name", "")),
                    str(m.get("meeting_id", "")),
                    str(m.get("meeting_date", "")),
                    str(m.get("location", "")),
                    str(m.get("prepared_by", "")),
                    str(m.get("confirmed_by", "")),
                    str(m.get("transcript_md", "")),
                    str(m.get("summary_md", "")),
                    str(m.get("table_items", ""))
                ]).lower()
                if q_clean not in searchable_corpus:
                    continue
            
            filtered_meetings.append(m)

        is_filtered = bool(q_clean or active_type != "All Meetings" or active_dr)
        if is_filtered:
            st.caption(f"Showing **{len(filtered_meetings)}** matching meeting archive(s)")
        else:
            st.caption(f"Showing all **{len(filtered_meetings)}** meeting archive(s)")

        st.markdown("<hr style='margin: 0.5rem 0 1rem 0; border: none; border-top: 1px solid rgba(0,0,0,0.06);'>", unsafe_allow_html=True)

        if not filtered_meetings:
            st.warning("No meeting records matched your search parameters.")
        else:
            for idx, m in enumerate(filtered_meetings):
                m_id_val = m.get("meeting_id", f"MOM-{idx}")
                client_lbl = m.get("client_name") or "Meeting Record"
                d_val = get_iso_date_str(m) or "____________"
                loc_val = m.get("location") or "____________"
                prep_val = m.get("prepared_by") or "CRD Team"
                
                summary_raw = str(m.get("summary_md", "")).replace("### Summary", "").strip()
                if not summary_raw:
                    summary_raw = "No summary recorded. Minutes generated and stored in Supabase archive."
                preview_text = summary_raw[:220] + ("..." if len(summary_raw) > 220 else "")

                with st.container(border=True):
                    c_info, c_act = st.columns([8.2, 1.8])
                    with c_info:
                        st.markdown(f"<p class='card-title'>{client_lbl}</p>", unsafe_allow_html=True)
                        st.markdown(f"<p class='card-meta'>Date: {d_val} &bull; {loc_val} &bull; Prepared by: {prep_val}</p>", unsafe_allow_html=True)
                        st.markdown(f"<p class='card-desc'>{preview_text}</p>", unsafe_allow_html=True)
                    with c_act:
                        st.markdown('<div class="view-btn-wrapper" style="flex-direction:column; gap:6px; justify-content:center;">', unsafe_allow_html=True)
                        if st.button("View Meeting", key=f"view_btn_{m_id_val}_{idx}", use_container_width=True):
                            st.session_state["selected_meeting_id"] = m_id_val
                            st.session_state["view_mode"] = "details"
                            st.session_state["edit_meeting_details"] = False
                            st.rerun()
                        st.markdown('</div>', unsafe_allow_html=True)

# ==============================================================================
# MODE 2: FULL-SCREEN MEETING VIEWER & INSPECTOR
# ==============================================================================
elif st.session_state["view_mode"] == "details":
    target_id = st.session_state.get("selected_meeting_id")
    active_meeting = next((m for m in meetings if m.get("meeting_id") == target_id), None)

    if not active_meeting:
        st.session_state["view_mode"] = "gallery"
        st.rerun()

    m_id = active_meeting.get("meeting_id")

    # Header Navigation
    top_nav1, top_nav2 = st.columns([2.5, 7.5])
    with top_nav1:
        if st.button("← Back to Gallery", key="btn_back_gallery"):
            st.session_state["view_mode"] = "gallery"
            st.session_state["edit_meeting_details"] = False
            st.rerun()

    # Editable Meeting Metadata Card
    with st.container(border=True):
        m_head1, m_head2 = st.columns([7.5, 2.5])
        with m_head1:
            st.markdown(f"<h3>{active_meeting.get('client_name', 'Client Meeting')}</h3>", unsafe_allow_html=True)
            st.caption(f"Meeting ID: `{m_id}`")
        with m_head2:
            st.write("<div style='height: 4px;'></div>", unsafe_allow_html=True)
            if not st.session_state["edit_meeting_details"]:
                if st.button("Edit Meeting Details", key="btn_toggle_edit_details"):
                    st.session_state["edit_meeting_details"] = True
                    st.rerun()
            else:
                if st.button("Cancel Edit", key="btn_cancel_edit_details"):
                    st.session_state["edit_meeting_details"] = False
                    st.rerun()

        if not st.session_state["edit_meeting_details"]:
            md_ro = active_meeting.get("raw_payload", {})
            md_ro = md_ro.get("meeting_details", {}) if isinstance(md_ro, dict) else {}
            mtype_ro = str(active_meeting.get("meeting_type", "") or md_ro.get("meeting_type", "N/A"))
            tr_ro = str(md_ro.get("time_range", "") or "N/A")
            d_r1_c1, d_r1_c2 = st.columns(2)
            with d_r1_c1:
                st.write(f"**Date:** {active_meeting.get('meeting_date', 'N/A')}")
                st.write(f"**Meeting Type:** {mtype_ro}")
                st.write(f"**Time:** {tr_ro}")
                st.write(f"**Prepared By:** {active_meeting.get('prepared_by', 'N/A')}")
            with d_r1_c2:
                st.write(f"**Location:** {active_meeting.get('location', 'N/A')}")
                st.write(f"**Confirmed By:** {active_meeting.get('confirmed_by', 'N/A')}")
        else:
            # Richer editing: load existing meeting_details (raw_payload)
            rp_ed = active_meeting.get("raw_payload", {})
            md_ed = rp_ed.get("meeting_details", {}) if isinstance(rp_ed, dict) else {}

            # Start / End times from time_range ("1:00 AM to 2:00 PM")
            start_t, end_t = _parse_time_range(str(md_ed.get("time_range", "") or ""))

            st.markdown("##### Meeting Details")
            e_r1_c1, e_r1_c2 = st.columns(2)
            with e_r1_c1:
                edit_client = st.text_input("Client / Company / Department", value=str(active_meeting.get("client_name", "")), key=f"e_client_{m_id}")
                edit_date = st.text_input("Date", value=str(active_meeting.get("meeting_date", "")), key=f"e_date_{m_id}")
                edit_prep = st.text_input("Prepared By", value=str(active_meeting.get("prepared_by", "") or md_ed.get("prep_name", "")), key=f"e_prep_{m_id}")
                edit_prep_desig = st.text_input("Prep Designation", value=str(md_ed.get("prep_desig", "")), key=f"e_prep_desig_{m_id}")
            with e_r1_c2:
                edit_loc = st.text_input("Location", value=str(active_meeting.get("location", "")), key=f"e_loc_{m_id}")
                mtype_opts = ["Internal", "External", "Team"]
                cur_type = str(active_meeting.get("meeting_type", "") or md_ed.get("meeting_type", "Internal"))
                type_idx = mtype_opts.index(cur_type) if cur_type in mtype_opts else 0
                edit_type = st.selectbox("Meeting Type", options=mtype_opts, index=type_idx, key=f"e_type_{m_id}")
                edit_conf = st.text_input("Confirmed By", value=str(active_meeting.get("confirmed_by", "") or md_ed.get("conf_name", "")), key=f"e_conf_{m_id}")
                edit_conf_desig = st.text_input("Conf Designation", value=str(md_ed.get("conf_desig", "")), key=f"e_conf_desig_{m_id}")

            st.markdown("##### Schedule")
            s_c1, s_c2 = st.columns(2)
            with s_c1:
                edit_start = st.time_input("Start Time", value=start_t or time(9, 0), key=f"e_start_{m_id}")
            with s_c2:
                edit_end = st.time_input("End Time", value=end_t or time(10, 0), key=f"e_end_{m_id}")

            st.markdown("##### Attendees")
            a_c1, a_c2 = st.columns(2)
            with a_c1:
                edit_crd = st.multiselect(
                    "CRD Team Attendees",
                    options=CRD_MEMBERS,
                    default=[str(x) for x in (md_ed.get("prime_attendees") or [])],
                    key=f"e_crd_{m_id}",
                )
            with a_c2:
                edit_ext = st.text_area(
                    "External Attendees",
                    value=", ".join(str(x) for x in (md_ed.get("external_attendees") or [])),
                    height=90,
                    key=f"e_ext_{m_id}",
                )

            st.write("")
            sm_c1, sm_c2 = st.columns([7.8, 2.2])
            with sm_c2:
                if st.button("Save Meeting Details", key="btn_save_meta"):
                    with st.spinner("Saving metadata to Supabase..."):
                        client = get_supabase_client()
                        if not client:
                            st.error("Supabase client uninitialized.")
                        else:
                            try:
                                ext_list = [x.strip() for x in edit_ext.split(",") if x.strip()]
                                time_range_str = f"{edit_start.strftime('%I:%M %p')} to {edit_end.strftime('%I:%M %p')}"

                                new_md = dict(md_ed)
                                new_md.update({
                                    "date": edit_date.strip(),
                                    "time_range": time_range_str,
                                    "meeting_type": edit_type,
                                    "location": edit_loc.strip(),
                                    "company_name": edit_client.strip(),
                                    "prime_attendees": list(edit_crd),
                                    "external_attendees": ext_list,
                                    "prep_name": edit_prep.strip(),
                                    "prep_desig": edit_prep_desig.strip(),
                                    "conf_name": edit_conf.strip(),
                                    "conf_desig": edit_conf_desig.strip(),
                                })
                                new_raw = dict(rp_ed) if isinstance(rp_ed, dict) else {}
                                new_raw["meeting_details"] = new_md

                                client.table("meeting_archives").update({
                                    "client_name": edit_client.strip(),
                                    "meeting_date": edit_date.strip(),
                                    "meeting_type": edit_type,
                                    "location": edit_loc.strip(),
                                    "prepared_by": edit_prep.strip(),
                                    "confirmed_by": edit_conf.strip(),
                                    "raw_payload": new_raw,
                                }).eq("meeting_id", m_id).execute()

                                active_meeting["client_name"] = edit_client.strip()
                                active_meeting["meeting_date"] = edit_date.strip()
                                active_meeting["meeting_type"] = edit_type
                                active_meeting["location"] = edit_loc.strip()
                                active_meeting["prepared_by"] = edit_prep.strip()
                                active_meeting["confirmed_by"] = edit_conf.strip()
                                active_meeting["raw_payload"] = new_raw

                                st.session_state["edit_meeting_details"] = False
                                st.success("Meeting details updated successfully!")
                                st.rerun()
                            except Exception as e:
                                st.error(f"Metadata update failed: {e}")

    # Tabs
    tab_editor, tab_transcript = st.tabs(["Minutes of Meeting Editor", "Full Transcript"])

    with tab_editor:
        with st.container(border=True):
            st.markdown("<h3>Minutes of Meeting Items</h3>", unsafe_allow_html=True)
            st.caption("Inline editable cards. Changes are synchronized directly to Supabase.")

            editor_key = f"mom_rows_{m_id}"
            if editor_key not in st.session_state:
                raw_items = active_meeting.get("table_items", [])
                if isinstance(raw_items, list) and len(raw_items) > 0:
                    st.session_state[editor_key] = raw_items
                else:
                    st.session_state[editor_key] = [{
                        "Discussion Points": "", "Action Plan": "",
                        "Indicative Delivery Date": "", "Person-in-charge": ""
                    }]

            rows = st.session_state[editor_key]
            rows_to_keep = []

            for idx, row in enumerate(rows):
                with st.container(border=True):
                    c_disc, c_act, c_date, c_pic, c_del = st.columns([3.2, 3.2, 1.8, 1.8, 0.6])

                    with c_disc:
                        st.markdown('<span class="playfair-label">Discussion Points</span>', unsafe_allow_html=True)
                        st.text_area("DP", value=str(row.get("Discussion Points", "")), key=f"dp_{m_id}_{idx}", height=75, label_visibility="collapsed")
                    with c_act:
                        st.markdown('<span class="playfair-label">Action Plan</span>', unsafe_allow_html=True)
                        st.text_area("AP", value=str(row.get("Action Plan", "")), key=f"ap_{m_id}_{idx}", height=75, label_visibility="collapsed")
                    with c_date:
                        st.markdown('<span class="playfair-label">Delivery Date</span>', unsafe_allow_html=True)
                        st.text_area("DD", value=str(row.get("Indicative Delivery Date", "")), key=f"date_{m_id}_{idx}", height=75, label_visibility="collapsed")
                    with c_pic:
                        st.markdown('<span class="playfair-label">Person-in-charge</span>', unsafe_allow_html=True)
                        st.text_area("PIC", value=str(row.get("Person-in-charge", "")), key=f"pic_{m_id}_{idx}", height=75, label_visibility="collapsed")
                    with c_del:
                        st.write("<div style='height: 38px;'></div>", unsafe_allow_html=True)
                        if st.button("Delete", key=f"del_{m_id}_{idx}", help="Delete Row"):
                            continue

                    rows_to_keep.append({
                        "Discussion Points": st.session_state[f"dp_{m_id}_{idx}"],
                        "Action Plan": st.session_state[f"ap_{m_id}_{idx}"],
                        "Indicative Delivery Date": st.session_state[f"date_{m_id}_{idx}"],
                        "Person-in-charge": st.session_state[f"pic_{m_id}_{idx}"]
                    })

            if len(rows_to_keep) != len(rows):
                st.session_state[editor_key] = rows_to_keep
                st.rerun()

            add_c1, _ = st.columns([2, 8])
            with add_c1:
                if st.button("+ Add Item", key=f"btn_add_{m_id}"):
                    rows_to_keep.append({
                        "Discussion Points": "", "Action Plan": "",
                        "Indicative Delivery Date": "", "Person-in-charge": ""
                    })
                    st.session_state[editor_key] = rows_to_keep
                    st.rerun()

            st.markdown('<span class="playfair-label" style="margin-top:0.75rem;">Summary & Other Discussions</span>', unsafe_allow_html=True)
            current_summary = str(active_meeting.get("summary_md", "")).replace("### Summary", "").strip()
            summary_val = st.text_area(
                "Summary Content",
                value=current_summary,
                height=110,
                label_visibility="collapsed",
                key=f"summary_{m_id}"
            )

            df_export = pd.DataFrame(rows_to_keep, columns=["Discussion Points", "Action Plan", "Indicative Delivery Date", "Person-in-charge"])
            raw_payload = active_meeting.get("raw_payload", {})
            md = raw_payload.get("meeting_details", {}) if isinstance(raw_payload, dict) else {}
            comp_name = active_meeting.get("client_name", "")
            
            meeting_details = {
                "date": active_meeting.get("meeting_date", ""),
                "time_range": md.get("time_range", ""),
                "location": active_meeting.get("location", ""),
                "company_name": comp_name,
                "prime_attendees": md.get("prime_attendees", []),
                "external_attendees": md.get("external_attendees", []),
                "prep_name": active_meeting.get("prepared_by", ""),
                "prep_desig": md.get("prep_desig", "PRIME Philippines"),
                "conf_name": active_meeting.get("confirmed_by", ""),
                "conf_desig": md.get("conf_desig", "Client")
            }

            st.markdown('<span class="playfair-label" style="margin-top:1.5rem;">Export Options</span>', unsafe_allow_html=True)
            template_selection = st.selectbox(
                "Select MoM Template Format",
                options=["Template 1 - Standard Corporate (Combined Table)", "Template 2 - Detailed General Meeting (Vertical Layout)"],
                label_visibility="collapsed",
                key=f"tpl_sel_{m_id}"
            )

            exp_col1, exp_col2 = st.columns(2)
            if "Template 1" in template_selection:
                with exp_col1:
                    doc_bio = export_to_word_template_1(df_export, meeting_details, summary_val)
                    st.download_button(label="Download Word Document (.docx)", data=doc_bio, file_name=f"MOM_{comp_name.replace(' ', '_') if comp_name else 'Report'}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"btn_dl_docx_1_{m_id}")
                with exp_col2:
                    pdf_bio = export_to_pdf_template_1(df_export, meeting_details, summary_val)
                    st.download_button(label="Download PDF Document (.pdf)", data=pdf_bio, file_name=f"MOM_{comp_name.replace(' ', '_') if comp_name else 'Report'}.pdf", mime="application/pdf", key=f"btn_dl_pdf_1_{m_id}")
            else:
                with exp_col1:
                    doc_bio = export_to_word_template_2(df_export, meeting_details, summary_val)
                    st.download_button(label="Download Word Document (.docx)", data=doc_bio, file_name=f"MOM_Detailed_{comp_name.replace(' ', '_') if comp_name else 'Report'}.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", key=f"btn_dl_docx_2_{m_id}")
                with exp_col2:
                    pdf_bio = export_to_pdf_template_2(df_export, meeting_details, summary_val)
                    st.download_button(label="Download PDF Document (.pdf)", data=pdf_bio, file_name=f"MOM_Detailed_{comp_name.replace(' ', '_') if comp_name else 'Report'}.pdf", mime="application/pdf", key=f"btn_dl_pdf_2_{m_id}")

            st.write("")
            sv_col1, sv_col2 = st.columns([7.5, 2.5])
            with sv_col2:
                if st.button("Save All Changes", key=f"btn_save_{m_id}"):
                    with st.spinner("Saving updates to Supabase..."):
                        client = get_supabase_client()
                        if not client:
                            st.error("Supabase client uninitialized.")
                        else:
                            try:
                                client.table("meeting_archives").update({
                                    "table_items": rows_to_keep,
                                    "summary_md": f"### Summary\n{summary_val}"
                                }).eq("meeting_id", m_id).execute()

                                st.success("Meeting record updated successfully!")
                                if editor_key in st.session_state:
                                    del st.session_state[editor_key]
                            except Exception as e:
                                st.error(f"Update failed: {e}")

    with tab_transcript:
        with st.container(border=True):
            raw_tx = active_meeting.get("transcript_md", "No transcript stored.")
            clean_tx = raw_tx.replace("### Transcript", "").strip()
            st.text_area(
                "Transcript Stream",
                value=clean_tx,
                height=520,
                disabled=True,
                label_visibility="collapsed"
            )
