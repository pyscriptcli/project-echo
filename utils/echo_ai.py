import streamlit as st
import requests
import json
import re
import base64
import io
import pandas as pd
from pypdf import PdfReader
from PIL import Image
from datetime import datetime
import docx
from rapidfuzz import process, fuzz
from utils.db import fetch_meeting_archives, fetch_echo_context, upsert_echo_context
from utils.skills import load_prompt
from utils.auth import get_current_user
from utils.agent import TOOLS, check_agent_access
from utils.limits import check_rate_limit, record_usage, token_balance


def _cur_user_id():
    user = get_current_user()
    return user["id"] if user and user.get("id") else None


def _cur_username():
    user = get_current_user()
    return str(user.get("username") or "").strip() if user else ""


def _fetch_all_tasks():
    """Return all tasks (assignee/status/due/meeting) for deliverables matching."""
    client = get_supabase_client_or_none()
    if not client:
        return []
    try:
        resp = client.table("tasks").select("title,status,assignee,due_date,meeting_id").execute()
        return resp.data if resp.data else []
    except Exception:  # noqa: BLE001
        return []


def get_supabase_client_or_none():
    from utils.db import get_supabase_client
    try:
        return get_supabase_client()
    except Exception:  # noqa: BLE001
        return None


def build_user_deliverables_context(username: str, archive_records: list) -> str:
    """Build a 'this week' deliverables block for the given user, from minutes + tasks.

    Matches the username (case-insensitive) to minutes action items' Person-in-charge
    and to tasks' assignee, so Echo can answer 'what are my deliverables this week'.
    """
    if not username:
        return ""
    name = username.lower()
    today = datetime.datetime.today().date()
    week_start = today - datetime.timedelta(days=today.weekday())
    week_end = week_start + datetime.timedelta(days=6)

    dm = []  # deliverables from minutes
    for m in archive_records or []:
        for item in (m.get("table_items") or []):
            pic = str(item.get("Person-in-charge") or item.get("person_in_charge") or "").strip()
            if pic and name in pic.lower():
                dm.append({
                    "source": "Minutes",
                    "client": str(m.get("client_name") or "")[:40],
                    "date": str(m.get("meeting_date") or "")[:10],
                    "point": str(item.get("Discussion Points") or item.get("discussion_point") or "")[:120],
                    "action": str(item.get("Action Plan") or item.get("action_plan") or "")[:120],
                    "due": str(item.get("Indicative Delivery Date") or item.get("indicative_delivery_date") or "TBD")[:32],
                    "pic": pic,
                })

    dt = []  # deliverables from tasks
    for t in _fetch_all_tasks() or []:
        asg = str(t.get("assignee") or "").strip()
        due_raw = t.get("due_date")
        due_date = None
        try:
            due_date = datetime.datetime.fromisoformat(str(due_raw)[:10]).date() if due_raw else None
        except Exception:  # noqa: BLE001
            due_date = None
        if asg and name in asg.lower():
            dt.append({
                "source": "Task",
                "title": str(t.get("title") or "")[:120],
                "status": str(t.get("status") or ""),
                "due": str(due_raw or "TBD")[:12],
                "in_week": bool(due_date and week_start <= due_date <= week_end),
            })

    if not dm and not dt:
        return f"CURRENT USER: {username}\nNo deliverables matched for this user in this week's data.\n"

    lines = [f"CURRENT USER: {username} (match tasks/minutes to this name)", "THIS WEEK'S DELIVERABLES:"]
    for d in (dm + dt)[:20]:
        if d["source"] == "Minutes":
            lines.append(f"- [Minutes] {d['client']} ({d['date']}) | {d['point']} | Action: {d['action']} | PIC: {d['pic']} | Due: {d['due']}")
        else:
            flag = " (this week)" if d.get("in_week") else ""
            lines.append(f"- [Task] {d['title']} | Status: {d['status']} | Due: {d['due']}{flag}")
    return "\n".join(lines)


# --- Pure SVG Icon Assets ---
SVG_ECHO_LOGO = """
<svg xmlns="http://www.w3.org/2000/svg" width="18" height="18" viewBox="0 0 24 24" fill="none" stroke="#0D1B3E" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align: middle;">
    <polygon points="12 2 2 7 12 12 22 7 12 2"></polygon>
    <polyline points="2 17 12 22 22 17"></polyline>
    <polyline points="2 12 12 17 22 12"></polyline>
</svg>
"""

SVG_USER_ICON = """
<svg xmlns="http://www.w3.org/2000/svg" width="12" height="12" viewBox="0 0 24 24" fill="none" stroke="#0D1B3E" stroke-width="2" stroke-linecap="round" stroke-linejoin="round">
    <path d="M20 21v-2a4 4 0 0 0-4-4H8a4 4 0 0 0-4 4v2"></path>
    <circle cx="12" cy="7" r="4"></circle>
</svg>
"""

SVG_GLOBE_ICON = """
<svg xmlns="http://www.w3.org/2000/svg" width="11" height="11" viewBox="0 0 24 24" fill="none" stroke="#0D1B3E" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align: middle; margin-right: 3px;">
    <circle cx="12" cy="12" r="10"></circle>
    <line x1="2" y1="12" x2="22" y2="12"></line>
    <path d="M12 2a15.3 15.3 0 0 1 4 10 15.3 15.3 0 0 1-4 10 15.3 15.3 0 0 1-4-10 15.3 15.3 0 0 1 4-10z"></path>
</svg>
"""

SVG_BRAIN_ICON = """
<svg xmlns="http://www.w3.org/2000/svg" width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#0D1B3E" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align: middle; margin-right: 4px;">
    <path d="M9.5 2A2.5 2.5 0 0 1 12 4.5v15a2.5 2.5 0 0 1-4.96.44 2.5 2.5 0 0 1-2.96-3.08 3 3 0 0 1-.34-5.58 2.5 2.5 0 0 1 1.32-4.24 2.5 2.5 0 0 1 4.44-2.04z"></path>
    <path d="M14.5 2A2.5 2.5 0 0 0 12 4.5v15a2.5 2.5 0 0 0 4.96.44 2.5 2.5 0 0 0 2.96-3.08 3 3 0 0 0 .34-5.58 2.5 2.5 0 0 0-1.32-4.24 2.5 2.5 0 0 0-4.44-2.04z"></path>
</svg>
"""

SVG_ALERT_ICON = """
<svg xmlns="http://www.w3.org/2000/svg" width="14" height="14" viewBox="0 0 24 24" fill="none" stroke="#854D0E" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align: middle; margin-right: 4px;">
    <path d="m21.73 18-8-14a2 2 0 0 0-3.48 0l-8 14A2 2 0 0 0 4 21h16a2 2 0 0 0 1.73-3Z"></path>
    <line x1="12" y1="9" x2="12" y2="13"></line>
    <line x1="12" y1="17" x2="12.01" y2="17"></line>
</svg>
"""

SVG_FILE_ICON = """
<svg xmlns="http://www.w3.org/2000/svg" width="13" height="13" viewBox="0 0 24 24" fill="none" stroke="#0D1B3E" stroke-width="2" stroke-linecap="round" stroke-linejoin="round" style="vertical-align: middle; margin-right: 3px;">
    <path d="M14 2H6a2 2 0 0 0-2 2v16a2 2 0 0 0 2 2h12a2 2 0 0 0 2-2V8z"></path>
    <polyline points="14 2 14 8 20 8"></polyline>
</svg>
"""

MODEL_REGISTRY = {
    "⚡ Fast - deepseek chat": "deepseek/deepseek-chat",
    "🧠 Thinking - deepseek reasoning": "deepseek/deepseek-reasoner",
    "👁 Vision - qwen vl": "qwen/qwen2.5-vl-72b-instruct"
}

ALLOWED_ATTACHMENT_TYPES = ["png", "jpg", "jpeg", "webp", "pdf", "docx", "doc", "txt", "csv"]

CHAT_COMPACT_ALIGNED_CSS = """
<style>
@import url('https://fonts.googleapis.com/css2?family=Playfair+Display:ital,wght@1,500;1,600&family=Inter:wght@400;500;600&display=swap');

html, body, [data-testid="stAppViewContainer"], .main, .block-container {
    overflow-y: auto !important;
    overflow-x: hidden !important;
    scrollbar-width: none !important;
    -ms-overflow-style: none !important;
    padding-top: 0.3rem !important;
    padding-bottom: 0.3rem !important;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
}

html::-webkit-scrollbar, 
body::-webkit-scrollbar, 
[data-testid="stAppViewContainer"]::-webkit-scrollbar, 
.main::-webkit-scrollbar, 
.block-container::-webkit-scrollbar,
div[data-testid="stVerticalBlockBorderWrapper"]::-webkit-scrollbar,
.echo-chat-box-container div[data-testid="stVerticalBlockBorderWrapper"]::-webkit-scrollbar {
    display: none !important;
    width: 0 !important;
    height: 0 !important;
    background: transparent !important;
}

div[data-testid="stVerticalBlockBorderWrapper"]:has(.echo-main-card-scope) {
    background-color: transparent !important;
    border: 1px solid rgba(0, 0, 0, 0.08) !important;
    border-radius: 8px !important;
    padding: 0 !important;
    box-shadow: none !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    scrollbar-width: none !important;
    -ms-overflow-style: none !important;
    max-width: 960px !important;
    margin: 0 auto !important;
    height: calc(100vh - 120px) !important;
    max-height: calc(100vh - 120px) !important;
}

div[data-testid="stVerticalBlockBorderWrapper"]:has(.echo-main-card-scope) > div[data-testid="stVerticalBlock"] {
    display: flex !important;
    flex-direction: column !important;
    height: 100% !important;
    max-height: 100% !important;
    padding: 0.45rem 0.75rem !important;
    gap: 0 !important;
    box-sizing: border-box !important;
}

.echo-header-bar {
    display: flex;
    align-items: center;
    gap: 8px;
    height: 34px;
    border-bottom: 1px solid rgba(212, 175, 55, 0.25);
    padding-bottom: 4px;
    margin-bottom: 4px;
    flex-shrink: 0 !important;
}

.echo-title {
    font-family: 'Playfair Display', Georgia, serif !important;
    font-style: italic !important;
    font-size: 1.15rem !important;
    font-weight: 600 !important;
    color: #1A2B4C !important;
    margin: 0 !important;
    line-height: 1 !important;
    letter-spacing: 0.01em !important;
}

.echo-top-controls div[data-testid="stPopover"] > button,
.echo-top-controls div[data-testid="stButton"] > button {
    background-color: #a3acd5 !important;
    color: #0D1B3E !important;
    border: 1px solid #a3acd5 !important;
    border-radius: 0 !important;
    height: 28px !important;
    min-height: 28px !important;
    padding: 0 0.55rem !important;
    transition: all 0.2s ease !important;
    display: inline-flex !important;
    align-items: center !important;
    justify-content: center !important;
    font-size: 0.85rem !important;
}

.echo-top-controls div[data-testid="stPopover"] > button:hover,
.echo-top-controls div[data-testid="stButton"] > button:hover {
    border-color: #8fa2d6 !important;
    background-color: #8fa2d6 !important;
    box-shadow: none !important;
}

.echo-chat-box-container {
    flex: 1 1 auto !important;
    min-height: 0 !important;
    overflow: hidden !important;
    display: flex !important;
    flex-direction: column !important;
}

.echo-chat-box-container div[data-testid="stVerticalBlockBorderWrapper"] {
    background: rgba(255, 255, 255, 0.8) !important;
    backdrop-filter: blur(4px) !important;
    -webkit-backdrop-filter: blur(4px) !important;
    border: 1px solid rgba(0, 0, 0, 0.06) !important;
    border-radius: 6px !important;
    overflow-y: auto !important;
    overflow-x: hidden !important;
    scrollbar-width: none !important;
    -ms-overflow-style: none !important;
    padding: 0.5rem 0.75rem !important;
    height: 100% !important;
}

.echo-msg-row-user {
    display: flex;
    justify-content: flex-end;
    align-items: flex-start;
    gap: 6px;
    width: 100%;
    margin-bottom: 0.55rem;
}

.echo-user-bubble {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
    background: #111A2B;
    color: #FFFFFF !important;
    border: 1px solid #D4AF37;
    padding: 0.35rem 0.65rem;
    border-radius: 10px 2px 10px 10px;
    max-width: 75%;
    font-size: 0.80rem;
    line-height: 1.4;
    word-break: break-word;
}
.echo-user-bubble p {
    color: #FFFFFF !important;
    margin: 0;
}

.echo-avatar-user {
    width: 18px;
    height: 18px;
    border-radius: 50%;
    background: #111A2B;
    border: 1px solid #D4AF37;
    display: flex;
    align-items: center;
    justify-content: center;
    flex-shrink: 0;
}

.echo-msg-row-assistant {
    display: flex;
    flex-direction: column;
    width: 100%;
    margin-bottom: 0.65rem;
    background: transparent;
}

.echo-assistant-header {
    display: flex;
    align-items: center;
    gap: 5px;
    margin-bottom: 0.12rem;
}

.echo-avatar-assistant {
    width: 16px;
    height: 16px;
    border-radius: 50%;
    background: #111A2B;
    border: 1px solid #D4AF37;
    display: flex;
    align-items: center;
    justify-content: center;
    flex-shrink: 0;
}

.echo-assistant-title {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
    font-size: 0.72rem;
    font-weight: 600;
    color: #1A2B4C;
}

.echo-assistant-badge-gold {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
    font-size: 0.52rem;
    padding: 1px 4px;
    border-radius: 2px;
    background: #FEF3C7;
    color: #92400E;
    font-weight: 600;
    border: 0.5px solid #FDE68A;
}

.echo-assistant-body {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
    padding-left: 21px;
    color: #374151;
    font-size: 0.80rem;
    line-height: 1.45;
}
.echo-assistant-body strong {
    color: #111827;
}

.echo-sources-container {
    display: flex;
    flex-wrap: wrap;
    gap: 4px;
    margin-top: 0.3rem;
    padding-left: 21px;
}
.echo-source-pill {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
    display: inline-flex;
    align-items: center;
    background: #111A2B;
    border: 1px solid #D4AF37;
    border-radius: 12px;
    padding: 1px 6px;
    font-size: 0.65rem;
    color: #D4AF37 !important;
    text-decoration: none !important;
    font-weight: 500;
    transition: all 0.2s ease;
}
.echo-source-pill:hover {
    border-color: #F1C40F;
    color: #FFFFFF !important;
    box-shadow: 0 0 6px rgba(212, 175, 55, 0.3);
}

.echo-assistant-body table {
    width: 100%;
    border-collapse: collapse;
    margin: 0.35rem 0;
    font-size: 0.75rem;
    background: #FFFFFF;
    border-radius: 4px;
    overflow: hidden;
    border: 1px solid #E2E8F0;
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
}
.echo-assistant-body th {
    background: #111A2B;
    color: #D4AF37;
    font-weight: 600;
    border: 1px solid #334155;
    padding: 3px 6px;
    text-align: left;
}
.echo-assistant-body td {
    border: 1px solid #E2E8F0;
    padding: 3px 6px;
    color: #374151;
}

.echo-thinking-wrapper {
    display: flex;
    align-items: center;
    gap: 5px;
    padding-left: 0.2rem;
    margin-bottom: 0.4rem;
}
.echo-thinking-pill {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
    display: inline-flex;
    align-items: center;
    gap: 4px;
    padding: 2px 6px;
    border-radius: 12px;
    background: #F8FAFC;
    border: 1px solid rgba(212, 175, 55, 0.35);
    font-size: 0.68rem;
    color: #854D0E;
    font-weight: 500;
}
.echo-pulse-dot {
    width: 4px;
    height: 4px;
    background-color: #D4AF37;
    border-radius: 50%;
    animation: echo-pulse 1.4s infinite ease-in-out both;
}
@keyframes echo-pulse {
    0%, 80%, 100% { transform: scale(0); opacity: 0.2; }
    40% { transform: scale(1); opacity: 1; }
}

.echo-context-candidate-card {
    background: #FFFFFF;
    border: 1px solid #D4AF37;
    border-radius: 6px;
    padding: 0.4rem 0.6rem;
    margin-top: 0.3rem;
    box-shadow: 0 2px 5px rgba(0, 0, 0, 0.04);
}

div[data-testid="stHorizontalBlock"]:has(.echo-input-col-target) {
    align-items: center !important;
    margin-top: 0.35rem !important;
    gap: 8px !important;
}

.echo-input-col-target div[data-testid="stChatInput"] {
    width: 100% !important;
    margin: 0 !important;
}

.echo-input-col-target div[data-testid="stChatInput"] > div {
    background: #FFFFFF !important;
    border: 1px solid rgba(212, 175, 55, 0.55) !important;
    border-radius: 20px !important;
    box-shadow: 0 2px 5px rgba(0, 0, 0, 0.03) !important;
    padding: 2px 8px !important;
    min-height: 40px !important;
}

.echo-input-col-target div[data-testid="stChatInput"] textarea {
    font-family: 'Inter', -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, sans-serif !important;
    color: #0F172A !important;
    font-size: 0.85rem !important;
}

.echo-input-col-target div[data-testid="stChatInput"] button {
    background-color: #a3acd5 !important;
    color: #0D1B3E !important;
    border-radius: 50% !important;
    transition: all 0.2s ease-in-out !important;
}

.echo-input-col-target div[data-testid="stChatInput"] button:hover {
    background-color: #8fa2d6 !important;
    color: #0D1B3E !important;
    box-shadow: none !important;
}

.echo-attach-col-target div[data-testid="stPopover"] {
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    height: 100% !important;
}

.echo-attach-col-target div[data-testid="stPopover"] > button {
    width: 40px !important;
    height: 40px !important;
    min-height: 40px !important;
    max-height: 40px !important;
    border-radius: 50% !important;
    padding: 0 !important;
    display: flex !important;
    align-items: center !important;
    justify-content: center !important;
    font-size: 1.1rem !important;
    background: #111A2B !important;
    border: 1px solid #D4AF37 !important;
    color: #D4AF37 !important;
    box-shadow: 0 2px 5px rgba(0, 0, 0, 0.03) !important;
    margin: 0 !important;
}

.echo-attach-col-target div[data-testid="stPopover"] > button:hover {
    border-color: #F1C40F !important;
    box-shadow: 0 0 8px rgba(212, 175, 55, 0.4) !important;
}

.echo-attached-tags {
    display: flex;
    flex-wrap: wrap;
    gap: 4px;
    margin-bottom: 4px;
    padding: 0 4px;
}

.echo-attached-tag {
    font-size: 0.68rem;
    color: #854D0E;
    background: #FEF3C7;
    border: 1px solid #FDE68A;
    border-radius: 10px;
    padding: 1px 7px;
    display: inline-flex;
    align-items: center;
}
</style>
"""

def _extract_text_from_pdf(uploaded_file) -> str:
    try:
        reader = PdfReader(uploaded_file)
        text_content = []
        for i, page in enumerate(reader.pages):
            page_text = page.extract_text()
            if page_text:
                text_content.append(f"--- Page {i+1} ---\n{page_text}")
        return "\n\n".join(text_content)
    except Exception as e:
        st.error(f"Failed to read PDF file: {e}")
        return ""

def _extract_text_from_docx(uploaded_file) -> str:
    try:
        doc = docx.Document(uploaded_file)
        full_text = []
        for p in doc.paragraphs:
            if p.text.strip():
                full_text.append(p.text.strip())
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text))
        return "\n".join(full_text)
    except Exception as e:
        st.error(f"Failed to read DOCX document: {e}")
        return ""

def _encode_image_to_base64(uploaded_file) -> tuple:
    try:
        image = Image.open(uploaded_file)
        fmt = image.format.lower() if image.format else "jpeg"
        if fmt == "jpg":
            fmt = "jpeg"
        buffered = io.BytesIO()
        image.save(buffered, format=image.format if image.format else "JPEG")
        img_str = base64.b64encode(buffered.getvalue()).decode("utf-8")
        return f"data:image/{fmt};base64,{img_str}", image
    except Exception as e:
        st.error(f"Failed to process image: {e}")
        return None, None

def _get_existing_knowledge_map() -> dict:
    try:
        data = fetch_echo_context()
        knowledge_map = {}
        for cat, val in data.items():
            cat_clean = str(cat).lower().strip()
            if isinstance(val, dict):
                for k, v in val.items():
                    val_str = json.dumps(v) if isinstance(v, (dict, list)) else str(v)
                    knowledge_map[(cat_clean, str(k).lower().strip())] = (str(k).strip(), val_str)
            elif isinstance(val, list):
                for item in val:
                    item_str = json.dumps(item) if isinstance(item, (dict, list)) else str(item)
                    knowledge_map[(cat_clean, item_str.lower().strip())] = (item_str.strip(), item_str)
        return knowledge_map
    except Exception:
        return {}

def _safe_upsert_and_verify(category: str, key: str, value: any, priority: int) -> tuple[bool, str]:
    try:
        c_clean = str(category).strip().lower()
        k_clean = str(key).strip()
        p_clean = int(priority) if pd.notna(priority) else 2

        # Standardize structured values for Supabase write
        if isinstance(value, (dict, list)):
            v_payload = json.dumps(value)
        else:
            v_str = str(value).strip()
            if (v_str.startswith("{") and v_str.endswith("}")) or (v_str.startswith("[") and v_str.endswith("]")):
                try:
                    parsed = json.loads(v_str)
                    v_payload = json.dumps(parsed)
                except Exception:
                    v_payload = v_str
            else:
                v_payload = v_str

        write_success = upsert_echo_context(
            category=c_clean,
            key=k_clean,
            value=v_payload,
            priority=p_clean
        )

        if not write_success:
            return False, f"Database write returned False for key: '{k_clean}'"

        return True, ""
    except Exception as e:
        return False, f"Exception during write of '{key}': {str(e)}"

def _check_duplicates_against_db(candidate_df: pd.DataFrame) -> tuple:
    existing_map = _get_existing_knowledge_map()
    clean_rows = []
    conflicts = []

    for idx, row in candidate_df.iterrows():
        if pd.isna(row.get('category')) or pd.isna(row.get('key')) or pd.isna(row.get('value')):
            continue
            
        cat = str(row['category']).strip()
        key = str(row['key']).strip()
        val = str(row['value']).strip()
        prio = int(row['priority']) if pd.notna(row.get('priority')) else 2
        lookup_key = (cat.lower(), key.lower())

        if lookup_key in existing_map:
            orig_key, orig_val = existing_map[lookup_key]
            conflicts.append({
                "category": cat,
                "key": key,
                "current_value": orig_val,
                "new_value": val,
                "priority": prio,
                "resolution": "Keep Current"
            })
        else:
            clean_rows.append({
                "category": cat,
                "key": key,
                "value": val,
                "priority": prio
            })

    return clean_rows, conflicts

@st.dialog("Echo Context Manager", width="large")
def render_context_popup_dialog():
    if "extracted_context_df" not in st.session_state:
        st.session_state["extracted_context_df"] = None
    if "detected_conflicts" not in st.session_state:
        st.session_state["detected_conflicts"] = None
    if "clean_staged_rows" not in st.session_state:
        st.session_state["clean_staged_rows"] = []

    if st.session_state["detected_conflicts"] is not None and len(st.session_state["detected_conflicts"]) > 0:
        st.markdown(f"<p style='font-size:0.95rem; font-weight:600; color:#854D0E;'>{SVG_ALERT_ICON} Duplicate Entries Flagged in Knowledge Base</p>", unsafe_allow_html=True)
        st.caption("The following items already exist in the database with different or identical values. Choose how each key should be resolved:")

        b_c1, b_c2, _ = st.columns([1, 1, 2])
        with b_c1:
            if st.button("Set All to Overwrite", key="btn_all_overwrite", use_container_width=True):
                for item in st.session_state["detected_conflicts"]:
                    item["resolution"] = "Overwrite with New"
                st.rerun()
        with b_c2:
            if st.button("Set All to Keep Current", key="btn_all_keep", use_container_width=True):
                for item in st.session_state["detected_conflicts"]:
                    item["resolution"] = "Keep Current"
                st.rerun()

        for idx, item in enumerate(st.session_state["detected_conflicts"]):
            with st.container(border=True):
                c_lbl, c_res = st.columns([3, 1.5])
                with c_lbl:
                    st.markdown(f"**Key:** `{item['key']}` | **Category:** `{item['category']}`")
                    v_col1, v_col2 = st.columns(2)
                    with v_col1:
                        st.markdown("<span style='font-size:0.75rem; color:#64748B;'>Current Value:</span>", unsafe_allow_html=True)
                        st.code(item['current_value'][:200] + ("..." if len(item['current_value']) > 200 else ""), language="json")
                    with v_col2:
                        st.markdown("<span style='font-size:0.75rem; color:#854D0E;'>New Incoming Value:</span>", unsafe_allow_html=True)
                        st.code(item['new_value'][:200] + ("..." if len(item['new_value']) > 200 else ""), language="json")
                with c_res:
                    res_choice = st.radio(
                        "Action",
                        options=["Keep Current", "Overwrite with New", "Skip Item"],
                        index=["Keep Current", "Overwrite with New", "Skip Item"].index(item["resolution"]),
                        key=f"res_radio_{idx}"
                    )
                    item["resolution"] = res_choice

        btn_act1, btn_act2 = st.columns(2)
        with btn_act1:
            if st.button("Confirm Resolution & Save to DB", type="primary", use_container_width=True):
                saved_count = 0
                error_messages = []
                
                with st.spinner("Writing and actively verifying records..."):
                    for row in st.session_state["clean_staged_rows"]:
                        success, err = _safe_upsert_and_verify(row['category'], row['key'], row['value'], row['priority'])
                        if success:
                            saved_count += 1
                        else:
                            error_messages.append(err)
                    
                    for item in st.session_state["detected_conflicts"]:
                        if item["resolution"] == "Overwrite with New":
                            success, err = _safe_upsert_and_verify(item['category'], item['key'], item['new_value'], item['priority'])
                            if success:
                                saved_count += 1
                            else:
                                error_messages.append(err)

                if error_messages:
                    st.error(f"Failed to save {len(error_messages)} record(s):")
                    for err_msg in error_messages:
                        st.caption(f"[Error] {err_msg}")
                        
                if saved_count > 0:
                    st.success(f"Successfully saved and verified {saved_count} record(s).")
                    
                if saved_count > 0 or not error_messages:
                    st.session_state["detected_conflicts"] = None
                    st.session_state["clean_staged_rows"] = []
                    st.session_state["extracted_context_df"] = None
                    st.rerun()

        with btn_act2:
            if st.button("Cancel & Back to Staging", use_container_width=True):
                st.session_state["detected_conflicts"] = None
                st.session_state["clean_staged_rows"] = []
                st.rerun()

        return

    mode = st.radio(
        "Mode",
        options=["Multimodal AI Extraction (Text/PDF/DOCX/Vision)", "Manual Row Entry"],
        horizontal=True,
        label_visibility="collapsed"
    )

    if mode == "Multimodal AI Extraction (Text/PDF/DOCX/Vision)":
        source_type = st.segmented_control(
            "Input Format",
            options=["Text Notes", "Document (PDF / DOCX)", "Image / Vision Scan"],
            default="Text Notes"
        )

        extracted = []
        if source_type == "Text Notes":
            raw_text = st.text_area(
                "Raw Context / Knowledge Dump",
                height=130,
                placeholder="Paste raw unstructured notes, tables, specs, or logs here..."
            )
            c1, c2 = st.columns([1.5, 1])
            with c1:
                if st.button("Extract Knowledge", key="btn_run_text_ext", type="primary", use_container_width=True):
                    if raw_text.strip():
                        with st.spinner("Extracting structured records..."):
                            extracted = _extract_context_with_ai(raw_text=raw_text)
                    else:
                        st.warning("Please supply context notes.")
            with c2:
                if st.button("Reset Staged Table", key="btn_rst_text_ext", use_container_width=True):
                    st.session_state["extracted_context_df"] = None
                    st.rerun()

        elif source_type == "Document (PDF / DOCX)":
            doc_file = st.file_uploader("Upload PDF or Word Document", type=["pdf", "docx", "doc"], key="dlg_doc_uploader")
            c1, c2 = st.columns([1.5, 1])
            with c1:
                if st.button("Parse & Extract Document", key="btn_run_doc_ext", type="primary", use_container_width=True):
                    if doc_file is not None:
                        with st.spinner("Reading and structuring document content..."):
                            if doc_file.name.lower().endswith(".pdf"):
                                doc_text = _extract_text_from_pdf(doc_file)
                            else:
                                doc_text = _extract_text_from_docx(doc_file)

                            if doc_text.strip():
                                extracted = _extract_context_with_ai(raw_text=doc_text)
                            else:
                                st.error("No extractable text stream found in document.")
                    else:
                        st.warning("Please upload a document file first.")
            with c2:
                if st.button("Reset Staged Table", key="btn_rst_doc_ext", use_container_width=True):
                    st.session_state["extracted_context_df"] = None
                    st.rerun()

        elif source_type == "Image / Vision Scan":
            img_file = st.file_uploader("Upload Image/Scan/Blueprint", type=["png", "jpg", "jpeg", "webp"], key="dlg_img_uploader")
            if img_file:
                st.image(img_file, caption="Scan Preview", use_container_width=True)
            c1, c2 = st.columns([1.5, 1])
            with c1:
                if st.button("Scan with Vision AI", key="btn_run_vision_ext", type="primary", use_container_width=True):
                    if img_file is not None:
                        with st.spinner("Analyzing image visual data with AI..."):
                            img_data_url, _ = _encode_image_to_base64(img_file)
                            if img_data_url:
                                extracted = _extract_context_with_ai(image_data_url=img_data_url)
                    else:
                        st.warning("Please upload an image.")
            with c2:
                if st.button("Reset Staged Table", key="btn_rst_vis_ext", use_container_width=True):
                    st.session_state["extracted_context_df"] = None
                    st.rerun()

        if extracted:
            st.session_state["extracted_context_df"] = pd.DataFrame(extracted)
            st.rerun()

    else:
        m_col1, m_col2, m_col3, m_col4, m_col5 = st.columns([1.2, 1.5, 2.5, 0.8, 1])
        with m_col1:
            m_cat = st.selectbox("Category", options=["knowledge", "team", "jargon", "projects"], key="dlg_manual_cat")
        with m_col2:
            m_key = st.text_input("Key / Entity Identifier", placeholder="e.g. SRC-01162026-001 or Topic", key="dlg_manual_key")
        with m_col3:
            m_val = st.text_input("Value / Structured JSON", placeholder='e.g. {"property": "Sct. Gandia", "rate": 249}', key="dlg_manual_val")
        with m_col4:
            m_prio = st.number_input("Priority", min_value=1, max_value=5, value=2, key="dlg_manual_prio")
        with m_col5:
            st.write("<div style='height: 28px;'></div>", unsafe_allow_html=True)
            if st.button("Add Row", key="btn_dlg_add_row", use_container_width=True):
                if m_key.strip() and m_val.strip():
                    new_entry = pd.DataFrame([{
                        "category": m_cat,
                        "key": m_key.strip(),
                        "value": m_val.strip(),
                        "priority": int(m_prio)
                    }])
                    if st.session_state["extracted_context_df"] is None:
                        st.session_state["extracted_context_df"] = new_entry
                    else:
                        st.session_state["extracted_context_df"] = pd.concat(
                            [st.session_state["extracted_context_df"], new_entry], ignore_index=True
                        )
                    st.rerun()
                else:
                    st.error("Key and Value required.")

    if st.session_state["extracted_context_df"] is not None and not st.session_state["extracted_context_df"].empty:
        st.markdown("---")
        st.markdown("<p style='font-size:0.80rem; font-weight:600; color:#1A2B4C;'>Staged Knowledge Rows (JSON Structured Values)</p>", unsafe_allow_html=True)

        column_config = {
            "category": st.column_config.SelectboxColumn("Category", options=["knowledge", "team", "jargon", "projects"], required=True),
            "key": st.column_config.TextColumn("Key / Entity Name", required=True),
            "value": st.column_config.TextColumn("Value (JSON / String)", required=True, width="large"),
            "priority": st.column_config.NumberColumn("Priority (1-5)", min_value=1, max_value=5, default=2)
        }

        edited_df = st.data_editor(
            st.session_state["extracted_context_df"],
            column_config=column_config,
            num_rows="dynamic",
            use_container_width=True,
            hide_index=True,
            key="dlg_data_editor"
        )

        if st.button("Validate & Commit to Knowledge Base", key="btn_dlg_commit", type="primary", use_container_width=True):
            clean_rows, conflicts = _check_duplicates_against_db(edited_df)
            
            if conflicts:
                st.session_state["detected_conflicts"] = conflicts
                st.session_state["clean_staged_rows"] = clean_rows
                st.rerun()
            else:
                saved = 0
                failed = 0
                with st.spinner("Saving and actively verifying records..."):
                    for row in clean_rows:
                        success, err = _safe_upsert_and_verify(
                            category=row['category'], 
                            key=row['key'], 
                            value=row['value'], 
                            priority=row.get('priority', 2)
                        )
                        if success:
                            saved += 1
                        else:
                            failed += 1
                
                if failed > 0:
                    st.error(f"Failed to save {failed} record(s). Check DB connection logs.")
                if saved > 0:
                    st.success(f"Successfully saved and verified {saved} new record(s).")
                    st.session_state["extracted_context_df"] = None
                    st.rerun()


def _handle_inline_file_upload():
    new_files = st.session_state.get("echo_chat_dock_uploader", [])
    if new_files:
        st.session_state["echo_ui_uploaded_files"] = new_files


def _get_context_lookup_corpus() -> dict[str, str]:
    corpus = {}
    try:
        context_data = fetch_echo_context()
        
        # 1. Team Members
        for member in context_data.get("team", []):
            member_str = json.dumps(member) if isinstance(member, (dict, list)) else str(member)
            if len(member_str.strip()) > 2:
                corpus[member_str.strip()] = f"{member_str.strip()} (Team Member)"
                for part in member_str.strip().split():
                    if len(part) > 2:
                        corpus[part] = f"{member_str.strip()} (Team Member)"

        # 2. Technical Jargon & Acronyms
        for term in context_data.get("jargon", {}).keys():
            if len(str(term).strip()) > 1:
                corpus[str(term).strip()] = f"{str(term).strip()} (Jargon/Acronym)"

        # 3. Active Projects
        for proj in context_data.get("projects", []):
            proj_str = json.dumps(proj) if isinstance(proj, (dict, list)) else str(proj)
            if len(proj_str.strip()) > 2:
                corpus[proj_str.strip()] = f"{proj_str.strip()} (Project)"

        # 4. Enterprise Knowledge
        for key in context_data.get("knowledge", {}).keys():
            if len(str(key).strip()) > 2:
                corpus[str(key).strip()] = f"{str(key).strip()} (Knowledge Entity)"
    except Exception:
        pass
    return corpus


def _fuzzy_align_query(question: str, threshold: int = 82) -> tuple[str, list[str]]:
    corpus = _get_context_lookup_corpus()
    if not corpus or not question.strip():
        return question, []

    known_keys = list(corpus.keys())
    tokens = re.findall(r"\b[A-Za-z0-9\-\./_]+\b", question)
    
    candidates = []
    n = len(tokens)
    for length in [3, 2, 1]:
        for i in range(n - length + 1):
            phrase = " ".join(tokens[i : i + length])
            if len(phrase) >= 3:
                candidates.append(phrase)

    matched_replacements = []
    seen_targets = set()

    for candidate in candidates:
        if candidate.lower() in {"the", "and", "for", "with", "what", "where", "who", "when", "how", "list", "show"}:
            continue

        match = process.extractOne(
            candidate,
            known_keys,
            scorer=fuzz.token_sort_ratio,
            score_cutoff=threshold
        )

        if match:
            best_term, score, _ = match
            canonical_repr = corpus[best_term]

            if canonical_repr not in seen_targets:
                if candidate.lower() != best_term.lower():
                    matched_replacements.append(f"'{candidate}' -> '{best_term}' ({score:.0f}% match)")
                seen_targets.add(canonical_repr)

    if matched_replacements:
        context_hints = "\n".join([f"- {r}" for r in matched_replacements])
        enhanced_prompt = (
            f"{question}\n\n"
            f"[SYSTEM NOTE: Fuzzy entity alignment resolved these intended names/terms from Knowledge Base:\n"
            f"{context_hints}]"
        )
        return enhanced_prompt, matched_replacements

    return question, []


def render_echo_chat(container=None, height=650, title="Ask Echo", caption=None, subtitle=None):
    """
    Renders the unified Echo executive chat workspace with full multi-format file attachment support.
    """
    target = container if container else st
    st.markdown(CHAT_COMPACT_ALIGNED_CSS, unsafe_allow_html=True)

    if "global_chat_history" not in st.session_state:
        st.session_state["global_chat_history"] = []
    if "echo_selected_model_label" not in st.session_state:
        st.session_state["echo_selected_model_label"] = "⚡ Fast - deepseek chat"
    if "echo_source_archives" not in st.session_state:
        st.session_state["echo_source_archives"] = True
    if "echo_source_knowledge" not in st.session_state:
        st.session_state["echo_source_knowledge"] = True
    if "echo_source_web" not in st.session_state:
        st.session_state["echo_source_web"] = False
    if "knowledge_proposal" not in st.session_state:
        st.session_state["knowledge_proposal"] = None
    if "echo_ui_uploaded_files" not in st.session_state:
        st.session_state["echo_ui_uploaded_files"] = []
    if "echo_agent_mode" not in st.session_state:
        st.session_state["echo_agent_mode"] = False

    safe_scroll_height = max(400, int(height) - 130) if height else 520

    with target.container(border=True):
        st.markdown('<div class="echo-main-card-scope"></div>', unsafe_allow_html=True)

        h_left, h_right = st.columns([0.86, 0.14])
        with h_left:
            st.markdown(
                f'<div class="echo-header-bar">'
                f'{SVG_ECHO_LOGO}<span class="echo-title">{title}</span>'
                f'</div>',
                unsafe_allow_html=True
            )
            if subtitle:
                st.caption(subtitle)
            if caption:
                st.caption(caption)

        with h_right:
            st.markdown('<div class="echo-top-controls">', unsafe_allow_html=True)
            c_settings, c_clr = st.columns(2)
            with c_settings:
                with st.popover("⚙", help="Settings"):
                    st.markdown("<span style='font-size:0.75rem; font-weight:600; color:#854D0E;'>AI MODEL</span>", unsafe_allow_html=True)
                    
                    model_options = [
                        "⚡ Fast - deepseek chat",
                        "🧠 Thinking - deepseek reasoning",
                        "👁 Vision - qwen vl"
                    ]
                    current_model = st.session_state.get("echo_selected_model_label", "⚡ Fast - deepseek chat")
                    if current_model not in model_options:
                        current_model = "⚡ Fast - deepseek chat"
                        
                    st.session_state["echo_selected_model_label"] = st.selectbox(
                        "Model",
                        options=model_options,
                        index=model_options.index(current_model),
                        label_visibility="collapsed"
                    )
                    
                    st.markdown("---")
                    st.markdown("<span style='font-size:0.75rem; font-weight:600; color:#854D0E;'>DATA SOURCES</span>", unsafe_allow_html=True)
                    st.session_state["echo_source_archives"] = st.checkbox("Meeting Archives", value=st.session_state["echo_source_archives"])
                    st.session_state["echo_source_knowledge"] = st.checkbox("Echo Knowledge Base", value=st.session_state["echo_source_knowledge"])
                    st.session_state["echo_source_web"] = st.checkbox("Search Web", value=st.session_state["echo_source_web"])

                    st.markdown("---")
                    st.markdown("<span style='font-size:0.75rem; font-weight:600; color:#854D0E;'>AGENT MODE</span>", unsafe_allow_html=True)
                    _agent_allowed = check_agent_access(_cur_user_id())
                    if _agent_allowed:
                        st.session_state["echo_agent_mode"] = st.checkbox(
                            "Enable Agent mode",
                            value=st.session_state.get("echo_agent_mode", False),
                            help="Let Echo take actions (create/edit tasks, log entries, add knowledge) with your approval.",
                        )
                    else:
                        st.session_state["echo_agent_mode"] = False
                        st.markdown(
                            "<span style='font-size:0.75rem; color:#666;'>Agent mode is locked. "
                            "Please contact the developer for your access.</span>",
                            unsafe_allow_html=True,
                        )
                    if st.session_state.get("echo_agent_mode", False):
                        st.caption(f"{len(TOOLS)} actions available · writes require your approval")

                    st.markdown("---")
                    st.markdown("<span style='font-size:0.75rem; font-weight:600; color:#854D0E;'>TOKEN USAGE</span>", unsafe_allow_html=True)
                    _tb = token_balance(_cur_user_id())
                    st.markdown(
                        f"<div style='font-size:0.78rem; line-height:1.6; color:#333;'>"
                        f"<div>Today: <b>{_tb['day_used']:,}</b> / {_tb['day_limit']:,} tokens "
                        f"<span style='color:#666;'>({_tb['day_chats_remaining']} chats left)</span></div>"
                        f"<div>This week: <b>{_tb['week_used']:,}</b> / {_tb['week_limit']:,} tokens "
                        f"<span style='color:#666;'>({_tb['week_chats_remaining']} chats left)</span></div>"
                        f"</div>",
                        unsafe_allow_html=True,
                    )

                    st.markdown("---")
                    st.markdown("<span style='font-size:0.75rem; font-weight:600; color:#854D0E;'>KNOWLEDGE MANAGEMENT</span>", unsafe_allow_html=True)
                    if st.button("Open Context Manager", key="btn_trigger_context_dialog", use_container_width=True):
                        render_context_popup_dialog()

            with c_clr:
                if st.button("🗑", key="btn_clear_global_chat", help="Reset conversation"):
                    st.session_state["global_chat_history"] = []
                    st.session_state["knowledge_proposal"] = None
                    st.session_state["echo_ui_uploaded_files"] = []
                    st.rerun()
            st.markdown('</div>', unsafe_allow_html=True)

        st.markdown('<div class="echo-chat-box-container">', unsafe_allow_html=True)
        chat_box = st.container(height=safe_scroll_height)
        st.markdown('</div>', unsafe_allow_html=True)

        with chat_box:
            if not st.session_state["global_chat_history"]:
                st.markdown(
                    '<div class="echo-msg-row-assistant">'
                    '<div class="echo-assistant-header">'
                    f'<div class="echo-avatar-assistant">{SVG_ECHO_LOGO}</div>'
                    '<span class="echo-assistant-title">Echo</span>'
                    '<span class="echo-assistant-badge-gold">AI</span>'
                    '</div>'
                    '<div class="echo-assistant-body">'
                    'Hi Team, this is Echo, ask me anything...'
                    '</div>'
                    '</div>',
                    unsafe_allow_html=True
                )
            else:
                for msg in st.session_state["global_chat_history"]:
                    if msg["role"] == "user":
                        st.markdown(
                            f'<div class="echo-msg-row-user">'
                            f'<div class="echo-user-bubble">{msg["content"]}</div>'
                            f'<div class="echo-avatar-user">{SVG_USER_ICON}</div>'
                            f'</div>',
                            unsafe_allow_html=True
                        )
                    else:
                        st.markdown(
                            '<div class="echo-msg-row-assistant">'
                            '<div class="echo-assistant-header">'
                            f'<div class="echo-avatar-assistant">{SVG_ECHO_LOGO}</div>'
                            '<span class="echo-assistant-title">Echo</span>'
                            '<span class="echo-assistant-badge-gold">AI</span>'
                            '</div>'
                            '<div class="echo-assistant-body">',
                            unsafe_allow_html=True
                        )
                        st.markdown(msg["content"])
                        st.markdown('</div>', unsafe_allow_html=True)
                        
                        if msg.get("sources"):
                            sources_html = '<div class="echo-sources-container">'
                            for src in msg["sources"]:
                                sources_html += f'<a href="{src["url"]}" target="_blank" class="echo-source-pill">{SVG_GLOBE_ICON}{src["title"]}</a>'
                            sources_html += '</div>'
                            st.markdown(sources_html, unsafe_allow_html=True)
                            
                        st.markdown('</div>', unsafe_allow_html=True)

        if st.session_state["knowledge_proposal"]:
            prop = st.session_state["knowledge_proposal"]
            val_display = str(prop.get("value", ""))
            if len(val_display) > 180:
                val_display = val_display[:180] + "..."
                
            with st.container():
                st.markdown(
                    f'<div class="echo-context-candidate-card">'
                    f'<div style="font-size:0.75rem; font-weight:600; color:#854D0E; margin-bottom:2px;">'
                    f'{SVG_BRAIN_ICON} Knowledge Base Candidate Identified'
                    f'</div>'
                    f'<div style="font-size:0.78rem; color:#1F2937; margin-bottom:6px;">'
                    f'Save <b>{prop.get("key")}</b> [<i>{prop.get("category")}</i>] to Knowledge Base?<br/>'
                    f'<code style="font-size:0.72rem; color:#1A2B4C;">{val_display}</code>'
                    f'</div></div>',
                    unsafe_allow_html=True
                )
                kp_col1, kp_col2 = st.columns([0.5, 0.5])
                with kp_col1:
                    if st.button("Save to Knowledge Base", key="btn_confirm_auto_prop", use_container_width=True):
                        existing_map = _get_existing_knowledge_map()
                        cat_clean = str(prop["category"]).strip().lower()
                        key_clean = str(prop["key"]).strip()
                        
                        if (cat_clean, key_clean.lower()) in existing_map:
                            st.warning(f"Key `{key_clean}` already exists in `{cat_clean}`. Open Context Manager to review overwrites.")
                        else:
                            success, err = _safe_upsert_and_verify(
                                category=prop["category"],
                                key=key_clean,
                                value=prop["value"],
                                priority=prop.get("priority", 2)
                            )
                            if success:
                                st.session_state["global_chat_history"].append({
                                    "role": "assistant",
                                    "content": f"Confirmed: `{key_clean}` verified and saved to Echo Knowledge Base."
                                })
                            else:
                                st.session_state["global_chat_history"].append({
                                    "role": "assistant",
                                    "content": f"Error: Failed to register `{key_clean}`. Detail: {err}"
                                })
                                
                            st.session_state["knowledge_proposal"] = None
                            st.rerun()
                with kp_col2:
                    if st.button("Dismiss", key="btn_dismiss_auto_prop", use_container_width=True):
                        st.session_state["knowledge_proposal"] = None
                        st.rerun()

        # Display staged attachments tags above input bar
        if st.session_state["echo_ui_uploaded_files"]:
            tags_html = '<div class="echo-attached-tags">'
            for f in st.session_state["echo_ui_uploaded_files"]:
                tags_html += f'<span class="echo-attached-tag">{SVG_FILE_ICON} {f.name}</span>'
            tags_html += '</div>'
            st.markdown(tags_html, unsafe_allow_html=True)

        # Single Row: Chat Input & Circular Attachment Trigger
        input_col, attach_col = st.columns([0.94, 0.06])

        with input_col:
            st.markdown('<div class="echo-input-col-target">', unsafe_allow_html=True)
            active_prompt = st.chat_input("Ask Echo...")
            st.markdown('</div>', unsafe_allow_html=True)

        with attach_col:
            st.markdown('<div class="echo-attach-col-target">', unsafe_allow_html=True)
            with st.popover("📎", help="Attach Documents / Scans"):
                st.markdown("<span style='font-size:0.80rem; font-weight:600; color:#1A2B4C;'>Upload Attachments</span>", unsafe_allow_html=True)
                st.file_uploader(
                    "Upload files",
                    type=ALLOWED_ATTACHMENT_TYPES,
                    accept_multiple_files=True,
                    key="echo_chat_dock_uploader",
                    label_visibility="collapsed",
                    on_change=_handle_inline_file_upload
                )
            st.markdown('</div>', unsafe_allow_html=True)

        if active_prompt:
            # Enforce per-user rate limit before spending tokens
            _uid = _cur_user_id()
            _rl = check_rate_limit(_uid)
            if not _rl["allowed"]:
                with chat_box:
                    st.markdown(
                        f'<div class="echo-msg-row-assistant"><div class="echo-assistant-header">'
                        f'<div class="echo-avatar-assistant">{SVG_ECHO_LOGO}</div>'
                        f'<span class="echo-assistant-title">Echo</span></div>'
                        f'<div class="echo-assistant-body">{_rl["why"]}</div></div>',
                        unsafe_allow_html=True,
                    )
                st.rerun()

            attached_files_list = st.session_state["echo_ui_uploaded_files"] if st.session_state["echo_ui_uploaded_files"] else []
            prompt_content_display = active_prompt
            if attached_files_list:
                file_names = ", ".join([f.name for f in attached_files_list])
                prompt_content_display = f"{SVG_FILE_ICON} [{file_names}]<br/>{active_prompt}"

            st.session_state["global_chat_history"].append({"role": "user", "content": prompt_content_display})

            with chat_box:
                st.markdown(
                    f'<div class="echo-msg-row-user">'
                    f'<div class="echo-user-bubble">{prompt_content_display}</div>'
                    f'<div class="echo-avatar-user">{SVG_USER_ICON}</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )
                thinking_placeholder = st.empty()
                status_text = "Searching the web..." if st.session_state["echo_source_web"] else "Thinking..."
                thinking_placeholder.markdown(
                    f'<div class="echo-thinking-wrapper">'
                    f'<div class="echo-avatar-assistant">{SVG_ECHO_LOGO}</div>'
                    f'<div class="echo-thinking-pill">'
                    f'<div class="echo-pulse-dot"></div> {status_text}'
                    f'</div>'
                    f'</div>',
                    unsafe_allow_html=True
                )

            archives = fetch_meeting_archives(limit=100) if st.session_state["echo_source_archives"] else []
            web_context, web_sources = _perform_web_search(active_prompt) if st.session_state["echo_source_web"] else ("", [])

            selected_label = st.session_state.get("echo_selected_model_label", "⚡ Fast - deepseek chat")
            default_model = MODEL_REGISTRY.get(selected_label, "deepseek/deepseek-chat")

            if attached_files_list:
                target_model = "qwen/qwen2.5-vl-72b-instruct"
            else:
                target_model = default_model

            answer, proposed_fact = _query_echo_backend(
                question=active_prompt,
                archive_records=archives,
                chat_history=st.session_state["global_chat_history"],
                web_context=web_context,
                model_name=target_model,
                include_knowledge=st.session_state["echo_source_knowledge"],
                uploaded_files=attached_files_list,
                user_context=build_user_deliverables_context(_cur_username(), archives),
            )
            
            st.session_state["echo_ui_uploaded_files"] = []
            
            thinking_placeholder.empty()
            # Record usage for telemetry + token balance (estimate tokens)
            _est_tokens = (len(active_prompt or "") + len(answer or "")) // 4
            record_usage(_uid, _est_tokens, action="chat")

            st.session_state["global_chat_history"].append({
                "role": "assistant",
                "content": answer,
                "sources": web_sources
            })
            if proposed_fact:
                st.session_state["knowledge_proposal"] = proposed_fact

            st.rerun()


def _perform_web_search(query: str) -> tuple:
    sources = []
    text_snippets = []
    try:
        url = "https://html.duckduckgo.com/html/"
        headers = {"User-Agent": "Mozilla/5.0 (Windows NT 10.0; Win64; x64)"}
        resp = requests.post(url, data={"q": query}, headers=headers, timeout=10)
        if resp.status_code == 200:
            titles = re.findall(r'<a class="result__url"[^>]*>(.*?)</a>', resp.text)
            snippets = re.findall(r'<a class="result__snippet[^>]*>(.*?)</a>', resp.text, re.DOTALL)
            urls = re.findall(r'<a class="result__url[^>]*href="([^"]+)"', resp.text)
            
            for i in range(min(4, len(snippets))):
                clean_snippet = re.sub(r'<.*?>', '', snippets[i]).strip()
                link = urls[i] if i < len(urls) else "#"
                raw_title = re.sub(r'<.*?>', '', titles[i]).strip() if i < len(titles) else f"Source {i+1}"
                
                domain = re.sub(r'^https?://(www\.)?', '', link).split('/')[0]
                pill_title = domain if domain else raw_title[:20]

                sources.append({"title": pill_title, "url": link})
                text_snippets.append(f"[{i+1}] {clean_snippet} (URL: {link})")
    except Exception:
        pass
    return ("\n".join(text_snippets), sources)


def _extract_context_with_ai(raw_text: str = "", image_data_url: str = None) -> list:
    system_prompt = load_prompt("data_extractor")

    if image_data_url:
        api_key = str(st.secrets.get("OPENROUTER_API_KEY", st.secrets.get("OPENAI_API_KEY", ""))).strip()
        if not api_key:
            st.error("API Key is required for vision scanning.")
            return []

        url = "https://openrouter.ai/api/v1/chat/completions" if "OPENROUTER_API_KEY" in st.secrets else "https://api.openai.com/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": "qwen/qwen2.5-vl-72b-instruct" if "OPENROUTER_API_KEY" in st.secrets else "gpt-4o-mini",
            "messages": [
                {"role": "system", "content": system_prompt},
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": "Extract all structured knowledge and data records from this image."},
                        {"type": "image_url", "image_url": {"url": image_data_url}}
                    ]
                }
            ],
            "response_format": {"type": "json_object"},
            "temperature": 0.1,
            "max_tokens": 4000
        }
    else:
        api_key = str(st.secrets.get("DEEPSEEK_API_KEY", st.secrets.get("OPENROUTER_API_KEY", ""))).strip()
        if not api_key:
            st.error("API Key configuration missing.")
            return []

        url = "https://api.deepseek.com/chat/completions" if "DEEPSEEK_API_KEY" in st.secrets else "https://openrouter.ai/api/v1/chat/completions"
        headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
        payload = {
            "model": "deepseek-chat" if "DEEPSEEK_API_KEY" in st.secrets else "deepseek/deepseek-chat",
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": raw_text[:20000]}
            ],
            "response_format": {"type": "json_object"},
            "temperature": 0.1,
            "max_tokens": 8000
        }

    try:
        resp = requests.post(url, headers=headers, json=payload, timeout=90)
        if resp.status_code == 200:
            raw_content = resp.json()["choices"][0]["message"]["content"].strip()
            cleaned = re.sub(r"^```json\s*", "", raw_content, flags=re.MULTILINE)
            cleaned = re.sub(r"^```\s*", "", cleaned, flags=re.MULTILINE).strip()

            try:
                parsed = json.loads(cleaned)
                if isinstance(parsed, dict) and "items" in parsed:
                    return _normalize_extracted_items(parsed["items"])
                elif isinstance(parsed, list):
                    return _normalize_extracted_items(parsed)
            except json.JSONDecodeError:
                match = re.search(r"(\{.*\})", cleaned, re.DOTALL)
                if match:
                    try:
                        parsed = json.loads(match.group(1))
                        return _normalize_extracted_items(parsed.get("items", []))
                    except Exception:
                        pass
                
                item_matches = re.findall(
                    r'\{\s*"category"\s*:\s*"([^"]+)"\s*,\s*"key"\s*:\s*"([^"]+)"\s*,\s*"value"\s*:\s*(?:\"(.*?)\"|(\{.*?\}))\s*(?:,\s*"priority"\s*:\s*(\d+))?\s*\}',
                    cleaned,
                    re.DOTALL
                )
                if item_matches:
                    fallback_items = []
                    for cat, key, val_str, val_obj, prio in item_matches:
                        val = val_str if val_str else val_obj
                        fallback_items.append({
                            "category": cat,
                            "key": key,
                            "value": val,
                            "priority": int(prio) if prio else 2
                        })
                    return fallback_items

        st.error(f"Extraction service error ({resp.status_code}): {resp.text}")
        return []
    except Exception as e:
        st.error(f"Extraction error: {e}")
        return []


def _normalize_extracted_items(items: list) -> list:
    normalized = []
    for item in items:
        if isinstance(item, dict) and "key" in item and "value" in item:
            val = item["value"]
            if isinstance(val, (dict, list)):
                val = json.dumps(val)
            normalized.append({
                "category": str(item.get("category", "knowledge")),
                "key": str(item["key"]),
                "value": str(val),
                "priority": int(item.get("priority", 2))
            })
    return normalized


def _query_echo_backend(
    question: str, 
    archive_records: list, 
    chat_history: list, 
    web_context: str = "",
    model_name: str = "deepseek/deepseek-chat",
    include_knowledge: bool = True,
    uploaded_files: list = None,
    user_context: str = "",
) -> tuple:
    # 1. Fuzzy Context Alignment using Knowledge Base Corpus
    aligned_question, corrections = _fuzzy_align_query(question)

    openrouter_key = str(st.secrets.get("OPENROUTER_API_KEY", "")).strip()
    deepseek_key = str(st.secrets.get("DEEPSEEK_API_KEY", "")).strip()

    if "qwen" in model_name or "openrouter" in model_name or ("/" in model_name and openrouter_key):
        api_key = openrouter_key if openrouter_key else deepseek_key
        api_url = "https://openrouter.ai/api/v1/chat/completions"
    else:
        api_key = deepseek_key if deepseek_key else openrouter_key
        api_url = "https://api.deepseek.com/chat/completions" if deepseek_key else "https://openrouter.ai/api/v1/chat/completions"

    if not api_key:
        return "API Key configuration missing. Please verify Streamlit Secrets.", None

    headers = {"Authorization": f"Bearer {api_key}", "Content-Type": "application/json"}
    archive_context = json.dumps(archive_records, indent=1) if archive_records else "[]"

    if include_knowledge:
        context_data = fetch_echo_context()
        team_list = ", ".join([json.dumps(m) if isinstance(m, (dict, list)) else str(m) for m in context_data.get('team', [])])
        jargon_list = "\n".join([f"- {k}: {json.dumps(v) if isinstance(v, (dict, list)) else v}" for k, v in context_data.get('jargon', {}).items()])
        projects = ", ".join([json.dumps(p) if isinstance(p, (dict, list)) else str(p) for p in context_data.get('projects', [])])
        knowledge_list = "\n".join([f"- {k}: {json.dumps(v) if isinstance(v, (dict, list)) else v}" for k, v in context_data.get('knowledge', {}).items()])

        knowledge_section = f"""
ECHO KNOWLEDGE BASE (SOURCE OF TRUTH):
---------------------------------------
TEAM MEMBERS: {team_list}
ACTIVE PROJECTS: {projects}
TECHNICAL JARGON:
{jargon_list}
ENTERPRISE KNOWLEDGE / STRUCTURED ENTITIES:
{knowledge_list}
"""
    else:
        knowledge_section = ""

    current_date_str = datetime.now().strftime("%A, %B %d, %Y")
    web_section = f"\nLIVE WEB SEARCH RESULTS:\n{web_context}\n" if web_context else ""

    context_string = f"""
CURRENT DATE & TIME: {current_date_str}
{knowledge_section}
{web_section}
{user_context}
"""

    citation_rule = (
        "Incorporate web facts smoothly into the response. Link structures are managed by the UI pills."
        if web_context else ""
    )

    system_prompt = load_prompt(
        "echo_analyst",
        current_date=current_date_str,
        citation_rule=citation_rule,
        context=context_string,
    )

    messages = [{"role": "system", "content": f"{system_prompt}\n\nMeeting Archives:\n{archive_context[:24000]}"}]
    for msg in chat_history[-6:]:
        messages.append({"role": msg["role"], "content": msg["content"]})

    if uploaded_files:
        user_content_blocks = [{"type": "text", "text": aligned_question}]
        for f in uploaded_files:
            if f.type.startswith("image/"):
                img_url, _ = _encode_image_to_base64(f)
                if img_url:
                    user_content_blocks.append({"type": "image_url", "image_url": {"url": img_url}})
            elif f.type == "application/pdf" or f.name.lower().endswith(".pdf"):
                pdf_extracted = _extract_text_from_pdf(f)
                user_content_blocks.append({"type": "text", "text": f"\n\n[PDF Attachment Content ({f.name})]:\n{pdf_extracted}"})
            elif (
                f.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document" 
                or f.type == "application/msword" 
                or f.name.lower().endswith((".docx", ".doc"))
            ):
                docx_extracted = _extract_text_from_docx(f)
                user_content_blocks.append({"type": "text", "text": f"\n\n[DOCX Attachment Content ({f.name})]:\n{docx_extracted}"})
            else:
                try:
                    f_text = f.read().decode("utf-8")
                    user_content_blocks.append({"type": "text", "text": f"\n\n[Attachment Content ({f.name})]:\n{f_text}"})
                except Exception:
                    pass
        messages.append({"role": "user", "content": user_content_blocks})
    else:
        messages.append({"role": "user", "content": aligned_question})

    payload = {
        "model": model_name.replace("deepseek/", "") if "api.deepseek.com" in api_url else model_name,
        "messages": messages,
        "temperature": 0.2,
        "max_tokens": 2000
    }

    if "reasoner" not in model_name and "vl" not in model_name:
        payload["response_format"] = {"type": "json_object"}

    try:
        resp = requests.post(api_url, headers=headers, json=payload, timeout=60)
        if resp.status_code == 200:
            raw_content = resp.json()["choices"][0]["message"]["content"].strip()
            cleaned = re.sub(r"^```json\s*", "", raw_content, flags=re.MULTILINE)
            cleaned = re.sub(r"^```\s*", "", cleaned, flags=re.MULTILINE).strip()

            try:
                result = json.loads(cleaned)
                return result.get("response", cleaned), result.get("propose_knowledge")
            except json.JSONDecodeError:
                match = re.search(r"(\{.*\})", cleaned, re.DOTALL)
                if match:
                    try:
                        result = json.loads(match.group(1))
                        return result.get("response", cleaned), result.get("propose_knowledge")
                    except Exception:
                        pass
                return raw_content, None

        return f"Service notice ({resp.status_code}): {resp.text}", None
    except Exception as e:
        return f"Analysis exception: {e}", None
